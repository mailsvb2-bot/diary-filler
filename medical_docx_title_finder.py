from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from medical_constants import DATE_FMT
from medical_text_utils import normalize_match, normalize_text
from medical_docx_blocks import iter_block_items
from medical_docx_date_patterns import _TITLE_DATE_RE, _first_valid_full_date, _normalize_full_date_match
from medical_docx_title_context import (
    _date_match_has_birth_context,
    _is_primary_title_context,
    _is_strong_birth_date_context,
)
from medical_docx_xml_fragments import _docx_xml_text_fragments


def _best_title_date_in_text(value: str) -> str:
    """Pick the date nearest to the document title in one text fragment."""
    text = normalize_text(value or "")
    if not text or not _is_primary_title_context(text):
        return ""
    title_positions = []
    for marker in ("первичный осмотр", "направление на госпитализацию", "госпитализацию"):
        idx = normalize_match(text).find(marker)
        if idx >= 0:
            title_positions.append(idx)
    if not title_positions:
        return ""
    matches: list[tuple[int, str]] = []
    for match in _TITLE_DATE_RE.finditer(text):
        normalized = _normalize_full_date_match(match)
        if not normalized:
            continue
        if _date_match_has_birth_context(text, match):
            continue
        distance = min(abs(match.start() - pos) for pos in title_positions)
        # Заголовок и дата должны быть рядом. Широкий поиск запрещён: иначе
        # можно опять поймать дату рождения из карточки пациента.
        if distance <= 120:
            matches.append((distance, normalized))
    if not matches:
        return ""
    return sorted(matches, key=lambda item: item[0])[0][1]

def _admission_date_from_filename(path: str | Path) -> str:
    """Return admission date from file/folder name when date is written next to document name.

    This is intentionally strict: it accepts only a date placed close to
    "Первичный осмотр" or "Направление на госпитализацию" in the filename/path
    and never scans generic patient/date fragments as a fallback.
    """
    try:
        candidates = [Path(path).stem, Path(path).name]
    except Exception:
        candidates = [str(path)]
    for value in candidates:
        value = normalize_text(str(value))
        if not value:
            continue
        found = _best_title_date_in_text(value)
        if found:
            return found
        # Частый вариант: файл называется просто "12.01.2026 Первичный осмотр.docx".
        low = normalize_match(value)
        title_match = re.search(r"первичн\w*\s+осмотр|направлени[ея]\s+на\s+госпитализац\w+|госпитализац\w+", low)
        if title_match:
            window_start = max(0, title_match.start() - 80)
            window_end = min(len(value), title_match.end() + 80)
            window = value[window_start:window_end]
            for match in _TITLE_DATE_RE.finditer(window):
                normalized = _normalize_full_date_match(match)
                if normalized and not _date_match_has_birth_context(window, match):
                    return normalized
    return ""

def extract_admission_date_from_title_docx(path: str | Path) -> str:
    """Return admission date only from the document title/header area.

    Контракт: дата поступления — это дата рядом с названием документа
    (например, ``12.01.2026 Первичный осмотр``). Дата рядом с ФИО,
    датой рождения, годом рождения или возрастом не используется.

    Важно: если дата написана в имени загружаемого файла рядом с названием
    документа, это тоже считается заголовком документа.
    """
    filename_date = _admission_date_from_filename(path)
    if filename_date:
        return filename_date

    structured_entries: list[str] = []
    row_entries: list[list[str]] = []

    try:
        doc = Document(str(path))
    except Exception:
        doc = None

    def add(value: str) -> None:
        value = normalize_text(value or "")
        if value:
            structured_entries.append(value)

    def collect_cell_text(cell: _Cell) -> str:
        parts: list[str] = []
        for paragraph in cell.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text.strip())
        for table in cell.tables:
            for row in table.rows:
                nested_parts: list[str] = []
                seen_nested: set[int] = set()
                for nested_cell in row.cells:
                    tc_id = id(nested_cell._tc)
                    if tc_id in seen_nested:
                        continue
                    seen_nested.add(tc_id)
                    nested_text = collect_cell_text(nested_cell)
                    if nested_text:
                        nested_parts.append(nested_text)
                if nested_parts:
                    parts.append(" | ".join(nested_parts))
        return normalize_text(" ".join(parts))

    if doc is not None:
        for block in iter_block_items(doc):
            if len(structured_entries) >= 120:
                break
            if isinstance(block, Paragraph):
                add(block.text)
            elif isinstance(block, Table):
                for row in block.rows:
                    cells: list[str] = []
                    seen_cells: set[int] = set()
                    for cell in row.cells:
                        tc_id = id(cell._tc)
                        if tc_id in seen_cells:
                            continue
                        seen_cells.add(tc_id)
                        cell_text = collect_cell_text(cell)
                        if cell_text:
                            cells.append(cell_text)
                            add(cell_text)
                    if cells:
                        row_entries.append(cells)
                        # Отдельно добавляем всю строку: иногда Word отдаёт
                        # дату и заголовок через разные ячейки одной строки.
                        add(" | ".join(cells))

    # 1) Самый надёжный вариант: дата и заголовок в одном фрагменте.
    for entry in structured_entries[:120]:
        value = _best_title_date_in_text(entry)
        if value:
            return value

    # 2) Таблица: дата и заголовок в соседних ячейках одной строки.
    for cells in row_entries[:60]:
        title_indexes = [idx for idx, cell in enumerate(cells) if _is_primary_title_context(cell)]
        if not title_indexes:
            continue
        for title_idx in title_indexes:
            title_cell = cells[title_idx]
            direct = _best_title_date_in_text(title_cell)
            if direct:
                return direct
            for idx, cell in sorted(enumerate(cells), key=lambda item: abs(item[0] - title_idx)):
                if idx == title_idx:
                    continue
                if _is_strong_birth_date_context(cell):
                    continue
                for match in _TITLE_DATE_RE.finditer(cell):
                    if _date_match_has_birth_context(cell, match):
                        continue
                    normalized = _normalize_full_date_match(match)
                    if normalized:
                        return normalized

    # 3) Соседние фрагменты: дата отдельной строкой, заголовок отдельной строкой.
    # Поддерживаем те же форматы, что и общий парсер: 12.01.2026,
    # 12 . 01 . 26, 120126, 12012026 и короткие формы вроде 1126.
    date_only_re = re.compile(
        r"^\s*(?:\d{1,2}\s*[./-]\s*\d{1,2}\s*[./-]\s*\d{2,4}|\d{4,8})(?:\s+\d{1,2}:\d{2})?\s*$"
    )
    for idx, entry in enumerate(structured_entries[:120]):
        if not date_only_re.match(entry):
            continue
        if _is_strong_birth_date_context(entry):
            continue
        neighbors = structured_entries[max(0, idx - 2): min(len(structured_entries), idx + 3)]
        if any(_is_strong_birth_date_context(item) for item in neighbors):
            continue
        if any(_is_primary_title_context(item) for item in neighbors):
            value = _first_valid_full_date(entry)
            if value:
                return value

    # 4) Raw XML fallback: headers/textboxes/shapes. Это закрывает случаи,
    # когда python-docx не видит текст заголовка.
    xml_fragments = _docx_xml_text_fragments(path)
    for entry in xml_fragments[:180]:
        value = _best_title_date_in_text(entry)
        if value:
            return value
    if xml_fragments:
        joined = normalize_text(" | ".join(xml_fragments[:180]))
        # Ищем дату строго рядом с заголовком, не шире 120 символов.
        title_re = re.compile(r"первичн\w*\s+осмотр|направлени[ея]\s+на\s+госпитализац\w+", re.IGNORECASE)
        for title_match in title_re.finditer(joined):
            window_start = max(0, title_match.start() - 120)
            window_end = min(len(joined), title_match.end() + 120)
            window = joined[window_start:window_end]
            value = _best_title_date_in_text(window)
            if value:
                return value

    return ""
