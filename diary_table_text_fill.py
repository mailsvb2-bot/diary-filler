"""Разделённый слой заполнителя дневников.

Файл создан при архитектурной нарезке бывшего diary_filler.py.
"""

from __future__ import annotations

import re
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from diary_constants import HOLIDAY_SKIP_END_DAY, HOLIDAY_SKIP_MONTHS, HOLIDAY_SKIP_START_DAY, STATUS_FONT_SIZE_PT, STRUCTURAL_DIARY_PREFIXES
from diary_dates import add_month, format_month_year, parse_month_year, safe_row_date
from diary_text_parser import is_signature_paragraph_text, normalize_text, remove_examinee_words
from diary_table_cells import clear_paragraph_keep_properties, first_signature_paragraph_index, is_structural_diary_prefix, reset_cell_to_one_paragraph, write_diary_text_into_existing_paragraph

def fill_diary_text_cell(cell, diary_text: str, keep_signature: bool = True) -> None:
    """Fill only the diary area of a cell, preserving template signatures.

    Signatures are not extracted, deleted, re-created, moved, or reformatted.
    If the template row contains a structural note such as "Совместный осмотр...",
    that note remains above the generated diary text.
    """
    diary_text = remove_examinee_words(diary_text)
    _ = keep_signature  # kept for API compatibility; template signatures are always preserved when recognizable.

    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
        write_diary_text_into_existing_paragraph(paragraph, diary_text)
        return

    signature_index = first_signature_paragraph_index(cell)

    # No signature block: safe old strategy, because there is no template signature layout to preserve.
    if signature_index is None:
        paragraph = reset_cell_to_one_paragraph(cell)
        write_diary_text_into_existing_paragraph(paragraph, diary_text)
        return

    paragraphs = list(cell.paragraphs)
    before_signature = paragraphs[:signature_index]

    preserved_prefix_indexes = {
        index
        for index, paragraph in enumerate(before_signature)
        if is_structural_diary_prefix(paragraph.text)
    }

    search_from = max(preserved_prefix_indexes) + 1 if preserved_prefix_indexes else 0
    target_index: int | None = None

    for index in range(search_from, signature_index):
        if index in preserved_prefix_indexes:
            continue
        if not normalize_text(paragraphs[index].text):
            target_index = index
            break

    if target_index is None:
        for index in range(search_from, signature_index):
            if index not in preserved_prefix_indexes:
                target_index = index
                break

    if target_index is None:
        signature_paragraph = paragraphs[signature_index]
        target_paragraph = signature_paragraph.insert_paragraph_before("")
    else:
        target_paragraph = paragraphs[target_index]

    for index, paragraph in enumerate(before_signature):
        if index in preserved_prefix_indexes:
            continue
        if target_index is not None and index == target_index:
            continue
        if normalize_text(paragraph.text):
            clear_paragraph_keep_properties(paragraph)

    write_diary_text_into_existing_paragraph(target_paragraph, diary_text)
