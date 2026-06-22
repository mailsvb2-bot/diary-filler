from __future__ import annotations

from typing import Iterable, Sequence

from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from medical_text_utils import normalize_match


def paragraph_matches_marker(normalized_paragraph_text: str, marker: str) -> bool:
    marker = normalize_match(marker)
    if not marker or not normalized_paragraph_text:
        return False
    text = normalized_paragraph_text.lstrip()
    if text.startswith(marker):
        if len(text) == len(marker):
            return True
        # Важно: короткие маркеры вроде «ЭПИ» не должны срабатывать на «ЭПИКРИЗ».
        # Разрешаем совпадение только если после маркера идёт разделитель.
        next_char = text[len(marker)]
        return not (next_char.isalnum() or next_char == "_")
    if marker == "зарегистрирован по адресу" and " зарегистрирован по адресу" in text:
        return True
    return False

def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    try:
        paragraph.clear()
    except AttributeError:
        for run in list(paragraph.runs):
            run._element.getparent().remove(run._element)
    paragraph.add_run(text)

def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para

def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)

def iter_all_paragraphs(parent) -> Iterable[Paragraph]:
    """Все абзацы документа, включая абзацы внутри таблиц."""
    if isinstance(parent, DocxDocument):
        for paragraph in parent.paragraphs:
            yield paragraph
        for table in parent.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_all_paragraphs(cell)
    elif isinstance(parent, _Cell):
        for paragraph in parent.paragraphs:
            yield paragraph
        for table in parent.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_all_paragraphs(cell)

def remove_exact_paragraphs(doc: DocxDocument, values: Sequence[str]) -> int:
    """Remove standalone service paragraphs such as a trailing 'ЭЭГ'."""
    normalized_values = {normalize_match(value) for value in values}
    count = 0
    for paragraph in list(iter_all_paragraphs(doc)):
        if normalize_match(paragraph.text) in normalized_values:
            remove_paragraph(paragraph)
            count += 1
    return count
