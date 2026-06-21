from __future__ import annotations

import re
from typing import Sequence

from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from medical_text_utils import normalize_match, normalize_text
from medical_docx_editor_utils import insert_paragraph_after, paragraph_matches_marker, remove_paragraph, set_paragraph_text


class DocxEditorReplaceMixin:
    def replace_first_matching_paragraph(self, markers: Sequence[str], text: str) -> bool:
        idx = self.find_paragraph_index(markers)
        if idx is None:
            return False
        set_paragraph_text(self.paragraphs[idx], text)
        return True

    def replace_all_matching_paragraphs(self, markers: Sequence[str], text: str) -> int:
        count = 0
        for paragraph in self.paragraphs:
            text_norm = normalize_match(paragraph.text)
            if any(paragraph_matches_marker(text_norm, marker) for marker in markers):
                set_paragraph_text(paragraph, text)
                count += 1
        return count

    def insert_before_first_matching_paragraph(self, markers: Sequence[str], text: str) -> bool:
        idx = self.find_paragraph_index(markers)
        if idx is None:
            return False
        target = self.paragraphs[idx]
        new_p = OxmlElement("w:p")
        target._p.addprevious(new_p)
        new_para = Paragraph(new_p, target._parent)
        if text:
            new_para.add_run(text)
        return True

    def remove_all_matching_paragraphs(self, markers: Sequence[str]) -> int:
        """Удалить все абзацы, которые начинаются с указанных маркеров.

        Используется для ЭПИ: если файл ЭПИ не выбран, в итоговых документах
        не должно оставаться ни строки «ЭПИ - ...», ни шаблонной подсказки.
        """
        count = 0
        for paragraph in list(self.paragraphs):
            text_norm = normalize_match(paragraph.text)
            if any(paragraph_matches_marker(text_norm, marker) for marker in markers):
                remove_paragraph(paragraph)
                count += 1
        return count

    def replace_first_matching_regex(self, pattern: str, text: str) -> bool:
        rx = re.compile(pattern, flags=re.IGNORECASE)
        for paragraph in self.paragraphs:
            if rx.search(normalize_text(paragraph.text)):
                set_paragraph_text(paragraph, text)
                return True
        return False

    def replace_block(
        self,
        start_markers: Sequence[str],
        label: str,
        value: str,
        all_markers: Sequence[str],
        *,
        allow_empty: bool = False,
        preserve_when_empty: bool = True,
    ) -> bool:
        value = normalize_text(value)
        if not value and preserve_when_empty and not allow_empty:
            return False

        idx = self.find_paragraph_index(start_markers)
        if idx is None:
            return False

        paragraphs = self.paragraphs
        end_idx = self.find_next_marker_index(idx + 1, all_markers, exclude=start_markers)
        if end_idx is None:
            end_idx = idx + 1

        for remove_idx in range(end_idx - 1, idx, -1):
            remove_paragraph(paragraphs[remove_idx])

        lines = [line.strip() for line in value.splitlines() if line.strip()]
        if not lines:
            lines = [""]

        first = f"{label} {lines[0]}".rstrip()
        set_paragraph_text(paragraphs[idx], first)

        anchor = paragraphs[idx]
        for line in lines[1:]:
            anchor = insert_paragraph_after(anchor, line)
        return True
