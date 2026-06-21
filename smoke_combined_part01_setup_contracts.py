from pathlib import Path
import shutil
from itertools import combinations

from docx import Document

import medical_documents as _medical_documents_module
from medical_documents import MedicalDocumentService, MedicalTextParser, DOCUMENT_ORDER, PatientData, build_expert_anamnesis, extract_docx_text, sanitize_diagnosis, treatment_period_text, parse_date, format_military_commissariat_area, format_date_with_russian_year_suffix, format_birth_for_person_line
assert not hasattr(_medical_documents_module, "MedicalApp"), "medical_documents.py must not contain old UI class"
from diary_filler import fill_diary_batch, extract_statuses_from_docx, parse_full_date, parse_month_year, safe_filename_part
from icd10_f import search_icd10_f
from medical_docx_reader import _first_valid_full_date

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "test_run_combined"
if OUT.exists():
    shutil.rmtree(OUT)
OUT.mkdir(parents=True, exist_ok=True)

nav = OUT / "Направление_тест.docx"
nav_doc = Document()
nav_doc.add_paragraph("10.06.2026 Первичный осмотр")
nav_doc.add_paragraph("История болезни № 123")
nav_doc.add_paragraph("Ф.И.О.: Иванова Ирина Ивановна")
nav_doc.add_paragraph("Год рождения: 1980")
nav_doc.add_paragraph("Зарегистрирован: Н. Новгород, тестовый район")
nav_doc.add_paragraph("Работает в организации: не работает")
nav_doc.add_paragraph("В 3 отделение КДП поступает добровольно")
nav_doc.add_paragraph("Жалобы на момент осмотра: тревога, нарушение сна")
nav_doc.add_paragraph("Анамнез жизни: Со слов пациентки, росла и развивалась без особенностей.")
nav_doc.add_paragraph("Анамнез заболевания: Ухудшение состояния в течение месяца.")
nav_doc.add_paragraph("Психический статус: В сознании, ориентирована, контакт доступен.")
nav_doc.add_paragraph("Соматический статус: Нормального питания, без грубой соматической патологии.")
nav_doc.add_paragraph("План лечения: терапия по назначению врача")
nav_doc.add_paragraph("На основании данных анамнеза жизни и заболевания, психического статуса, данных клинических исследований был выставлен диагноз: F тестовый диагноз")
nav_doc.save(nav)

epi = OUT / "ЭПИ.docx"
epi_doc = Document()
epi_doc.add_paragraph("ЭПИ: ЭПИ тестовая информация.")
epi_doc.save(epi)


# --- Medical documents smoke with EPI and manual UI-like реквизиты ---
service = MedicalDocumentService()

# --- Two-digit title date must normalize to the 1900/2000 century once, not twice ---
assert _first_valid_full_date("12.01.26") == "12.01.2026"
assert _first_valid_full_date("12.01.76") == "12.01.1976"
assert _first_valid_full_date("10052026 Первичный осмотр") == "10.05.2026"
assert _first_valid_full_date("100526 Первичный осмотр") == "10.05.2026"
assert _first_valid_full_date("1126 Первичный осмотр") == "01.01.2026"

# --- Admission date regression: title date must win over birth date ---
title_date_doc = OUT / "Дата_заголовок_против_рождения.docx"
td = Document()
td.add_paragraph("12.01.2026 Первичный осмотр")
td.add_paragraph("Ф.И.О.: Сидоров Иван Михайлович, Дата рождения: 04.01.2000")
td.add_paragraph("Диагноз: F41.2 Тестовый диагноз")
td.add_paragraph("Жалобы: тест")
td.add_paragraph("Психический статус: тест")
td.save(title_date_doc)
parsed_title_date = service.parse_primary_document(title_date_doc)
assert parsed_title_date.admission_date == "12.01.2026", parsed_title_date.admission_date
assert parsed_title_date.birth in {"04.01.2000", "2000", ""}

table_title_date_doc = OUT / "Дата_заголовок_таблица.docx"
td = Document()
table = td.add_table(rows=1, cols=2)
table.cell(0, 0).text = "13.02.2026"
table.cell(0, 1).text = "Первичный осмотр"
td.add_paragraph("Ф.И.О.: Сидоров Иван Михайлович")
td.add_paragraph("Дата рождения: 04.01.2000")
td.add_paragraph("Диагноз: F41.2 Тестовый диагноз")
td.add_paragraph("Жалобы: тест")
td.add_paragraph("Психический статус: тест")
td.save(table_title_date_doc)
parsed_table_title_date = service.parse_primary_document(table_title_date_doc)
assert parsed_table_title_date.admission_date == "13.02.2026", parsed_table_title_date.admission_date

compact_title_date_doc = OUT / "Дата_заголовок_без_точек.docx"
td = Document()
td.add_paragraph("1126 Первичный осмотр")
td.add_paragraph("Ф.И.О.: Сидоров Иван Михайлович")
td.add_paragraph("Дата рождения: 04.01.2000")
td.add_paragraph("Диагноз: F41.2 Тестовый диагноз")
td.add_paragraph("Жалобы: тест")
td.add_paragraph("Психический статус: тест")
td.save(compact_title_date_doc)
parsed_compact_title_date = service.parse_primary_document(compact_title_date_doc)
assert parsed_compact_title_date.admission_date == "01.01.2026", parsed_compact_title_date.admission_date


# --- UI source regression: sick-leave selector defaults to "нет" and popup has no duplicate yes/no fields ---
import main as _main_module
assert hasattr(_main_module.CombinedMedicalDiaryApp, "_prompt_assigned_treatment_if_needed")

def _project_python_source() -> str:
    return "\n".join(
        path.read_text(encoding="utf-8")
        for path in sorted(ROOT.glob("*.py"))
        if not path.name.startswith(("smoke_test", "smoke_combined_"))
    )

main_source = _project_python_source()
assert 'expert_sick_leave_needed_var = tk.StringVar(value="нет")' in main_source
assert '("Место работы", self.vk_mse_work_org_var.get().strip() or shared_org)' in main_source
assert '("Должность", self.vk_mse_position_var.get().strip() or shared_position)' in main_source
assert '("Место работы", self.sick_leave_vk_work_org_var.get().strip() or shared_org)' in main_source
assert '("Должность", self.sick_leave_vk_position_var.get().strip() or shared_position)' in main_source
assert '("Место работы, должность", self.vk_mse_work_org_var.get().strip())' not in main_source
assert '_apply_primary_work_defaults(data)' in main_source
assert 'self._primary_work_org_default' in main_source
assert 'self._work_details_manually_edited' in main_source
assert '"Сормовский и Московский"' in main_source and '"Канавинский"' in main_source
assert 'kind == "discharge"' in main_source
assert 'def _ensure_discharge_date' in main_source
assert 'title="Дата выписки"' in main_source
assert 'self.output_vars[DIARY_KIND] = tk.BooleanVar(value=False)' in main_source
assert 'admission_doctor_referral' in main_source
assert 'Перетащите сюда первичный осмотр/направление на госпитализацию' in main_source
assert 'text="Нужен больничный лист?"' in main_source
assert 'command=self._on_expert_sick_leave_fill' in main_source
assert 'text="Да"' in main_source
assert 'text="Нет"' in main_source
assert '("Работает? да/нет"' not in main_source
assert '("Нужен больничный лист? да/нет"' not in main_source
assert '("С какого числа больничный"' in main_source
assert '("Где работает / организация"' in main_source
assert '("Должность"' in main_source
assert "Ничего не запоминаем между разными popup-окнами" in main_source
assert "self._last_committee_date,\n            self.vk_date_var.get().strip()" not in main_source
assert "self.vk_date_var.get().strip(),\n            self.sick_leave_vk_commission_date_var.get().strip()" not in main_source
assert 'single_line=self._compact_ui' in main_source
assert 'suffix not in {".docx", ".docm"}' in main_source
assert 'txt_low.strip().startswith("эпи")' in main_source

release_zip_source = (ROOT / "make_release_zip.py").read_text(encoding="utf-8")
for snippet in ['".spec"', '".DS_Store"', '"Thumbs.db"', '".vscode"', '".idea"']:
    assert snippet in release_zip_source, snippet

