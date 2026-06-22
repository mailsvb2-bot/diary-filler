"""Разделённый слой заполнителя дневников.

Файл создан при архитектурной нарезке бывшего diary_filler.py.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from diary_constants import (DATE_PREFIX_RE, EXAMINEE_ANY_RE, EXAMINEE_STANDARD_STATE_RE, EXAMINEE_START_RE, MIN_STATUS_LEN, SIGNATURE_MARKERS, STATUS_DATE_PREFIX_RE, STATUS_LABEL_PREFIX_RE, STATUS_NUMBER_BEFORE_DATE_RE, STATUS_NUMBER_PREFIX_RE, STATUS_STANDALONE_DAY_PREFIX_RE, STRUCTURAL_DIARY_PREFIXES, WHITESPACE_RE)

def normalize_text(text: str) -> str:
    text = (text or "")
    text = re.sub(r"[\u00ad\u200b\u200c\u200d\u2060\ufeff]", "", text)
    text = text.replace("\xa0", " ").replace("\n", " ")
    text = WHITESPACE_RE.sub(" ", text).strip()
    return text


def strip_leading_status_metadata(text: str) -> str:
    value = normalize_text(text)
    patterns = (
        STATUS_LABEL_PREFIX_RE,
        STATUS_NUMBER_BEFORE_DATE_RE,
        STATUS_DATE_PREFIX_RE,
        DATE_PREFIX_RE,
        STATUS_NUMBER_PREFIX_RE,
        STATUS_STANDALONE_DAY_PREFIX_RE,
    )
    for _ in range(12):
        before = value
        for pattern in patterns:
            updated = pattern.sub("", value, count=1).lstrip(" —-–—:.);,]\t")
            if updated != value:
                value = normalize_text(updated)
                break
        if value == before:
            break
    return normalize_text(value)


def remove_examinee_words(text: str) -> str:
    value = normalize_text(text)
    value = EXAMINEE_STANDARD_STATE_RE.sub(lambda match: match.group(1), value)
    for _ in range(3):
        updated = EXAMINEE_START_RE.sub("", value)
        if updated == value:
            break
        value = normalize_text(updated)
    value = EXAMINEE_ANY_RE.sub("", value)
    value = re.sub(r"^\s*[:,.;!\-–—]+\s*", "", value)
    value = re.sub(r"\s+([,.;:!?])", r"\1", value)
    value = re.sub(r"([\(\[\{])\s+", r"\1", value)
    value = re.sub(r"\s+([\)\]\}])", r"\1", value)
    value = re.sub(r"\s{2,}", " ", value)
    return normalize_text(value)


def clean_status_text(text: str) -> str:
    return remove_examinee_words(strip_leading_status_metadata(text))


def is_signature_paragraph_text(text: str) -> bool:
    """Return True for template signature paragraphs that must stay untouched."""
    low = normalize_text(text).lower().replace("ё", "е")
    return any(low.startswith(marker) for marker in SIGNATURE_MARKERS)


def looks_like_status(text: str) -> bool:
    text = clean_status_text(text)
    low = text.lower()
    if len(text) < MIN_STATUS_LEN:
        return False
    if is_signature_paragraph_text(text):
        return False
    if any(low.startswith(prefix) for prefix in STRUCTURAL_DIARY_PREFIXES):
        return False
    if low in {"дневник наблюдения", "день госпитализации", "число", "дата", "месяц/год", "месяц / год"}:
        return False
    if re.fullmatch(r"[\d\s./-]+", text):
        return False
    return True


def extract_statuses_from_docx(path: str | Path) -> list[str]:
    doc = Document(str(path))
    statuses: list[str] = []
    seen_statuses: set[str] = set()

    def add_candidate(text: str) -> None:
        cleaned = clean_status_text(text)
        key = cleaned.lower().replace("ё", "е")
        if looks_like_status(cleaned) and key not in seen_statuses:
            statuses.append(cleaned)
            seen_statuses.add(key)

    for paragraph in doc.paragraphs:
        add_candidate(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            seen_cells: set[int] = set()
            for cell in row.cells:
                # Merged cells are exposed repeatedly by python-docx; process each
                # physical cell once so one diary text does not consume several rows.
                tc_id = id(cell._tc)
                if tc_id in seen_cells:
                    continue
                seen_cells.add(tc_id)
                for paragraph in cell.paragraphs:
                    add_candidate(paragraph.text)
    return statuses
