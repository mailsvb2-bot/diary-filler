"""Разделённый слой заполнителя дневников.

Файл создан при архитектурной нарезке бывшего diary_filler.py.
"""

from __future__ import annotations

import re
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from diary_constants import HOLIDAY_SKIP_END_DAY, HOLIDAY_SKIP_MONTHS, HOLIDAY_SKIP_START_DAY, STATUS_FONT_SIZE_PT, STRUCTURAL_DIARY_PREFIXES
from diary_dates import add_month, format_month_year, parse_month_year, safe_row_date
from diary_text_parser import is_signature_paragraph_text, normalize_text, remove_examinee_words

def clear_paragraph_keep_properties(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}pPr"):
            continue
        p.remove(child)

def reset_cell_to_one_paragraph(cell):
    if not cell.paragraphs:
        return cell.add_paragraph()
    first = cell.paragraphs[0]
    clear_paragraph_keep_properties(first)
    for paragraph in list(cell.paragraphs[1:]):
        paragraph._element.getparent().remove(paragraph._element)
    return first

def is_structural_diary_prefix(text: str) -> bool:
    """Return True for template notes that should stay above generated diary text."""
    low = normalize_text(text).lower().replace("ё", "е")
    return any(low.startswith(prefix) for prefix in STRUCTURAL_DIARY_PREFIXES)

def first_signature_paragraph_index(cell) -> int | None:
    for index, paragraph in enumerate(cell.paragraphs):
        if is_signature_paragraph_text(paragraph.text):
            return index
    return None

def add_run_with_size(paragraph, text: str):
    run = paragraph.add_run(text)
    run.font.size = Pt(STATUS_FONT_SIZE_PT)
    return run

def fill_text_cell(cell, text: str, *, alignment=None) -> None:
    paragraph = reset_cell_to_one_paragraph(cell)
    if alignment is not None:
        paragraph.alignment = alignment
    add_run_with_size(paragraph, text)

def write_diary_text_into_existing_paragraph(paragraph, diary_text: str) -> None:
    clear_paragraph_keep_properties(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run_with_size(paragraph, diary_text)
