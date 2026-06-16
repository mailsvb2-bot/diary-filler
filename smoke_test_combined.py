import sys
from pathlib import Path as _Path
sys.path.insert(0, str(_Path(__file__).resolve().parent / "src"))
from pathlib import Path
import shutil

from docx import Document

from medical_documents import MedicalDocumentService, DOCUMENT_ORDER, PatientData, extract_docx_text
from diary_filler import fill_diary_batch, extract_statuses_from_docx
from icd10_f import search_icd10_f

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "test_run_combined"
if OUT.exists():
    shutil.rmtree(OUT)
OUT.mkdir(parents=True, exist_ok=True)

nav = OUT / "Направление_тест.docx"
nav_doc = Document()
nav_doc.add_paragraph("10.06.2026 Первичный осмотр")
nav_doc.add_paragraph("История болезни № 123")
nav_doc.add_paragraph("Ф.И.О.: Иванова Ирина Ивановна")
nav_doc.add_paragraph("Год рождения: 1980")
nav_doc.add_paragraph("Зарегистрирован: Н. Новгород, тестовый район")
nav_doc.add_paragraph("Работает в организации: не работает")
nav_doc.add_paragraph("В 3 отделение КДП поступает добровольно")
nav_doc.add_paragraph("Жалобы на момент осмотра: тревога, нарушение сна")
nav_doc.add_paragraph("Анамнез жизни: Со слов пациентки, росла и развивалась без особенностей.")
nav_doc.add_paragraph("Анамнез заболевания: Ухудшение состояния в течение месяца.")
nav_doc.add_paragraph("Психический статус: В сознании, ориентирована, контакт доступен.")
nav_doc.add_paragraph("Соматический статус: Нормального питания, без грубой соматической патологии.")
nav_doc.add_paragraph("План лечения: терапия по назначению врача")
nav_doc.add_paragraph("На основании данных анамнеза жизни и заболевания, психического статуса, данных клинических исследований был выставлен диагноз: F тестовый диагноз")
nav_doc.save(nav)

epi = OUT / "ЭПИ.docx"
epi_doc = Document()
epi_doc.add_paragraph("ЭПИ: ЭПИ тестовая информация.")
epi_doc.save(epi)

# --- Medical documents smoke with EPI and manual UI-like реквизиты ---
service = MedicalDocumentService()
manual_data = service.parse_navigation(nav)
manual_data.discharge_date = "11.06.2026"
manual_data.diagnosis = "F99.9 Тестовый диагноз из UI"
manual_data.rvk_act_number = "77-А"
manual_data.rvk_military_commissariat = "Ленинского"
manual_data.rvk_work_position = "ООО РВК, программист"
manual_data.vk_date = "16.06.2026"
manual_data.vk_protocol_number = "42"
manual_data.vk_protocol_date = "16.06.2026"
manual_data.vk_mse_work_org = "ГБУЗ НО Тест, санитар"
manual_data.vk_mse_position = ""
manual_data.sick_leave_vk_date = "18.06.2026"
manual_data.sick_leave_vk_protocol_number = "55"
manual_data.sick_leave_vk_protocol_date = "18.06.2026"
manual_data.sick_leave_vk_commission_date = "18.06.2026"
manual_data.sick_leave_vk_work_position = "ООО Тест, инженер"
created, data = service.create_documents(
    navigation_path=nav,
    output_dir=OUT / "medical_with_epi",
    discharge_date="11.06.2026",
    epi_path=epi,
    selected_docs=DOCUMENT_ORDER,
    override_data=manual_data,
)
assert len(created) == 6, created
assert all(path.exists() for path in created)
combined_text = "\n".join(extract_docx_text(path) for path in created)
assert "F99.9 Тестовый диагноз из UI" in combined_text
assert "О СОСТОЯНИИ ЗДОРОВЬЯ ГРАЖДАНИНА № 77-А" in combined_text
assert "Выписка из ПРОТОКОЛА № 42" in combined_text
assert "Выписка из ПРОТОКОЛА № 55" in combined_text
assert "Место работы: ГБУЗ НО Тест, санитар" in combined_text
assert "Место работы: ООО РВК, программист" in combined_text
assert "военного комиссариата Ленинского района" in combined_text
assert "Место работы, должность: ООО Тест, инженер" in combined_text
assert "Находится на лечении с 10.06.2026 (9 дней)" in combined_text
assert "От 16.06.2026 г." in combined_text
assert "ЭПИ тестовая информация" in combined_text

# --- Medical documents without EPI: no ЭПИ mentions should remain ---
manual_no_epi = service.parse_navigation(nav)
manual_no_epi.discharge_date = "11.06.2026"
manual_no_epi.diagnosis = "F88 Диагноз без дополнительного исследования"
manual_no_epi.rvk_act_number = "88-Б"
manual_no_epi.rvk_military_commissariat = "Советского района"
manual_no_epi.rvk_work_position = "не работает"
manual_no_epi.vk_date = "17.06.2026"
manual_no_epi.vk_protocol_number = "43"
manual_no_epi.vk_protocol_date = "17.06.2026"
manual_no_epi.vk_mse_work_org = "не работает"
manual_no_epi.vk_mse_position = ""
manual_no_epi.sick_leave_vk_date = "19.06.2026"
manual_no_epi.sick_leave_vk_protocol_number = "56"
manual_no_epi.sick_leave_vk_protocol_date = "19.06.2026"
manual_no_epi.sick_leave_vk_commission_date = "19.06.2026"
manual_no_epi.sick_leave_vk_work_position = "не работает"
created_no_epi, _ = service.create_documents(
    navigation_path=nav,
    output_dir=OUT / "medical_without_epi",
    discharge_date="11.06.2026",
    epi_path=None,
    selected_docs=DOCUMENT_ORDER,
    override_data=manual_no_epi,
)
for path in created_no_epi:
    text = extract_docx_text(path)
    assert "сюда подставлять" not in text.lower(), path
    assert "выбирается в ui" not in text.lower(), path
    assert not __import__("re").search(r"(?<![A-Za-zА-ЯЁа-яё])ЭПИ(?![A-Za-zА-ЯЁа-яё])", text), path

# --- Diary filler smoke ---
source = OUT / "texts.docx"
doc = Document()
doc.add_paragraph("01.06.2026 Пациент был спокоен, жалоб активно не предъявлял, в беседе доступен, инструкции выполнял.")
doc.add_paragraph("02.06.2026 Пациент сообщил об улучшении сна, фон настроения ровный, поведение упорядоченное.")
doc.save(source)
assert len(extract_statuses_from_docx(source)) == 2

table_file = OUT / "diary_table.docx"
doc = Document()
table = doc.add_table(rows=1, cols=4)
headers = ["№", "Число", "Месяц/год", "Дневник наблюдения"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
for day in [10, 11, 12, 13, 14, 15]:
    row = table.add_row()
    row.cells[0].text = str(day)
    row.cells[1].text = str(day)
    row.cells[2].text = ""
    row.cells[3].text = "Лечащий врач Балаганин С.В."
doc.save(table_file)

result = fill_diary_batch(
    status_files=[source],
    diary_files=[table_file],
    output_dir=OUT / "diaries",
    patient_name="Иванова И.И.",
    admission_value="10.06.2026",
    discharge_value="12.06.2026",
    repeat_statuses=True,
    reset_each_file=True,
    keep_signature=True,
    fill_months=True,
    force_final_diary=True,
    remove_holiday_rows=True,
)
assert result.processed_files == 1
assert result.created_files[0].exists()
assert result.filled_rows >= 1
assert result.final_rows_filled == 1
assert result.removed_after_discharge_rows >= 3
diary_text = "\n".join("\t".join(cell.text for cell in row.cells) for row in Document(result.created_files[0]).tables[0].rows)
assert "Пациентка была спокойна" in diary_text
assert "не предъявляла" in diary_text

# --- Diary gender source smoke: UI filename may be male, source document is female ---
result_filename_male = fill_diary_batch(
    status_files=[source],
    diary_files=[table_file],
    output_dir=OUT / "diaries_gender_source",
    patient_name="Иванов Иван Иванович",
    gender_source_name="Иванова Ирина Ивановна",
    admission_value="10.06.2026",
    discharge_value="12.06.2026",
    repeat_statuses=True,
    reset_each_file=True,
    keep_signature=True,
    fill_months=True,
    force_final_diary=True,
    remove_holiday_rows=True,
)
diary_text2 = "\n".join("\t".join(cell.text for cell in row.cells) for row in Document(result_filename_male.created_files[0]).tables[0].rows)
assert "Пациентка была спокойна" in diary_text2
assert "не предъявляла" in diary_text2
assert result_filename_male.created_files[0].name.startswith("Иванов Иван Иванович")

assert any(item.code == "F41.2" for item in search_icd10_f("41"))
assert any(item.code == "F41.2" for item in search_icd10_f("трев деп"))
print("OK")
print("Medical docs with EPI:", len(created))
print("Medical docs without EPI:", len(created_no_epi))
print("Diary files:", len(result.created_files))
print("Output:", OUT)
