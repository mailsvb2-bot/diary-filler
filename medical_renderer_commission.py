from __future__ import annotations

from pathlib import Path

from docx import Document

from medical_constants import TARGET_MEDICAL_FACILITY
from medical_docx_editor import (
    DocxBlockEditor,
    iter_all_paragraphs,
    remove_exact_paragraphs,
    set_paragraph_text,
)
from medical_expert import put_expert_anamnesis
from medical_formatting import (
    format_birth_for_person_line,
    format_date_with_russian_year_suffix,
    format_military_commissariat_area,
    treatment_period_text,
)
from medical_gender import finalize_medical_document
from medical_markers import (
    COMMISSION_MARKERS,
    DISCHARGE_MARKERS,
    PRIMARY_MARKERS,
    RVK_MARKERS,
    SICK_LEAVE_VK_MARKERS,
    VK_MSE_MARKERS,
)
from medical_models import PatientData
from medical_parser_sanitize import sanitize_diagnosis
from medical_text_utils import normalize_match


class MedicalRendererCommissionMixin:
    def render_commission(self, template_path: str | Path, output_path: str | Path, data: PatientData) -> None:
        """Комиссионный: шапку/первую строку оставляем, заполняем клиническую часть ниже."""
        doc = Document(str(template_path))
        editor = DocxBlockEditor(doc)
        dates = data.lab_dates()

        if data.commission_date or data.commission_number:
            commission_date = format_date_with_russian_year_suffix(data.commission_date or data.admission_date)
            header = (
                f"{commission_date} 10:00      "
                f"Совместный осмотр с зам глав врача Зуйковой А.А. № {data.commission_number}"
            ).rstrip()
            editor.replace_first_matching_regex(r"Совместный\s+осмотр", header)

        birth_text = format_birth_for_person_line(data.birth)
        person_line = f"{data.fio}, {birth_text}, зарегистрирован по адресу: {data.registered or 'Н. Новгород'}".strip(" ,")
        editor.replace_first_matching_paragraph(["г.р.,", "зарегистрирован по адресу"], person_line)
        put_expert_anamnesis(editor, data, COMMISSION_MARKERS, ["В 3 отделение КДП поступает"], include_sick_leave_number=False, include_return_to_work=False)
        editor.replace_block(["В 3 отделение КДП поступает"], "В 3 отделение КДП поступает", data.admission, COMMISSION_MARKERS)
        editor.replace_block(["Жалобы при поступлении", "Жалобы"], "Жалобы при поступлении:", data.complaints, COMMISSION_MARKERS)
        editor.replace_block(["Анамнез жизни"], "Анамнез жизни:", data.life_anamnesis, COMMISSION_MARKERS)
        editor.replace_block(["Анамнез заболевания"], "Анамнез заболевания:", data.disease_anamnesis, COMMISSION_MARKERS)
        editor.replace_block(["Психический статус при поступлении", "Психический статус"], "Психический статус при поступлении:", data.mental_status, COMMISSION_MARKERS)
        editor.replace_block(["Соматический статус", "Сомато-неврологический статус"], "Соматический статус:", data.somatic_status, COMMISSION_MARKERS)

        lab_lines = [
            (["ОАК"], f"ОАК ({dates['day1']}) - в норме"),
            (["ОАМ"], f"ОАМ ({dates['day1']}) - в норме"),
            (["RW"], f"RW (от {dates['day1']}) - в норме"),
            (["HCV"], f"HCV (от {dates['day1']}) - в норме"),
            (["HBsAg"], f"HBsAg (от {dates['day1']}) - в норме"),
            (["ВИЧ"], f"ВИЧ (от {dates['day2']}) - в норме"),
            (["Биохимия крови"], f"Биохимия крови ({dates['day1']}) - в норме"),
            (["Глюкоза крови"], f"Глюкоза крови ({dates['day1']}) – 3,40 ммоль/л"),
            (["Кал на яйца глист"], f"Кал на яйца глист ({dates['day1']}) - не обнаружены."),
            (["Флюорография"], f"Флюорография ({dates['flg']}) - патологии не выявлено."),
            (["ЭКГ"], f"ЭКГ ({dates['day1']}) – ритм синусовый, ЭОС нормальная."),
        ]
        for markers, text in lab_lines:
            editor.replace_first_matching_paragraph(markers, text)
        if data.epi_text:
            editor.replace_block(["ЭПИ"], f"ЭПИ ({dates['day2']}) -", data.epi_text, COMMISSION_MARKERS)
        else:
            editor.remove_all_matching_paragraphs(["ЭПИ"])

        diagnosis_sentence = ""
        diagnosis = sanitize_diagnosis(data.diagnosis)
        if diagnosis:
            diagnosis_sentence = (
                "На основании данных анамнеза жизни и заболевания, психического статуса, "
                f"данных клинических исследований установлен диагноз: {diagnosis}"
            )
        editor.replace_block(["На основании данных", "Диагноз"], "", diagnosis_sentence, COMMISSION_MARKERS)
        editor.replace_block(["Лечение"], "Лечение:", data.treatment_plan, COMMISSION_MARKERS)
        editor.replace_block(["Эпидемиологический анамнез"], "Эпидемиологический анамнез:", data.epidemiology, COMMISSION_MARKERS)
        finalize_medical_document(doc, data)
        doc.save(str(output_path))

    def render_admission_doctor_referral(self, template_path: str | Path, output_path: str | Path, data: PatientData) -> None:
        """Осмотр врача приёмного покоя по отдельному шаблону.

        Документ берёт те же клинические данные, что первичный осмотр, но имеет
        собственный заголовок и финальную фразу направления в учреждение.
        """
        doc = Document(str(template_path))
        editor = DocxBlockEditor(doc)

        header_date = data.admission_date or "Дата"
        # В шаблоне дата и заголовок могут быть в одной строке: "12.04.2026 10:00 Осмотр...".
        # Маркер находится не в начале абзаца, поэтому ищем по regex/contains, а не только startswith.
        header_done = editor.replace_first_matching_regex(
            r"Осмотр\s+врача\s+при[её]много\s+покоя",
            f"{header_date} 10:00 Осмотр врача приёмного покоя.",
        )
        if not header_done:
            editor.replace_first_matching_paragraph(["Дата, время"], f"{header_date} 10:00 Осмотр врача приёмного покоя.")
        person_line = f"{data.fio}, {data.birth}, {data.registered}".strip(" ,")
        if person_line.strip(" ,"):
            editor.replace_first_matching_paragraph(["Сидоров", "Ф.И.О.", "ФИО"], person_line)
        editor.replace_block(["На учёте у психиатров", "На учете у психиатров"], "На учёте у психиатров:", data.psych_account, PRIMARY_MARKERS, allow_empty=True)
        editor.replace_block(["Работает в организации"], "Работает в организации:", data.work_org, PRIMARY_MARKERS, allow_empty=True)
        editor.replace_block(["Должность"], "Должность:", data.position, PRIMARY_MARKERS, allow_empty=True)
        editor.replace_block(["Больничный лист"], "Больничный лист:", data.sick_leave, PRIMARY_MARKERS, allow_empty=True)
        editor.replace_block(["Оформление инвалидности"], "Оформление инвалидности:", data.disability, PRIMARY_MARKERS, allow_empty=True)
        editor.replace_block(["Направление от РВК"], "Направление от РВК:", data.rvk_referral, PRIMARY_MARKERS, allow_empty=True)
        editor.replace_block(["В 3 отделение КДП поступает"], "В 3 отделение КДП поступает:", data.admission, PRIMARY_MARKERS, allow_empty=True)
        editor.replace_block(["Жалобы на момент осмотра", "Жалобы"], "Жалобы на момент осмотра:", data.complaints, PRIMARY_MARKERS)
        editor.replace_block(["Анамнез жизни"], "Анамнез жизни:", data.life_anamnesis, PRIMARY_MARKERS)
        editor.replace_block(["Анамнез заболевания"], "Анамнез заболевания:", data.disease_anamnesis, PRIMARY_MARKERS)
        editor.replace_block(["Психический статус"], "Психический статус:", data.mental_status, PRIMARY_MARKERS)
        editor.replace_block(["Соматический статус"], "Соматический статус:", data.somatic_status, PRIMARY_MARKERS)
        editor.replace_block(["План обследования"], "План обследования:", data.examination_plan, PRIMARY_MARKERS)
        diagnosis = sanitize_diagnosis(data.diagnosis)
        diagnosis_sentence = ""
        if diagnosis:
            diagnosis_sentence = (
                "На основании данных анамнеза жизни и заболевания, психического статуса, "
                f"данных клинических исследований был выставлен диагноз: {diagnosis}"
            )
        editor.replace_block(["На основании данных", "Диагноз"], "", diagnosis_sentence, PRIMARY_MARKERS)
        editor.replace_block(["Эпидемиологический анамнез"], "Эпидемиологический анамнез:", data.epidemiology, PRIMARY_MARKERS)
        # Финальная фраза должна быть строго такой по пользовательскому требованию.
        target_referral_line = f"В связи с психическим состоянием, направляется на лечение в {TARGET_MEDICAL_FACILITY}"
        referral_done = False
        for paragraph in list(iter_all_paragraphs(doc)):
            if "направляется" in normalize_match(paragraph.text):
                set_paragraph_text(paragraph, target_referral_line)
                referral_done = True
                break
        if not referral_done:
            referral_done = editor.insert_before_first_matching_paragraph(
                ["Врач психиатр", "Врач-психиатр"],
                target_referral_line,
            )
        if not referral_done:
            doc.add_paragraph(target_referral_line)
        editor.replace_block(["Врач психиатр", "Врач-психиатр"], "Врач психиатр", data.doctor, PRIMARY_MARKERS, allow_empty=True)
        finalize_medical_document(doc, data)
        doc.save(str(output_path))
