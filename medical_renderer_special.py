from __future__ import annotations

from pathlib import Path

from docx import Document

from medical_constants import TARGET_MEDICAL_FACILITY
from medical_docx_editor import (
    DocxBlockEditor,
    iter_all_paragraphs,
    remove_exact_paragraphs,
    set_paragraph_text,
)
from medical_expert import put_expert_anamnesis
from medical_formatting import (
    format_birth_for_person_line,
    format_date_with_russian_year_suffix,
    format_military_commissariat_area,
    treatment_period_text,
)
from medical_gender import finalize_medical_document
from medical_markers import (
    COMMISSION_MARKERS,
    DISCHARGE_MARKERS,
    PRIMARY_MARKERS,
    RVK_MARKERS,
    SICK_LEAVE_VK_MARKERS,
    VK_MSE_MARKERS,
)
from medical_models import PatientData
from medical_parser_sanitize import sanitize_diagnosis
from medical_text_utils import normalize_match


class MedicalRendererSpecialMixin:
    def render_vk_mse(self, template_path: str | Path, output_path: str | Path, data: PatientData) -> None:
        """ВК на МСЭ: верхнюю шапку и текст решения не трогаем, заполняем поля пациента."""
        doc = Document(str(template_path))
        editor = DocxBlockEditor(doc)

        if data.vk_date:
            editor.replace_first_matching_regex(r"^\s*\.?\s*\d{4}\s*$", data.vk_date)
        if data.vk_protocol_number:
            editor.replace_first_matching_paragraph(["Выписка из ПРОТОКОЛА"], f"Выписка из ПРОТОКОЛА № {data.vk_protocol_number}")
        if data.vk_protocol_date:
            editor.replace_first_matching_regex(r"^\s*От\s*\.?\s*\d{4}\s*г\.?\s*$", f"От {data.vk_protocol_date} г.")

        editor.replace_all_matching_paragraphs(["Ф.И.О", "Ф.И.О:"], f"Ф.И.О: {data.fio}")
        editor.replace_all_matching_paragraphs(["Год рождения"], f"Год рождения: {data.birth}")
        editor.replace_all_matching_paragraphs(["Проживает"], f"Проживает: {data.registered or 'Н. Новгород'}")
        vk_work_parts = [
            (data.vk_mse_work_org or data.work_org).strip(),
            (data.vk_mse_position or data.position).strip(),
        ]
        vk_work_line = ", ".join(part for part in vk_work_parts if part) or "не работает"
        editor.replace_all_matching_paragraphs(["Место работы"], f"Место работы: {vk_work_line}")
        editor.replace_all_matching_paragraphs(["Диагноз"], f"Диагноз: {sanitize_diagnosis(data.diagnosis)}")

        editor.replace_block(["Жалобы"], "Жалобы:", data.complaints, VK_MSE_MARKERS)
        editor.replace_block(["Анамнез жизни"], "Анамнез жизни:", data.life_anamnesis, VK_MSE_MARKERS)
        editor.replace_block(["Анамнез заболевания"], "Анамнез заболевания:", data.disease_anamnesis, VK_MSE_MARKERS)
        editor.replace_block(["Психический статус при поступлении", "Психический статус"], "Психический статус при поступлении:", data.mental_status, VK_MSE_MARKERS)
        if data.epi_text:
            editor.replace_block(["ЭПИ"], "ЭПИ -", data.epi_text, VK_MSE_MARKERS)
        else:
            editor.remove_all_matching_paragraphs(["ЭПИ"])
        editor.replace_block(["Сомато-неврологический статус", "Соматический статус"], "Сомато-неврологический статус:", data.somatic_status, VK_MSE_MARKERS)
        editor.replace_block(["Получает лечение"], "Получает лечение:", data.treatment_plan, VK_MSE_MARKERS)
        finalize_medical_document(doc, data)
        doc.save(str(output_path))

    def render_sick_leave_vk(self, template_path: str | Path, output_path: str | Path, data: PatientData) -> None:
        """ВК больничный: отдельная форма ВК для продления лечения/больничного."""
        doc = Document(str(template_path))
        editor = DocxBlockEditor(doc)

        # Верхняя дата, номер протокола и дата протокола работают так же, как в ВК на МСЭ.
        if data.sick_leave_vk_date:
            editor.replace_first_matching_regex(r"^\s*\.?\s*\d{4}\s*$", data.sick_leave_vk_date)
        if data.sick_leave_vk_protocol_number:
            editor.replace_first_matching_paragraph(["Выписка из ПРОТОКОЛА"], f"Выписка из ПРОТОКОЛА № {data.sick_leave_vk_protocol_number}")
        if data.sick_leave_vk_protocol_date:
            editor.replace_first_matching_regex(r"^\s*От\s*\.?\s*\d{4}\s*г\.?\s*$", f"От {data.sick_leave_vk_protocol_date} г.")

        work_position = data.sick_leave_vk_work_position or ", ".join(
            part for part in [data.sick_leave_vk_work_org, data.sick_leave_vk_position] if part
        ).strip(", ") or ", ".join(part for part in [data.work_org, data.position] if part).strip(", ") or "не работает"
        treatment_line = treatment_period_text(data.admission_date, data.sick_leave_vk_commission_date or data.sick_leave_vk_date)

        editor.replace_all_matching_paragraphs(["Ф.И.О", "Ф.И.О:"], f"Ф.И.О: {data.fio}")
        editor.replace_all_matching_paragraphs(["Год рождения"], f"Год рождения: {data.birth}")
        editor.replace_all_matching_paragraphs(["Проживает"], f"Проживает: {data.registered or 'Н. Новгород'}")
        editor.replace_all_matching_paragraphs(["Место работы"], f"Место работы, должность: {work_position}")
        editor.replace_all_matching_paragraphs(["Находится на лечении"], treatment_line)
        editor.replace_all_matching_paragraphs(["Диагноз"], f"Диагноз: {sanitize_diagnosis(data.diagnosis)}")

        editor.replace_block(["Жалобы"], "Жалобы:", data.complaints, SICK_LEAVE_VK_MARKERS)
        editor.replace_block(["Анамнез жизни"], "Анамнез жизни:", data.life_anamnesis, SICK_LEAVE_VK_MARKERS)
        editor.replace_block(["Анамнез заболевания"], "Анамнез заболевания:", data.disease_anamnesis, SICK_LEAVE_VK_MARKERS)
        editor.replace_block(["Психический статус при поступлении", "Психический статус"], "Психический статус при поступлении:", data.mental_status, SICK_LEAVE_VK_MARKERS)
        if data.epi_text:
            editor.replace_block(["ЭПИ"], "ЭПИ -", data.epi_text, SICK_LEAVE_VK_MARKERS)
        else:
            editor.remove_all_matching_paragraphs(["ЭПИ"])
        editor.replace_block(["Сомато-неврологический статус", "Соматический статус"], "Сомато-неврологический статус:", data.somatic_status, SICK_LEAVE_VK_MARKERS)
        editor.replace_block(["Получает лечение"], "Получает лечение:", data.treatment_plan, SICK_LEAVE_VK_MARKERS)
        finalize_medical_document(doc, data)
        doc.save(str(output_path))

    def render_rvk(self, template_path: str | Path, output_path: str | Path, data: PatientData) -> None:
        """Акт для РВК: нормативную шапку оставляем, заполняем поля ниже."""
        doc = Document(str(template_path))
        editor = DocxBlockEditor(doc)
        dates = data.lab_dates()

        act_number = data.rvk_act_number or data.case_number
        editor.replace_first_matching_paragraph(["О СОСТОЯНИИ"], f"О СОСТОЯНИИ ЗДОРОВЬЯ ГРАЖДАНИНА № {act_number}".rstrip())
        editor.replace_block(["История болезни №"], "История болезни №", data.case_number, RVK_MARKERS, preserve_when_empty=False, allow_empty=True)
        editor.replace_block(["Ф.И.О.", "ФИО"], "Ф.И.О.:", data.fio, RVK_MARKERS, allow_empty=True)
        editor.replace_block(["Год рождения"], "Год рождения:", data.birth, RVK_MARKERS, allow_empty=True)
        editor.replace_block(["Проживает"], "Проживает:", data.registered, RVK_MARKERS, allow_empty=True)
        # В Акте для РВК строка "Место работы" не нужна: удаляем её из результата,
        # чтобы туда не попадали данные из направления или старые значения UI.
        editor.remove_all_matching_paragraphs(["Место работы"])
        period = f"Находился на обследовании в {TARGET_MEDICAL_FACILITY} с {data.admission_date} по {data.discharge_date}".strip()
        editor.replace_first_matching_paragraph(["Находился на обследовании"], period)
        military_area = format_military_commissariat_area(data.rvk_military_commissariat)
        if military_area:
            editor.replace_first_matching_paragraph(
                ["Госпитализируется по направлению военного комиссариата"],
                f"Госпитализируется по направлению военного комиссариата {military_area}."
            )
        if data.psych_account:
            editor.replace_first_matching_paragraph(["На учёте", "На учете"], f"На учёте у психиатров {data.psych_account}.")
        editor.replace_block(["Жалобы"], "Жалобы:", data.complaints or "не предъявляет", RVK_MARKERS, allow_empty=True)
        editor.replace_block(["Анамнез жизни"], "Анамнез жизни:", data.life_anamnesis, RVK_MARKERS)
        editor.replace_block(["Анамнез заболевания"], "Анамнез заболевания:", data.disease_anamnesis, RVK_MARKERS)
        editor.replace_block(["Психический статус"], "Психический статус:", data.mental_status, RVK_MARKERS)
        editor.replace_block(["Сомато-неврологический статус", "Соматический статус"], "Сомато-неврологический статус:", data.somatic_status, RVK_MARKERS)
        self._replace_lab_lines(editor, dates)
        if data.epi_text:
            editor.replace_block(["ЭПИ"], "ЭПИ -", data.epi_text, RVK_MARKERS)
        else:
            editor.remove_all_matching_paragraphs(["ЭПИ"])
        # В шаблоне Акта РВК после блока ЭПИ/ЭЭГ есть служебная одиночная строка "ЭЭГ".
        # Она не относится к результату исследования и должна исчезать из итогового документа.
        remove_exact_paragraphs(doc, ["ЭЭГ", "ЭПИ"])
        editor.replace_first_matching_paragraph(["Диагноз"], f"Диагноз: {sanitize_diagnosis(data.diagnosis)}")
        finalize_medical_document(doc, data)
        doc.save(str(output_path))
