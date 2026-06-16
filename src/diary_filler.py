"""
Заполнитель дневников: логика без UI.

Сохранённый функционал:
- выбор одного/нескольких DOCX-файлов с текстами дневников/статусами;
- выбор одного/нескольких DOCX-файлов-таблиц дневников;
- ФИО пациента для имени результата и определения пола;
- дата поступления как ММ.ГГГГ или ДД.ММ.ГГГГ;
- дата выписки как ДД.ММ.ГГГГ;
- заполнение текстов дневников;
- повтор текстов при нехватке;
- сброс/продолжение индекса текстов между несколькими таблицами;
- сохранение подписей «лечащий врач» в ячейках;
- заполнение столбца месяц/год;
- финальная запись выписки в последнюю строку;
- дата выписки всегда попадает в последнюю строку таблицы;
- удаление строк после даты выписки;
- удаление праздничных строк 1-9 января/мая;
- аккуратная замена мужских/женских форм для типовых клинических слов;
- отчёт по обработке.
"""

from __future__ import annotations

import calendar
import os
import re
import shutil
import subprocess
import sys
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


MIN_STATUS_LEN = 25
STATUS_FONT_SIZE_PT = 8
HOLIDAY_SKIP_MONTHS = {1, 5}
HOLIDAY_SKIP_START_DAY = 1
HOLIDAY_SKIP_END_DAY = 9
SIGNATURE_MARKERS = (
    "лечащий врач",
    "зав.отделением",
    "зав. отделением",
    "зав отделением",
)
STRUCTURAL_DIARY_PREFIXES = (
    "совместный осмотр",
)

DATE_PREFIX_RE = re.compile(
    r"^\s*(?:"
    r"\d{1,2}[./-]\d{1,2}[./-]\d{2,4}\s*(?:г\.?|год)?"
    r"|\d{1,2}\s+[а-яё]+\s+\d{2,4}\s*(?:г\.?|год)?"
    r")\s*",
    re.IGNORECASE,
)
WHITESPACE_RE = re.compile(r"[ \t\r\f\v]+")
MONTH_YEAR_RE = re.compile(r"^\s*(\d{1,2})\s*[./-]\s*(\d{4})\s*$")
FULL_DATE_RE = re.compile(r"^\s*(\d{1,2})\s*[./-]\s*(\d{1,2})\s*[./-]\s*(\d{2,4})\s*$")
INVALID_FILENAME_CHARS_RE = re.compile(r'[<>:"/\\|?*\x00-\x1f]+')

STATUS_LABEL_PREFIX_RE = re.compile(
    r"^\s*(?:дата|число|номер|№|n|no\.?|запись|дневник)"
    r"\s*(?:№?\s*\d{1,6})?\s*[:.\-–—]\s*",
    re.IGNORECASE,
)
STATUS_NUMBER_BEFORE_DATE_RE = re.compile(
    r"^\s*\d{1,6}\s+(?=(?:\d{4}[./-]\d{1,2}[./-]\d{1,2}|\d{1,2}[./-]\d{1,2}|\d{1,2}\s+[а-яё]+))",
    re.IGNORECASE,
)
STATUS_NUMBER_PREFIX_RE = re.compile(r"^\s*(?:№\s*)?\d{1,6}\s*(?:[.)\-–—:]|\])\s*")
STATUS_DATE_PREFIX_RE = re.compile(
    r"^\s*(?:(?:дата|число|от)\s*[:№.\-–—]?\s*)?"
    r"(?:"
    r"\d{4}[./-]\d{1,2}[./-]\d{1,2}"
    r"|\d{1,2}[./-]\d{1,2}(?:[./-]\d{2,4})?"
    r"|\d{1,2}\s+[а-яё]+\s+\d{2,4}"
    r")"
    r"\s*(?:г\.?|год)?\s*(?:\d{1,2}:\d{2})?\s*(?:[.)\]}\-–—:,;]+)?\s*",
    re.IGNORECASE,
)
STATUS_STANDALONE_DAY_PREFIX_RE = re.compile(r"^\s*\d{1,3}\s+(?=[А-ЯЁA-Z])")

EXAMINEE_FORMS_PATTERN = r"испытуем(?:ый|ая|ое|ые|ого|ой|ую|ому|ым|ыми|ых|уюся)?"
EXAMINEE_STANDARD_STATE_RE = re.compile(
    rf"(?<![A-Za-zА-ЯЁа-яё])(состояние)\s+{EXAMINEE_FORMS_PATTERN}(?![A-Za-zА-ЯЁа-яё])",
    re.IGNORECASE,
)
EXAMINEE_START_RE = re.compile(
    rf"^\s*(?<![A-Za-zА-ЯЁа-яё]){EXAMINEE_FORMS_PATTERN}(?![A-Za-zА-ЯЁа-яё])"
    r"\s*(?:[:,.;!\-–—]+)?\s*",
    re.IGNORECASE,
)
EXAMINEE_ANY_RE = re.compile(
    rf"(?<![A-Za-zА-ЯЁа-яё]){EXAMINEE_FORMS_PATTERN}(?![A-Za-zА-ЯЁа-яё])",
    re.IGNORECASE,
)

FINAL_DIARY_TEXT = (
    "Состояние улучшилось. Жалоб не предъявляет. Острой психотической симптоматики не продуцирует. "
    "Фон настроения ровный, суицидальных мыслей не высказывает. Критика к состоянию присутствует. "
    "На текущую дату оформлена выписка из стационара. Даны рекомендации"
)

RUSSIAN_VOWELS = "аеёиоуыэюя"


GENDER_WORD_PAIRS: tuple[tuple[str, str], ...] = (
    # Person / patient nouns and pronouns.
    ("пациент", "пациентка"),
    ("пациента", "пациентки"),
    ("больной", "больная"),
    ("он", "она"),
    ("его", "ее"),
    ("ему", "ей"),
    ("ним", "ней"),
    ("него", "нее"),
    ("сам", "сама"),

    # Common clinical verbs in past tense.
    ("был", "была"),
    ("поступил", "поступила"),
    ("госпитализирован", "госпитализирована"),
    ("осмотрен", "осмотрена"),
    ("обследован", "обследована"),
    ("выписан", "выписана"),
    ("переведен", "переведена"),
    ("переведён", "переведена"),
    ("направлен", "направлена"),
    ("доставлен", "доставлена"),
    ("обратился", "обратилась"),
    ("находился", "находилась"),
    ("выезжал", "выезжала"),
    ("выехал", "выехала"),
    ("проживал", "проживала"),
    ("прожил", "прожила"),
    ("работал", "работала"),
    ("учился", "училась"),
    ("обучался", "обучалась"),
    ("родился", "родилась"),
    ("рожден", "рождена"),
    ("рождён", "рождена"),
    ("воспитывался", "воспитывалась"),
    ("рос", "росла"),
    ("развивался", "развивалась"),
    ("окончил", "окончила"),
    ("закончил", "закончила"),
    ("служил", "служила"),
    ("состоял", "состояла"),
    ("занимался", "занималась"),
    ("употреблял", "употребляла"),
    ("перенес", "перенесла"),
    ("перенёс", "перенесла"),
    ("оставался", "оставалась"),
    ("жаловался", "жаловалась"),
    ("отказывался", "отказывалась"),
    ("согласился", "согласилась"),
    ("успокоился", "успокоилась"),
    ("вернулся", "вернулась"),
    ("ориентировался", "ориентировалась"),
    ("смеялся", "смеялась"),
    ("раздражался", "раздражалась"),
    ("нуждался", "нуждалась"),
    ("держался", "держалась"),
    ("отмечал", "отмечала"),
    ("отметил", "отметила"),
    ("сообщал", "сообщала"),
    ("сообщил", "сообщила"),
    ("рассказал", "рассказала"),
    ("указал", "указала"),
    ("сказал", "сказала"),
    ("заявил", "заявила"),
    ("назвал", "назвала"),
    ("подтвердил", "подтвердила"),
    ("отрицал", "отрицала"),
    ("высказывал", "высказывала"),
    ("предъявлял", "предъявляла"),
    ("имел", "имела"),
    ("являлся", "являлась"),
    ("получал", "получала"),
    ("получил", "получила"),
    ("принимал", "принимала"),
    ("просил", "просила"),
    ("требовал", "требовала"),
    ("отвечал", "отвечала"),
    ("вступал", "вступала"),
    ("контактировал", "контактировала"),
    ("спал", "спала"),
    ("просыпался", "просыпалась"),
    ("засыпал", "засыпала"),
    ("ел", "ела"),
    ("пил", "пила"),
    ("съел", "съела"),
    ("выпил", "выпила"),
    ("посещал", "посещала"),
    ("участвовал", "участвовала"),
    ("выполнял", "выполняла"),
    ("соблюдал", "соблюдала"),
    ("нарушал", "нарушала"),
    ("покидал", "покидала"),
    ("лежал", "лежала"),
    ("сидел", "сидела"),
    ("стоял", "стояла"),
    ("плакал", "плакала"),
    ("сохранял", "сохраняла"),
    ("проявлял", "проявляла"),
    ("демонстрировал", "демонстрировала"),
    ("страдал", "страдала"),
    ("переносил", "переносила"),
    ("реагировал", "реагировала"),
    ("понимал", "понимала"),
    ("осознавал", "осознавала"),
    ("считал", "считала"),
    ("планировал", "планировала"),
    ("выражал", "выражала"),
    ("подписал", "подписала"),
    ("завершил", "завершила"),
    ("прошел", "прошла"),
    ("прошёл", "прошла"),
    ("шел", "шла"),
    ("шёл", "шла"),
    ("вел", "вела"),
    ("вёл", "вела"),
    ("дал", "дала"),
    ("лег", "легла"),
    ("лёг", "легла"),
    ("сел", "села"),
    ("стал", "стала"),

    # Short adjectives / participles that are frequent in observation diaries.
    ("ориентирован", "ориентирована"),
    ("дезориентирован", "дезориентирована"),
    ("контактен", "контактна"),
    ("доступен", "доступна"),
    ("адекватен", "адекватна"),
    ("спокоен", "спокойна"),
    ("тревожен", "тревожна"),
    ("напряжен", "напряжена"),
    ("напряжён", "напряжена"),
    ("возбужден", "возбуждена"),
    ("возбуждён", "возбуждена"),
    ("заторможен", "заторможена"),
    ("стабилен", "стабильна"),
    ("нестабилен", "нестабильна"),
    ("активен", "активна"),
    ("пассивен", "пассивна"),
    ("опрятен", "опрятна"),
    ("аккуратен", "аккуратна"),
    ("критичен", "критична"),
    ("самокритичен", "самокритична"),
    ("эмоционален", "эмоциональна"),
    ("раздражителен", "раздражительна"),
    ("подозрителен", "подозрительна"),
    ("насторожен", "насторожена"),
    ("насторожён", "насторожена"),
    ("замкнут", "замкнута"),
    ("открыт", "открыта"),
    ("общителен", "общительна"),
    ("вежлив", "вежлива"),
    ("агрессивен", "агрессивна"),
    ("конфликтен", "конфликтна"),
    ("упорядочен", "упорядочена"),
    ("собран", "собрана"),
    ("растерян", "растеряна"),
    ("испуган", "испугана"),
    ("обеспокоен", "обеспокоена"),
    ("заинтересован", "заинтересована"),
    ("мотивирован", "мотивирована"),
    ("настроен", "настроена"),
    ("склонен", "склонна"),
    ("способен", "способна"),
    ("готов", "готова"),
    ("удовлетворен", "удовлетворена"),
    ("удовлетворён", "удовлетворена"),
    ("согласен", "согласна"),
    ("сонлив", "сонлива"),
    ("вял", "вяла"),
    ("уверен", "уверена"),
    ("расторможен", "расторможена"),
    ("расторможён", "расторможена"),
    ("ухожен", "ухожена"),
    ("неухожен", "неухожена"),
)


@dataclass
class FillResult:
    filled_rows: int
    detected_rows: int
    month_cells_filled: int
    final_rows_filled: int
    next_status_index: int
    gender_replacements: int = 0
    removed_holiday_rows: int = 0
    removed_after_discharge_rows: int = 0


@dataclass
class DiaryBatchResult:
    created_files: list[Path]
    report_path: Path
    processed_files: int
    filled_rows: int
    detected_rows: int
    month_cells_filled: int
    final_rows_filled: int
    gender_replacements: int
    removed_holiday_rows: int
    removed_after_discharge_rows: int


def normalize_text(text: str) -> str:
    text = (text or "")
    text = re.sub(r"[\u00ad\u200b\u200c\u200d\u2060\ufeff]", "", text)
    text = text.replace("\xa0", " ").replace("\n", " ")
    text = WHITESPACE_RE.sub(" ", text).strip()
    return text


def strip_leading_status_metadata(text: str) -> str:
    value = normalize_text(text)
    patterns = (
        STATUS_LABEL_PREFIX_RE,
        STATUS_NUMBER_BEFORE_DATE_RE,
        STATUS_DATE_PREFIX_RE,
        DATE_PREFIX_RE,
        STATUS_NUMBER_PREFIX_RE,
        STATUS_STANDALONE_DAY_PREFIX_RE,
    )
    for _ in range(12):
        before = value
        for pattern in patterns:
            updated = pattern.sub("", value, count=1).lstrip(" —-–—:.);,]\t")
            if updated != value:
                value = normalize_text(updated)
                break
        if value == before:
            break
    return normalize_text(value)


def remove_examinee_words(text: str) -> str:
    value = normalize_text(text)
    value = EXAMINEE_STANDARD_STATE_RE.sub(lambda match: match.group(1), value)
    for _ in range(3):
        updated = EXAMINEE_START_RE.sub("", value)
        if updated == value:
            break
        value = normalize_text(updated)
    value = EXAMINEE_ANY_RE.sub("", value)
    value = re.sub(r"^\s*[:,.;!\-–—]+\s*", "", value)
    value = re.sub(r"\s+([,.;:!?])", r"\1", value)
    value = re.sub(r"([\(\[\{])\s+", r"\1", value)
    value = re.sub(r"\s+([\)\]\}])", r"\1", value)
    value = re.sub(r"\s{2,}", " ", value)
    return normalize_text(value)


def clean_status_text(text: str) -> str:
    return remove_examinee_words(strip_leading_status_metadata(text))


def is_signature_paragraph_text(text: str) -> bool:
    """Return True for template signature paragraphs that must stay untouched."""
    low = normalize_text(text).lower().replace("ё", "е")
    return any(low.startswith(marker) for marker in SIGNATURE_MARKERS)


def looks_like_status(text: str) -> bool:
    text = clean_status_text(text)
    low = text.lower()
    if len(text) < MIN_STATUS_LEN:
        return False
    if is_signature_paragraph_text(text):
        return False
    if any(low.startswith(prefix) for prefix in STRUCTURAL_DIARY_PREFIXES):
        return False
    if low in {"дневник наблюдения", "день госпитализации", "число", "дата", "месяц/год", "месяц / год"}:
        return False
    if re.fullmatch(r"[\d\s./-]+", text):
        return False
    return True


def extract_statuses_from_docx(path: str | Path) -> list[str]:
    doc = Document(str(path))
    statuses: list[str] = []

    def add_candidate(text: str) -> None:
        cleaned = clean_status_text(text)
        if looks_like_status(cleaned):
            statuses.append(cleaned)

    for paragraph in doc.paragraphs:
        add_candidate(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    add_candidate(paragraph.text)
    return statuses


def parse_month_year(text: str) -> tuple[int, int]:
    match = MONTH_YEAR_RE.fullmatch(normalize_text(text))
    if not match:
        raise ValueError("Введите начальный месяц и год в формате ММ.ГГГГ, например 06.2026")
    month = int(match.group(1))
    year = int(match.group(2))
    if month < 1 or month > 12:
        raise ValueError("Месяц должен быть от 01 до 12")
    if year < 1900 or year > 2200:
        raise ValueError("Год выглядит некорректно")
    return month, year


def parse_full_date(text: str) -> date:
    value = normalize_text(text)
    match = FULL_DATE_RE.fullmatch(value)
    if not match:
        raise ValueError("Введите дату в формате ДД.ММ.ГГГГ, например 11.06.2026")
    day = int(match.group(1))
    month = int(match.group(2))
    year = int(match.group(3))
    if year < 100:
        year += 2000 if year < 70 else 1900
    if year < 1900 or year > 2200:
        raise ValueError("Год выглядит некорректно")
    try:
        return date(year, month, day)
    except ValueError:
        raise ValueError("Дата выглядит некорректно") from None


def parse_admission_month_year(text: str) -> tuple[int, int]:
    value = normalize_text(text)
    if FULL_DATE_RE.fullmatch(value):
        admission_date = parse_full_date(value)
        return admission_date.month, admission_date.year
    return parse_month_year(value)


def parse_optional_discharge_date(text: str) -> date | None:
    value = normalize_text(text)
    if not value:
        return None
    return parse_full_date(value)


def safe_row_date(year: int, month: int, day: int | None) -> date | None:
    if day is None:
        return None
    try:
        return date(year, month, day)
    except ValueError:
        return None


def add_month(month: int, year: int, delta: int = 1) -> tuple[int, int]:
    month += delta
    while month > 12:
        month -= 12
        year += 1
    while month < 1:
        month += 12
        year -= 1
    return month, year


def format_month_year(month: int, year: int) -> str:
    return f"{month:02d}.{year:04d}"


def safe_filename_part(text: str) -> str:
    value = normalize_text(text)
    value = INVALID_FILENAME_CHARS_RE.sub(" ", value)
    value = WHITESPACE_RE.sub(" ", value).strip(" .")
    if not value:
        raise ValueError("Введите ФИО пациента")
    return value[:120].strip(" .")


def make_diary_output_name(patient_name: str, *, file_index: int, total_files: int) -> str:
    base = f"{patient_name} дневники"
    if total_files > 1:
        base = f"{base} {file_index:02d}"
    return f"{base}.docx"


def available_path(path: Path) -> Path:
    if not path.exists():
        return path
    stem = path.with_suffix("")
    ext = path.suffix
    for idx in range(2, 10000):
        candidate = Path(f"{stem}_{idx}{ext}")
        if not candidate.exists():
            return candidate
    raise RuntimeError(f"Не удалось подобрать имя файла для {path}")


def detect_gender_from_patient_name(patient_name: str) -> str | None:
    """Detect patient gender by the first word of the entered full name.

    Matches the GitHub diary-filler workflow for Russian surnames:
    surname ending with a vowel -> female, otherwise male.
    """
    value = normalize_text(patient_name)
    if not value:
        return None
    surname = re.sub(r"[^A-Za-zА-Яа-яЁё-]+", "", value.split()[0]).strip("-")
    if not surname:
        return None
    last = surname[-1].lower()
    return "female" if last in RUSSIAN_VOWELS else "male"


def gender_label(gender: str | None) -> str:
    if gender == "male":
        return "мужской"
    if gender == "female":
        return "женский"
    return "не определён"


def _preserve_case(source: str, target: str) -> str:
    if source.isupper():
        return target.upper()
    if source[:1].isupper():
        return target[:1].upper() + target[1:]
    return target


def _replace_gender_pair(text: str, source: str, target: str) -> tuple[str, int]:
    pattern = re.compile(rf"(?<![A-Za-zА-Яа-яЁё]){re.escape(source)}(?![A-Za-zА-Яа-яЁё])", re.IGNORECASE)
    count = 0

    def repl(match: re.Match) -> str:
        nonlocal count
        count += 1
        return _preserve_case(match.group(0), target)

    return pattern.sub(repl, text), count


def adapt_text_to_patient_gender(text: str, gender: str | None) -> tuple[str, int]:
    """Adapt known gendered diary words to the detected patient gender."""
    if gender not in {"male", "female"} or not text:
        return text, 0

    pairs = sorted(GENDER_WORD_PAIRS, key=lambda pair: max(len(pair[0]), len(pair[1])), reverse=True)
    result = text
    replacements = 0
    for male, female in pairs:
        source, target = (female, male) if gender == "male" else (male, female)
        result, changed = _replace_gender_pair(result, source, target)
        replacements += changed
    return result, replacements


def convert_text_gender(text: str, gender: str | None) -> tuple[str, int]:
    """Backward-compatible alias used by the combined application."""
    return adapt_text_to_patient_gender(text, gender)


def detect_first_month_year_from_docx(path: str | Path) -> tuple[int, int] | None:
    try:
        doc = Document(str(path))
        for table in doc.tables:
            month_year_col = find_month_year_column(table)
            day_col = find_day_column(table)
            if month_year_col is None:
                continue
            for row in table.rows:
                if not is_data_row(row, day_col):
                    continue
                if len(row.cells) <= month_year_col:
                    continue
                value = normalize_text(row.cells[month_year_col].text)
                try:
                    return parse_month_year(value)
                except ValueError:
                    continue
    except Exception:
        return None
    return None


def cell_int(text: str) -> int | None:
    value = normalize_text(text)
    match = re.fullmatch(r"0*(\d{1,2})", value)
    if not match:
        return None
    result = int(match.group(1))
    return result if 1 <= result <= 31 else None


def is_holiday_skip_date(day: int | None, month: int) -> bool:
    """Return True for rows dated 01.01-09.01 and 01.05-09.05."""
    return (
        day is not None
        and month in HOLIDAY_SKIP_MONTHS
        and HOLIDAY_SKIP_START_DAY <= day <= HOLIDAY_SKIP_END_DAY
    )


def is_data_row(row, day_col: int | None = None) -> bool:
    if not row.cells:
        return False
    if day_col is not None and len(row.cells) > day_col:
        return cell_int(row.cells[day_col].text) is not None
    first = normalize_text(row.cells[0].text)
    return bool(re.fullmatch(r"\d+", first))


def find_column_by_header(table, keywords: tuple[str, ...], *, fallback: int | None = None) -> int | None:
    if not table.rows:
        return fallback
    max_header_rows = min(5, len(table.rows))
    for row in table.rows[:max_header_rows]:
        for index, cell in enumerate(row.cells):
            text = normalize_text(cell.text).lower().replace(" ", "")
            if all(keyword.replace(" ", "") in text for keyword in keywords):
                return index
    return fallback


def find_diary_column(table) -> int | None:
    if not table.rows:
        return None
    for row in table.rows[: min(5, len(table.rows))]:
        for index, cell in enumerate(row.cells):
            text = normalize_text(cell.text).lower()
            if "дневник" in text or "наблюдения" in text:
                return index
    return len(table.rows[0].cells) - 1 if table.rows[0].cells else None


def find_day_column(table) -> int | None:
    if not table.rows or not table.rows[0].cells:
        return None
    col = find_column_by_header(table, ("число",), fallback=None)
    if col is not None:
        return col
    return 1 if len(table.rows[0].cells) >= 3 else 0


def find_month_year_column(table) -> int | None:
    if not table.rows or not table.rows[0].cells:
        return None
    for row in table.rows[: min(5, len(table.rows))]:
        for index, cell in enumerate(row.cells):
            text = normalize_text(cell.text).lower().replace(" ", "")
            if "месяц" in text and ("год" in text or "/" in text):
                return index
    return 2 if len(table.rows[0].cells) >= 4 else None


def clear_paragraph_keep_properties(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}pPr"):
            continue
        p.remove(child)


def reset_cell_to_one_paragraph(cell):
    if not cell.paragraphs:
        return cell.add_paragraph()
    first = cell.paragraphs[0]
    clear_paragraph_keep_properties(first)
    for paragraph in list(cell.paragraphs[1:]):
        paragraph._element.getparent().remove(paragraph._element)
    return first


def is_structural_diary_prefix(text: str) -> bool:
    """Return True for template notes that should stay above generated diary text."""
    low = normalize_text(text).lower().replace("ё", "е")
    return any(low.startswith(prefix) for prefix in STRUCTURAL_DIARY_PREFIXES)


def first_signature_paragraph_index(cell) -> int | None:
    for index, paragraph in enumerate(cell.paragraphs):
        if is_signature_paragraph_text(paragraph.text):
            return index
    return None


def add_run_with_size(paragraph, text: str):
    run = paragraph.add_run(text)
    run.font.size = Pt(STATUS_FONT_SIZE_PT)
    return run


def fill_text_cell(cell, text: str, *, alignment=None) -> None:
    paragraph = reset_cell_to_one_paragraph(cell)
    if alignment is not None:
        paragraph.alignment = alignment
    add_run_with_size(paragraph, text)


def write_diary_text_into_existing_paragraph(paragraph, diary_text: str) -> None:
    clear_paragraph_keep_properties(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run_with_size(paragraph, diary_text)


def fill_diary_text_cell(cell, diary_text: str, keep_signature: bool = True) -> None:
    """Fill only the diary area of a cell, preserving template signatures.

    Signatures are not extracted, deleted, re-created, moved, or reformatted.
    If the template row contains a structural note such as "Совместный осмотр...",
    that note remains above the generated diary text.
    """
    diary_text = remove_examinee_words(diary_text)
    _ = keep_signature  # kept for API compatibility; template signatures are always preserved when recognizable.

    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
        write_diary_text_into_existing_paragraph(paragraph, diary_text)
        return

    signature_index = first_signature_paragraph_index(cell)

    # No signature block: safe old strategy, because there is no template signature layout to preserve.
    if signature_index is None:
        paragraph = reset_cell_to_one_paragraph(cell)
        write_diary_text_into_existing_paragraph(paragraph, diary_text)
        return

    paragraphs = list(cell.paragraphs)
    before_signature = paragraphs[:signature_index]

    preserved_prefix_indexes = {
        index
        for index, paragraph in enumerate(before_signature)
        if is_structural_diary_prefix(paragraph.text)
    }

    search_from = max(preserved_prefix_indexes) + 1 if preserved_prefix_indexes else 0
    target_index: int | None = None

    for index in range(search_from, signature_index):
        if index in preserved_prefix_indexes:
            continue
        if not normalize_text(paragraphs[index].text):
            target_index = index
            break

    if target_index is None:
        for index in range(search_from, signature_index):
            if index not in preserved_prefix_indexes:
                target_index = index
                break

    if target_index is None:
        signature_paragraph = paragraphs[signature_index]
        target_paragraph = signature_paragraph.insert_paragraph_before("")
    else:
        target_paragraph = paragraphs[target_index]

    for index, paragraph in enumerate(before_signature):
        if index in preserved_prefix_indexes:
            continue
        if target_index is not None and index == target_index:
            continue
        if normalize_text(paragraph.text):
            clear_paragraph_keep_properties(paragraph)

    write_diary_text_into_existing_paragraph(target_paragraph, diary_text)


def remove_row(row) -> None:
    tr = row._tr
    parent = tr.getparent()
    if parent is not None:
        parent.remove(tr)


def collect_dated_entries(doc, start_month: int, start_year: int) -> list[dict]:
    entries: list[dict] = []
    current_month = start_month
    current_year = start_year
    previous_day: int | None = None

    for table in doc.tables:
        diary_col = find_diary_column(table)
        day_col = find_day_column(table)
        month_year_col = find_month_year_column(table)
        if diary_col is None:
            continue
        for row in table.rows:
            if not is_data_row(row, day_col):
                continue
            if len(row.cells) <= diary_col:
                continue
            day_value: int | None = None
            if day_col is not None and len(row.cells) > day_col:
                day_value = cell_int(row.cells[day_col].text)
            if day_value is not None and previous_day is not None and day_value < previous_day:
                current_month, current_year = add_month(current_month, current_year)
            current_date = safe_row_date(current_year, current_month, day_value)
            entries.append(
                {
                    "row": row,
                    "diary_col": diary_col,
                    "day_col": day_col,
                    "month_year_col": month_year_col,
                    "day": day_value,
                    "month": current_month,
                    "year": current_year,
                    "date": current_date,
                }
            )
            if day_value is not None:
                previous_day = day_value
    return entries


def should_remove_holiday(row_date: date | None) -> bool:
    if row_date is None:
        return False
    return (
        row_date.month in HOLIDAY_SKIP_MONTHS
        and HOLIDAY_SKIP_START_DAY <= row_date.day <= HOLIDAY_SKIP_END_DAY
    )


def fill_diary_file(
    path: str | Path,
    statuses: Sequence[str],
    *,
    start_idx: int = 0,
    repeat_statuses: bool = True,
    keep_signature: bool = True,
    fill_months: bool = True,
    start_month: int,
    start_year: int,
    discharge_date: date | None = None,
    force_final_diary: bool = True,
    remove_holiday_rows: bool = True,
    patient_gender: str | None = None,
) -> FillResult:
    """Fill all diary tables in a DOCX file.

    This version mirrors the GitHub diary-filler algorithm:
    - Month/year is calculated from the admission month and rolls forward when the day number decreases.
    - With discharge date, the chosen final row is the last data row not later than discharge.
    - Every data row after the chosen final row is removed by document order, even if it has an invalid date such as 29.02.
    - The final row is overwritten with the exact discharge day and month/year and receives the final discharge diary text.
    - 01.01-09.01 and 01.05-09.05 rows are removed, except the final discharge row.
    - Removed rows do not consume source diary texts.
    - Existing signature paragraphs in cells are preserved, not re-created.
    """
    doc = Document(str(path))
    idx = start_idx
    filled_rows = 0
    detected_rows = 0
    month_cells_filled = 0
    final_rows_filled = 0
    gender_replacements = 0
    removed_holiday_rows = 0
    removed_after_discharge_rows = 0

    data_entries: list[tuple[object, int, int | None, int | None]] = []
    for table in doc.tables:
        diary_col = find_diary_column(table)
        day_col = find_day_column(table)
        month_year_col = find_month_year_column(table)
        if diary_col is None:
            continue
        for row in table.rows:
            if not is_data_row(row, day_col):
                continue
            if len(row.cells) <= diary_col:
                continue
            data_entries.append((row, diary_col, day_col, month_year_col))

    dated_entries: list[dict[str, object]] = []
    current_month = start_month
    current_year = start_year
    previous_day: int | None = None

    for entry_index, (row, _diary_col, day_col, _month_year_col) in enumerate(data_entries):
        day_value: int | None = None
        if day_col is not None and len(row.cells) > day_col:
            day_value = cell_int(row.cells[day_col].text)

        if day_value is not None and previous_day is not None and day_value < previous_day:
            current_month, current_year = add_month(current_month, current_year, 1)
        if day_value is not None:
            previous_day = day_value

        row_date = safe_row_date(current_year, current_month, day_value)
        dated_entries.append(
            {
                "month": current_month,
                "year": current_year,
                "day": day_value,
                "date": row_date,
                "after_discharge": False,
                "skip_holiday": False,
                "skip_after_discharge": False,
            }
        )

    final_entry_index: int | None = None
    if discharge_date is not None:
        for entry_index in range(len(data_entries) - 1, -1, -1):
            row_date = dated_entries[entry_index]["date"]
            if isinstance(row_date, date) and row_date <= discharge_date:
                final_entry_index = entry_index
                break
        if final_entry_index is None and data_entries:
            raise ValueError(
                "В выбранной таблице не найдено ни одной строки до даты выписки. "
                "Проверьте месяц/год поступления и дату выписки."
            )
    else:
        for entry_index in range(len(data_entries) - 1, -1, -1):
            day_value = dated_entries[entry_index]["day"]
            row_month = int(dated_entries[entry_index]["month"])
            if not (remove_holiday_rows and is_holiday_skip_date(day_value if isinstance(day_value, int) else None, row_month)):
                final_entry_index = entry_index
                break

    for entry_index, entry in enumerate(dated_entries):
        day_value = entry["day"]
        row_month = int(entry["month"])
        is_final_row = final_entry_index is not None and entry_index == final_entry_index
        after_final_discharge_row = (
            discharge_date is not None
            and final_entry_index is not None
            and entry_index > final_entry_index
        )
        entry["after_discharge"] = after_final_discharge_row
        entry["skip_after_discharge"] = after_final_discharge_row
        entry["skip_holiday"] = (
            remove_holiday_rows
            and is_holiday_skip_date(day_value if isinstance(day_value, int) else None, row_month)
            and not is_final_row
            and not after_final_discharge_row
        )

    for entry_index, ((row, diary_col, day_col, month_year_col), entry) in enumerate(zip(data_entries, dated_entries)):
        detected_rows += 1
        skip_after_discharge = bool(entry["skip_after_discharge"])
        skip_holiday = bool(entry["skip_holiday"])

        if skip_after_discharge or skip_holiday:
            remove_row(row)
            if skip_after_discharge:
                removed_after_discharge_rows += 1
            else:
                removed_holiday_rows += 1
            continue

        row_month = int(entry["month"])
        row_year = int(entry["year"])
        is_final_diary_row = force_final_diary and final_entry_index is not None and entry_index == final_entry_index

        if is_final_diary_row and discharge_date is not None:
            if day_col is not None and len(row.cells) > day_col:
                fill_text_cell(
                    row.cells[day_col],
                    f"{discharge_date.day:02d}",
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                )
            row_month = discharge_date.month
            row_year = discharge_date.year

        if fill_months and month_year_col is not None and len(row.cells) > month_year_col:
            fill_text_cell(
                row.cells[month_year_col],
                format_month_year(row_month, row_year),
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )
            month_cells_filled += 1

        if is_final_diary_row:
            adapted_final_text, changed = adapt_text_to_patient_gender(FINAL_DIARY_TEXT, patient_gender)
            adapted_final_text = remove_examinee_words(adapted_final_text)
            gender_replacements += changed
            fill_diary_text_cell(row.cells[diary_col], adapted_final_text, keep_signature=keep_signature)
            filled_rows += 1
            final_rows_filled += 1
            continue

        if not statuses:
            continue
        if idx >= len(statuses):
            if repeat_statuses:
                idx = 0
            else:
                continue

        adapted_status, changed = adapt_text_to_patient_gender(statuses[idx], patient_gender)
        adapted_status = clean_status_text(adapted_status)
        gender_replacements += changed
        fill_diary_text_cell(row.cells[diary_col], adapted_status, keep_signature=keep_signature)
        filled_rows += 1
        idx += 1

    doc.save(str(path))
    return FillResult(
        filled_rows=filled_rows,
        detected_rows=detected_rows,
        month_cells_filled=month_cells_filled,
        final_rows_filled=final_rows_filled,
        next_status_index=idx,
        gender_replacements=gender_replacements,
        removed_holiday_rows=removed_holiday_rows,
        removed_after_discharge_rows=removed_after_discharge_rows,
    )

def read_statuses_from_files(paths: Iterable[str | Path]) -> list[str]:
    statuses: list[str] = []
    for path in paths:
        statuses.extend(extract_statuses_from_docx(path))
    return statuses


def open_folder(path: str | Path) -> None:
    path = str(path)
    try:
        if sys.platform.startswith("win"):
            os.startfile(path)  # type: ignore[attr-defined]
        elif sys.platform == "darwin":
            subprocess.Popen(["open", path])
        else:
            subprocess.Popen(["xdg-open", path])
    except Exception:
        pass


def fill_diary_batch(
    *,
    status_files: Sequence[str | Path],
    diary_files: Sequence[str | Path],
    output_dir: str | Path | None,
    patient_name: str,
    admission_value: str,
    gender_source_name: str | None = None,
    discharge_value: str = "",
    repeat_statuses: bool = True,
    reset_each_file: bool = True,
    keep_signature: bool = True,
    fill_months: bool = True,
    force_final_diary: bool = True,
    remove_holiday_rows: bool = True,
    open_result_folder: bool = False,
) -> DiaryBatchResult:
    if not diary_files:
        raise ValueError("Сначала выберите файлы-таблицы дневников, которые нужно заполнить.")
    if not status_files and not fill_months and not force_final_diary:
        raise ValueError("Сначала выберите файл(ы) с текстами дневников, включите месяц/год или финальную запись выписки.")

    start_month, start_year = parse_admission_month_year(admission_value)
    discharge_date_value = parse_optional_discharge_date(discharge_value)
    patient_filename = safe_filename_part(patient_name)
    gender_name = safe_filename_part(gender_source_name or patient_name)
    patient_gender = detect_gender_from_patient_name(gender_name)
    if patient_gender is None:
        raise ValueError("Введите ФИО так, чтобы первым словом была фамилия пациента. Например: Иванов И.И. или Петрова А.А.")

    statuses = read_statuses_from_files(status_files)
    if status_files and not statuses:
        raise ValueError("В выбранных файлах с текстами дневников не найдено подходящих текстов.")

    first_dir = Path(diary_files[0]).parent
    result_dir = Path(output_dir) if output_dir else first_dir
    result_dir.mkdir(parents=True, exist_ok=True)

    idx = 0
    created_files: list[Path] = []
    lines = [
        "ОТЧЁТ: заполнение дневников",
        f"Дата запуска: {datetime.now():%d.%m.%Y %H:%M:%S}",
        f"Пациент / имя файлов: {patient_filename}",
        f"ФИО для определения рода: {gender_name}",
        f"Поступление: {admission_value}",
        f"Выписка: {discharge_value or 'не указана'}",
        f"Папка результата: {result_dir}",
        f"Текстов дневников найдено: {len(statuses)}",
        "",
    ]

    total_filled = 0
    total_detected = 0
    total_months = 0
    total_final = 0
    total_gender = 0
    total_holidays = 0
    total_after_discharge = 0

    for n, src in enumerate(diary_files, start=1):
        src_path = Path(src)
        out_name = make_diary_output_name(patient_filename, file_index=n, total_files=len(diary_files))
        dst = available_path(result_dir / out_name)
        shutil.copy2(src_path, dst)
        effective_start_idx = 0 if reset_each_file else idx
        result = fill_diary_file(
            dst,
            statuses,
            start_idx=effective_start_idx,
            repeat_statuses=repeat_statuses,
            keep_signature=keep_signature,
            fill_months=fill_months,
            start_month=start_month,
            start_year=start_year,
            discharge_date=discharge_date_value,
            force_final_diary=force_final_diary,
            remove_holiday_rows=remove_holiday_rows,
            patient_gender=patient_gender,
        )
        if not reset_each_file:
            idx = result.next_status_index
        created_files.append(dst)
        total_filled += result.filled_rows
        total_detected += result.detected_rows
        total_months += result.month_cells_filled
        total_final += result.final_rows_filled
        total_gender += result.gender_replacements
        total_holidays += result.removed_holiday_rows
        total_after_discharge += result.removed_after_discharge_rows
        lines.append(
            f"{src_path.name}: строк найдено {result.detected_rows}; дневников заполнено {result.filled_rows}; "
            f"месяц/год {result.month_cells_filled}; финальных записей {result.final_rows_filled}; "
            f"замен пола {result.gender_replacements}; удалено праздников {result.removed_holiday_rows}; "
            f"удалено после выписки {result.removed_after_discharge_rows}"
        )

    lines.extend(
        [
            "",
            f"Файлов обработано: {len(created_files)}/{len(diary_files)}",
            f"Дневников заполнено: {total_filled}",
            f"Строк дневников найдено: {total_detected}",
            f"Дат месяц/год заполнено: {total_months}",
            f"Финальных записей: {total_final}",
            f"Грамматических замен по полу: {total_gender}",
            f"Удалено праздничных строк: {total_holidays}",
            f"Удалено строк после выписки: {total_after_discharge}",
        ]
    )
    report_path = available_path(result_dir / "ОТЧЁТ_дневники.txt")
    report_path.write_text("\n".join(lines), encoding="utf-8")
    if open_result_folder:
        open_folder(result_dir)

    return DiaryBatchResult(
        created_files=created_files,
        report_path=report_path,
        processed_files=len(created_files),
        filled_rows=total_filled,
        detected_rows=total_detected,
        month_cells_filled=total_months,
        final_rows_filled=total_final,
        gender_replacements=total_gender,
        removed_holiday_rows=total_holidays,
        removed_after_discharge_rows=total_after_discharge,
    )
