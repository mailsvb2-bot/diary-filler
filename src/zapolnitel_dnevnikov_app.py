import calendar
import os
import re
import sys
import shutil
import subprocess
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError:
    Document = None
    WD_ALIGN_PARAGRAPH = None
    Pt = None


APP_TITLE = "Заполнитель дневников"
MIN_STATUS_LEN = 25
STATUS_FONT_SIZE_PT = 8
SIGNATURE_MARKER = "лечащий врач"
HOLIDAY_SKIP_MONTHS = {1, 5}
HOLIDAY_SKIP_START_DAY = 1
HOLIDAY_SKIP_END_DAY = 9
SIGNATURE_MARKERS = (
    "лечащий врач",
    "зав.отделением",
    "зав. отделением",
    "зав отделением",
)
STRUCTURAL_DIARY_PREFIXES = (
    "совместный осмотр",
)

DATE_PREFIX_RE = re.compile(
    r"^\s*(?:"
    r"\d{1,2}[./-]\d{1,2}[./-]\d{2,4}\s*(?:г\.?|год)?"
    r"|\d{1,2}\s+[а-яё]+\s+\d{2,4}\s*(?:г\.?|год)?"
    r")\s*",
    re.IGNORECASE,
)
WHITESPACE_RE = re.compile(r"[ \t\r\f\v]+")
MONTH_YEAR_RE = re.compile(r"^\s*(\d{1,2})\s*[./-]\s*(\d{4})\s*$")
FULL_DATE_RE = re.compile(r"^\s*(\d{1,2})\s*[./-]\s*(\d{1,2})\s*[./-]\s*(\d{2,4})\s*$")
INVALID_FILENAME_CHARS_RE = re.compile(r'[<>:"/\\|?*\x00-\x1f]+')
# Source files sometimes contain the role label in different cases:
# "Испытуемый", "Испытуемая", "испытуемого", "испытуемой" etc.
# These words are service labels, not diary content, so they are removed before insertion.
EXAMINEE_FORMS_PATTERN = r"испытуем(?:ый|ая|ое|ые|ого|ой|ую|ому|ым|ыми|ых|уюся)?"
EXAMINEE_STANDARD_STATE_RE = re.compile(
    rf"(?<![A-Za-zА-ЯЁа-яё])(состояние)\s+{EXAMINEE_FORMS_PATTERN}(?![A-Za-zА-ЯЁа-яё])",
    re.IGNORECASE,
)
EXAMINEE_START_RE = re.compile(
    rf"^\s*(?<![A-Za-zА-ЯЁа-яё]){EXAMINEE_FORMS_PATTERN}(?![A-Za-zА-ЯЁа-яё])"
    r"\s*(?:[:,.;!\-–—]+)?\s*",
    re.IGNORECASE,
)
EXAMINEE_ANY_RE = re.compile(
    rf"(?<![A-Za-zА-ЯЁа-яё]){EXAMINEE_FORMS_PATTERN}(?![A-Za-zА-ЯЁа-яё])",
    re.IGNORECASE,
)

# Prefixes from source Word files are often service metadata, not diary text:
# dates, list numbers, record numbers, labels like "Дневник №...".
# They are stripped only at the beginning of extracted diary paragraphs.
STATUS_LABEL_PREFIX_RE = re.compile(
    r"^\s*(?:дата|число|номер|№|n|no\.?|запись|дневник)"
    r"\s*(?:№?\s*\d{1,6})?\s*[:.\-–—]\s*",
    re.IGNORECASE,
)
STATUS_NUMBER_BEFORE_DATE_RE = re.compile(
    r"^\s*\d{1,6}\s+(?=(?:\d{4}[./-]\d{1,2}[./-]\d{1,2}|\d{1,2}[./-]\d{1,2}|\d{1,2}\s+[а-яё]+))",
    re.IGNORECASE,
)
STATUS_NUMBER_PREFIX_RE = re.compile(r"^\s*(?:№\s*)?\d{1,6}\s*(?:[.)\-–—:]|\])\s*")
STATUS_DATE_PREFIX_RE = re.compile(
    r"^\s*(?:(?:дата|число|от)\s*[:№.\-–—]?\s*)?"
    r"(?:"
    r"\d{4}[./-]\d{1,2}[./-]\d{1,2}"
    r"|\d{1,2}[./-]\d{1,2}(?:[./-]\d{2,4})?"
    r"|\d{1,2}\s+[а-яё]+\s+\d{2,4}"
    r")"
    r"\s*(?:г\.?|год)?\s*(?:\d{1,2}:\d{2})?\s*(?:[.)\]}\-–—:,;]+)?\s*",
    re.IGNORECASE,
)
STATUS_STANDALONE_DAY_PREFIX_RE = re.compile(r"^\s*\d{1,3}\s+(?=[А-ЯЁA-Z])")

FINAL_DIARY_TEXT = (
    "Состояние улучшилось. Жалоб не предъявляет. Острой психотической симптоматики не продуцирует. "
    "Фон настроения ровный, суицидальных мыслей не высказывает. Критика к состоянию присутствует. "
    "На текущую дату оформлена выписка из стационара. Даны рекомендации"
)



@dataclass
class FillResult:
    filled_rows: int
    detected_rows: int
    month_cells_filled: int
    final_rows_filled: int
    next_status_index: int
    gender_replacements: int = 0
    removed_holiday_rows: int = 0
    removed_after_discharge_rows: int = 0


def normalize_text(text: str) -> str:
    """Normalize text copied from Word while preserving meaningful sentences."""
    # Word may contain non-breaking spaces, soft hyphens and zero-width marks.
    # They can visually hide inside words, so remove them before regex cleanup.
    text = (text or "")
    text = re.sub(r"[\u00ad\u200b\u200c\u200d\u2060\ufeff]", "", text)
    text = text.replace("\xa0", " ").replace("\n", " ")
    text = WHITESPACE_RE.sub(" ", text).strip()
    return text


def strip_leading_status_metadata(text: str) -> str:
    """Remove service prefixes before the actual diary text.

    Examples stripped from the start: "1.", "05.06.2026",
    "12 05.06.2026", "Дневник №3:", "05 Пациент...".
    The function deliberately works only at the beginning of the paragraph so
    dates and numbers inside the clinical text remain untouched.
    """
    value = normalize_text(text)
    patterns = (
        STATUS_LABEL_PREFIX_RE,
        STATUS_NUMBER_BEFORE_DATE_RE,
        STATUS_DATE_PREFIX_RE,
        DATE_PREFIX_RE,
        STATUS_NUMBER_PREFIX_RE,
        STATUS_STANDALONE_DAY_PREFIX_RE,
    )

    for _ in range(12):
        before = value
        for pattern in patterns:
            updated = pattern.sub("", value, count=1).lstrip(" —-–—:.);,]\t")
            if updated != value:
                value = normalize_text(updated)
                break
        if value == before:
            break
    return normalize_text(value)


def remove_examinee_words(text: str) -> str:
    """Remove source labels "Испытуемый" / "Испытуемая" from diary text.

    This is intentionally applied both while loading source texts and directly
    before writing into Word, so the words cannot leak into the final document
    even if they came from a table cell, old cache, or a later text transform.
    """
    value = normalize_text(text)

    # Standard source phrase: "Состояние испытуемого/испытуемой ...".
    # Keep the meaningful word "Состояние" and remove only the service label.
    value = EXAMINEE_STANDARD_STATE_RE.sub(lambda match: match.group(1), value)

    # Most sources use these words as a leading label, often with punctuation:
    # "Испытуемый:", "Испытуемая —", "Испытуемого,". Remove the whole label.
    for _ in range(3):
        updated = EXAMINEE_START_RE.sub("", value)
        if updated == value:
            break
        value = normalize_text(updated)

    # Also remove occurrences in all common grammatical forms elsewhere in the paragraph.
    value = EXAMINEE_ANY_RE.sub("", value)

    # Tidy punctuation/spaces that may remain after removing the word.
    value = re.sub(r"^\s*[:,.;!\-–—]+\s*", "", value)
    value = re.sub(r"\s+([,.;:!?])", r"\1", value)
    value = re.sub(r"([\(\[\{])\s+", r"\1", value)
    value = re.sub(r"\s+([\)\]\}])", r"\1", value)
    value = re.sub(r"\s{2,}", " ", value)
    return normalize_text(value)


def clean_status_text(text: str) -> str:
    """Prepare one diary/status paragraph for inserting into a table cell."""
    return remove_examinee_words(strip_leading_status_metadata(text))


def looks_like_status(text: str) -> bool:
    text = clean_status_text(text)
    low = text.lower()
    if len(text) < MIN_STATUS_LEN:
        return False
    if is_signature_paragraph_text(text):
        return False
    if low in {"дневник наблюдения", "день госпитализации", "число", "дата", "месяц/год", "месяц / год"}:
        return False
    if re.fullmatch(r"[\d\s./-]+", text):
        return False
    return True


def extract_statuses_from_docx(path: str) -> list[str]:
    """Extract diary texts from paragraphs and, if needed, from table cells."""
    doc = Document(path)
    statuses: list[str] = []

    def add_candidate(text: str) -> None:
        cleaned = clean_status_text(text)
        if looks_like_status(cleaned):
            statuses.append(cleaned)

    for paragraph in doc.paragraphs:
        add_candidate(paragraph.text)

    # Source documents are usually plain paragraphs, but this lets the app also
    # read statuses from Word tables if such a source appears later.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    add_candidate(paragraph.text)

    return statuses


def parse_month_year(text: str) -> tuple[int, int]:
    """Parse UI month/year value like 06.2026, 6.2026, 06/2026, 06-2026."""
    match = MONTH_YEAR_RE.fullmatch(normalize_text(text))
    if not match:
        raise ValueError("Введите начальный месяц и год в формате ММ.ГГГГ, например 06.2026")
    month = int(match.group(1))
    year = int(match.group(2))
    if month < 1 or month > 12:
        raise ValueError("Месяц должен быть от 01 до 12")
    if year < 1900 or year > 2200:
        raise ValueError("Год выглядит некорректно")
    return month, year




def parse_full_date(text: str) -> date:
    """Parse UI date value like 10.02.2026, 10/02/2026, 10-02-2026."""
    value = normalize_text(text)
    match = FULL_DATE_RE.fullmatch(value)
    if not match:
        raise ValueError("Введите дату в формате ДД.ММ.ГГГГ, например 11.06.2026")
    day = int(match.group(1))
    month = int(match.group(2))
    year = int(match.group(3))
    if year < 100:
        year += 2000 if year < 70 else 1900
    if year < 1900 or year > 2200:
        raise ValueError("Год выглядит некорректно")
    try:
        return date(year, month, day)
    except ValueError:
        raise ValueError("Дата выглядит некорректно") from None


def parse_admission_month_year(text: str) -> tuple[int, int]:
    """Parse admission value. Accepts MM.YYYY or full DD.MM.YYYY."""
    value = normalize_text(text)
    if FULL_DATE_RE.fullmatch(value):
        admission_date = parse_full_date(value)
        return admission_date.month, admission_date.year
    return parse_month_year(value)


def parse_optional_discharge_date(text: str) -> date | None:
    """Parse discharge date. Empty value keeps the old behavior without cutoff."""
    value = normalize_text(text)
    if not value:
        return None
    return parse_full_date(value)


def safe_row_date(year: int, month: int, day: int | None) -> date | None:
    if day is None:
        return None
    try:
        return date(year, month, day)
    except ValueError:
        return None


def add_month(month: int, year: int, delta: int = 1) -> tuple[int, int]:
    month += delta
    while month > 12:
        month -= 12
        year += 1
    while month < 1:
        month += 12
        year -= 1
    return month, year


def format_month_year(month: int, year: int) -> str:
    return f"{month:02d}.{year:04d}"


def safe_filename_part(text: str) -> str:
    """Return a Windows-safe filename fragment from a patient name."""
    value = normalize_text(text)
    value = INVALID_FILENAME_CHARS_RE.sub(" ", value)
    value = WHITESPACE_RE.sub(" ", value).strip(" .")
    if not value:
        raise ValueError("Введите ФИО пациента")
    return value[:120].strip(" .")


def make_diary_output_name(patient_name: str, *, file_index: int, total_files: int) -> str:
    """Create output filename requested by the UI: '<ФИО пациента> дневники.docx'."""
    base = f"{patient_name} дневники"
    if total_files > 1:
        base = f"{base} {file_index:02d}"
    return f"{base}.docx"


RUSSIAN_VOWELS = "аеёиоуыэюя"

# Explicit bidirectional pairs. The converter intentionally uses a controlled
# medical/diary vocabulary instead of a blind "change every -л to -ла" rule:
# this keeps ordinary nouns and template service text from being damaged.
GENDER_WORD_PAIRS: tuple[tuple[str, str], ...] = (
    # Person / patient nouns and pronouns.
    ("пациент", "пациентка"),
    ("пациента", "пациентки"),
    ("больной", "больная"),
    ("он", "она"),
    ("его", "ее"),
    ("ему", "ей"),
    ("ним", "ней"),
    ("него", "нее"),
    ("сам", "сама"),

    # Common clinical verbs in past tense.
    ("был", "была"),
    ("поступил", "поступила"),
    ("госпитализирован", "госпитализирована"),
    ("осмотрен", "осмотрена"),
    ("обследован", "обследована"),
    ("выписан", "выписана"),
    ("переведен", "переведена"),
    ("переведён", "переведена"),
    ("направлен", "направлена"),
    ("доставлен", "доставлена"),
    ("обратился", "обратилась"),
    ("находился", "находилась"),
    ("оставался", "оставалась"),
    ("жаловался", "жаловалась"),
    ("отказывался", "отказывалась"),
    ("согласился", "согласилась"),
    ("успокоился", "успокоилась"),
    ("вернулся", "вернулась"),
    ("ориентировался", "ориентировалась"),
    ("смеялся", "смеялась"),
    ("раздражался", "раздражалась"),
    ("нуждался", "нуждалась"),
    ("держался", "держалась"),
    ("отмечал", "отмечала"),
    ("отметил", "отметила"),
    ("сообщал", "сообщала"),
    ("сообщил", "сообщила"),
    ("рассказал", "рассказала"),
    ("указал", "указала"),
    ("сказал", "сказала"),
    ("заявил", "заявила"),
    ("назвал", "назвала"),
    ("подтвердил", "подтвердила"),
    ("отрицал", "отрицала"),
    ("высказывал", "высказывала"),
    ("предъявлял", "предъявляла"),
    ("имел", "имела"),
    ("являлся", "являлась"),
    ("получал", "получала"),
    ("получил", "получила"),
    ("принимал", "принимала"),
    ("просил", "просила"),
    ("требовал", "требовала"),
    ("отвечал", "отвечала"),
    ("вступал", "вступала"),
    ("контактировал", "контактировала"),
    ("спал", "спала"),
    ("просыпался", "просыпалась"),
    ("засыпал", "засыпала"),
    ("ел", "ела"),
    ("пил", "пила"),
    ("съел", "съела"),
    ("выпил", "выпила"),
    ("посещал", "посещала"),
    ("участвовал", "участвовала"),
    ("выполнял", "выполняла"),
    ("соблюдал", "соблюдала"),
    ("нарушал", "нарушала"),
    ("покидал", "покидала"),
    ("лежал", "лежала"),
    ("сидел", "сидела"),
    ("стоял", "стояла"),
    ("плакал", "плакала"),
    ("сохранял", "сохраняла"),
    ("проявлял", "проявляла"),
    ("демонстрировал", "демонстрировала"),
    ("страдал", "страдала"),
    ("переносил", "переносила"),
    ("реагировал", "реагировала"),
    ("понимал", "понимала"),
    ("осознавал", "осознавала"),
    ("считал", "считала"),
    ("планировал", "планировала"),
    ("выражал", "выражала"),
    ("подписал", "подписала"),
    ("завершил", "завершила"),
    ("прошел", "прошла"),
    ("прошёл", "прошла"),
    ("шел", "шла"),
    ("шёл", "шла"),
    ("вел", "вела"),
    ("вёл", "вела"),
    ("дал", "дала"),
    ("лег", "легла"),
    ("лёг", "легла"),
    ("сел", "села"),
    ("стал", "стала"),

    # Short adjectives / participles that are frequent in observation diaries.
    ("ориентирован", "ориентирована"),
    ("дезориентирован", "дезориентирована"),
    ("контактен", "контактна"),
    ("доступен", "доступна"),
    ("адекватен", "адекватна"),
    ("спокоен", "спокойна"),
    ("тревожен", "тревожна"),
    ("напряжен", "напряжена"),
    ("напряжён", "напряжена"),
    ("возбужден", "возбуждена"),
    ("возбуждён", "возбуждена"),
    ("заторможен", "заторможена"),
    ("стабилен", "стабильна"),
    ("нестабилен", "нестабильна"),
    ("активен", "активна"),
    ("пассивен", "пассивна"),
    ("опрятен", "опрятна"),
    ("аккуратен", "аккуратна"),
    ("критичен", "критична"),
    ("самокритичен", "самокритична"),
    ("эмоционален", "эмоциональна"),
    ("раздражителен", "раздражительна"),
    ("подозрителен", "подозрительна"),
    ("насторожен", "насторожена"),
    ("насторожён", "насторожена"),
    ("замкнут", "замкнута"),
    ("открыт", "открыта"),
    ("общителен", "общительна"),
    ("вежлив", "вежлива"),
    ("агрессивен", "агрессивна"),
    ("конфликтен", "конфликтна"),
    ("упорядочен", "упорядочена"),
    ("собран", "собрана"),
    ("растерян", "растеряна"),
    ("испуган", "испугана"),
    ("обеспокоен", "обеспокоена"),
    ("заинтересован", "заинтересована"),
    ("мотивирован", "мотивирована"),
    ("настроен", "настроена"),
    ("склонен", "склонна"),
    ("способен", "способна"),
    ("готов", "готова"),
    ("удовлетворен", "удовлетворена"),
    ("удовлетворён", "удовлетворена"),
    ("согласен", "согласна"),
    ("сонлив", "сонлива"),
    ("вял", "вяла"),
    ("уверен", "уверена"),
    ("расторможен", "расторможена"),
    ("расторможён", "расторможена"),
    ("ухожен", "ухожена"),
    ("неухожен", "неухожена"),
)


def detect_gender_from_patient_name(fio: str) -> str | None:
    """
    Detect patient gender by the first word of the entered full name.

    The rule matches the user's workflow for Russian surnames:
    - surname ending with а/я or another vowel -> female;
    - surname ending with a consonant/й/ь -> male.
    """
    value = normalize_text(fio)
    if not value:
        return None
    surname = re.sub(r"[^A-Za-zА-Яа-яЁё-]+", "", value.split()[0]).strip("-")
    if not surname:
        return None
    last = surname[-1].lower()
    return "female" if last in RUSSIAN_VOWELS else "male"


def gender_label(gender: str | None) -> str:
    if gender == "male":
        return "мужской"
    if gender == "female":
        return "женский"
    return "не определён"


def _preserve_case(source: str, target: str) -> str:
    if source.isupper():
        return target.upper()
    if source[:1].isupper():
        return target[:1].upper() + target[1:]
    return target


def _replace_gender_pair(text: str, source: str, target: str) -> tuple[str, int]:
    pattern = re.compile(rf"(?<![A-Za-zА-Яа-яЁё]){re.escape(source)}(?![A-Za-zА-Яа-яЁё])", re.IGNORECASE)
    count = 0

    def repl(match: re.Match) -> str:
        nonlocal count
        count += 1
        return _preserve_case(match.group(0), target)

    return pattern.sub(repl, text), count


def adapt_text_to_patient_gender(text: str, gender: str | None) -> tuple[str, int]:
    """Adapt known gendered diary words to the detected patient gender."""
    if gender not in {"male", "female"} or not text:
        return text, 0

    pairs = sorted(GENDER_WORD_PAIRS, key=lambda pair: max(len(pair[0]), len(pair[1])), reverse=True)
    result = text
    replacements = 0
    for male, female in pairs:
        source, target = (female, male) if gender == "male" else (male, female)
        result, changed = _replace_gender_pair(result, source, target)
        replacements += changed
    return result, replacements


def detect_first_month_year_from_docx(path: str) -> tuple[int, int] | None:
    """Detect the first existing Month/Year value in the diary table, if present."""
    try:
        doc = Document(path)
        for table in doc.tables:
            month_year_col = find_month_year_column(table)
            day_col = find_day_column(table)
            if month_year_col is None:
                continue
            for row in table.rows:
                if not is_data_row(row, day_col):
                    continue
                if len(row.cells) <= month_year_col:
                    continue
                value = normalize_text(row.cells[month_year_col].text)
                try:
                    return parse_month_year(value)
                except ValueError:
                    continue
    except Exception:
        return None
    return None


def cell_int(text: str) -> int | None:
    text = normalize_text(text)
    if not text:
        return None
    match = re.fullmatch(r"0*(\d{1,2})", text)
    if not match:
        return None
    value = int(match.group(1))
    if 1 <= value <= 31:
        return value
    return None


def is_holiday_skip_date(day: int | None, month: int) -> bool:
    """Return True for rows that must be removed from the output diary.

    The requested skip ranges are 01.01-09.01 and 01.05-09.05.
    Year does not matter.
    """
    return (
        day is not None
        and month in HOLIDAY_SKIP_MONTHS
        and HOLIDAY_SKIP_START_DAY <= day <= HOLIDAY_SKIP_END_DAY
    )


def remove_table_row(row) -> None:
    """Remove a Word table row while preserving the rest of the table."""
    tr = row._tr
    parent = tr.getparent()
    if parent is not None:
        parent.remove(tr)


def is_data_row(row, day_col: int | None = None) -> bool:
    if not row.cells:
        return False

    # Main current template: first column is hospitalization day, second column is day of month.
    if day_col is not None and len(row.cells) > day_col:
        return cell_int(row.cells[day_col].text) is not None

    # Backward-compatible fallback for older templates where the first column was the only marker.
    first = normalize_text(row.cells[0].text)
    return bool(re.fullmatch(r"\d+", first))


def find_column_by_header(table, keywords: tuple[str, ...], *, fallback: int | None = None) -> int | None:
    if not table.rows:
        return fallback

    max_header_rows = min(5, len(table.rows))
    for row in table.rows[:max_header_rows]:
        for index, cell in enumerate(row.cells):
            text = normalize_text(cell.text).lower().replace(" ", "")
            if all(keyword.replace(" ", "") in text for keyword in keywords):
                return index
    return fallback


def find_diary_column(table) -> int | None:
    """Find the observation/diary column by header text. Fallback: last column."""
    if not table.rows:
        return None

    for row in table.rows[: min(5, len(table.rows))]:
        for index, cell in enumerate(row.cells):
            text = normalize_text(cell.text).lower()
            if "дневник" in text or "наблюдения" in text:
                return index

    # In these diary templates the diary text is the rightmost/main wide column.
    return len(table.rows[0].cells) - 1 if table.rows[0].cells else None


def find_day_column(table) -> int | None:
    """Find the day-of-month column. In the provided templates it is the second column."""
    if not table.rows or not table.rows[0].cells:
        return None

    # Header is usually row 1, column 1: "Число".
    col = find_column_by_header(table, ("число",), fallback=None)
    if col is not None:
        return col

    # Fallback for the current 4-column template: [hospital day, day of month, month/year, diary].
    return 1 if len(table.rows[0].cells) >= 3 else 0


def find_month_year_column(table) -> int | None:
    """Find the month/year column. In the provided templates it is the third column."""
    if not table.rows or not table.rows[0].cells:
        return None

    for row in table.rows[: min(5, len(table.rows))]:
        for index, cell in enumerate(row.cells):
            text = normalize_text(cell.text).lower().replace(" ", "")
            if "месяц" in text and ("год" in text or "/" in text):
                return index

    # Fallback for the current 4-column template.
    return 2 if len(table.rows[0].cells) >= 4 else None


def clear_paragraph_keep_properties(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        # Keep paragraph properties, remove runs/bookmarks/content.
        if child.tag.endswith("}pPr"):
            continue
        p.remove(child)


def reset_cell_to_one_paragraph(cell):
    """Clear cell content without destroying table/cell properties."""
    if not cell.paragraphs:
        return cell.add_paragraph()

    first = cell.paragraphs[0]
    clear_paragraph_keep_properties(first)

    for paragraph in list(cell.paragraphs[1:]):
        paragraph._element.getparent().remove(paragraph._element)

    return first


def is_signature_paragraph_text(text: str) -> bool:
    """Return True for template signature paragraphs that must stay untouched."""
    low = normalize_text(text).lower().replace("ё", "е")
    return any(low.startswith(marker) for marker in SIGNATURE_MARKERS)


def is_structural_diary_prefix(text: str) -> bool:
    """Return True for template notes that should stay above the generated diary text."""
    low = normalize_text(text).lower().replace("ё", "е")
    return any(low.startswith(prefix) for prefix in STRUCTURAL_DIARY_PREFIXES)


def first_signature_paragraph_index(cell) -> int | None:
    for index, paragraph in enumerate(cell.paragraphs):
        if is_signature_paragraph_text(paragraph.text):
            return index
    return None


def add_run_with_size(paragraph, text: str):
    run = paragraph.add_run(text)
    run.font.size = Pt(STATUS_FONT_SIZE_PT)
    return run


def fill_text_cell(cell, text: str, *, alignment=None) -> None:
    paragraph = reset_cell_to_one_paragraph(cell)
    if alignment is not None:
        paragraph.alignment = alignment
    add_run_with_size(paragraph, text)


def write_diary_text_into_existing_paragraph(paragraph, diary_text: str) -> None:
    clear_paragraph_keep_properties(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run_with_size(paragraph, diary_text)


def fill_diary_text_cell(cell, diary_text: str, keep_signature: bool = True) -> None:
    diary_text = remove_examinee_words(diary_text)
    """
    Fill only the diary area of a cell.

    Important: signatures are part of the Word template. They are not extracted,
    deleted, re-created, moved, or reformatted here. The function writes the new
    diary text into an existing paragraph before the first signature paragraph.
    If the template row contains a structural note such as "Совместный осмотр...",
    that note is preserved above the generated diary text.
    """
    # Kept for backward compatibility with older calls/reports. In the safe UI
    # signatures are always preserved because the templates already contain them.
    _ = keep_signature

    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
        write_diary_text_into_existing_paragraph(paragraph, diary_text)
        return

    signature_index = first_signature_paragraph_index(cell)

    # If the cell has no recognizable signature block, use the old safe clearing
    # strategy because there is no template signature layout to preserve.
    if signature_index is None:
        paragraph = reset_cell_to_one_paragraph(cell)
        write_diary_text_into_existing_paragraph(paragraph, diary_text)
        return

    paragraphs = list(cell.paragraphs)
    before_signature = paragraphs[:signature_index]

    preserved_prefix_indexes = {
        index
        for index, paragraph in enumerate(before_signature)
        if is_structural_diary_prefix(paragraph.text)
    }

    search_from = max(preserved_prefix_indexes) + 1 if preserved_prefix_indexes else 0

    target_index: int | None = None

    # Prefer an existing empty paragraph between template note and signatures.
    for index in range(search_from, signature_index):
        if index in preserved_prefix_indexes:
            continue
        if not normalize_text(paragraphs[index].text):
            target_index = index
            break

    # If all paragraphs before signatures contain old diary text, reuse the first
    # non-structural paragraph before the signature block.
    if target_index is None:
        for index in range(search_from, signature_index):
            if index not in preserved_prefix_indexes:
                target_index = index
                break

    # Very rare case: signature is the first paragraph. Insert text before it;
    # the signature paragraph itself still remains untouched.
    if target_index is None:
        signature_paragraph = paragraphs[signature_index]
        target_paragraph = signature_paragraph.insert_paragraph_before("")
    else:
        target_paragraph = paragraphs[target_index]

    # Clear stale generated diary text before signatures, but do not remove any
    # paragraphs. Keeping paragraph shells helps preserve the template spacing.
    for index, paragraph in enumerate(before_signature):
        if index in preserved_prefix_indexes:
            continue
        if target_index is not None and index == target_index:
            continue
        if normalize_text(paragraph.text):
            clear_paragraph_keep_properties(paragraph)

    write_diary_text_into_existing_paragraph(target_paragraph, diary_text)



def fill_diary_file(
    filepath: str,
    statuses: list[str],
    *,
    start_idx: int = 0,
    repeat_statuses: bool = True,
    keep_signature: bool = True,
    fill_months: bool = True,
    start_month: int = 1,
    start_year: int = 2026,
    force_final_diary: bool = True,
    final_diary_text: str = FINAL_DIARY_TEXT,
    patient_gender: str | None = None,
    discharge_date: date | None = None,
) -> FillResult:
    """
    Fill all diary tables in a docx.

    The current diary template has 4 columns:
    1) hospitalization day, 2) day of month, 3) month.year, 4) diary text.

    Month/year logic:
    - take only the month/year from the admission field, e.g. 02.2026 or 10.02.2026 -> 02.2026;
    - write this value into the Month/Year column for every remaining data row;
    - when the day-of-month value becomes smaller than the previous one, move to the next month.

    Discharge cutoff logic:
    - if a discharge date is entered, the program chooses the last template row
      whose calculated date is not later than the discharge date;
    - absolutely every data row after that chosen row is removed, even if the
      template row itself contains an invalid calendar date like 29.02 in a
      non-leap year;
    - the remaining last data row is always rewritten as the discharge date:
      day column = DD, month/year column = MM.YYYY;
    - this last row receives the fixed discharge diary text; an exact matching source row is not required.

    Holiday row logic:
    - rows dated 01.01-09.01 and 01.05-09.05 are removed;
    - the final discharge row is never removed as a holiday, because it must close the table;
    - removed rows do not consume diary texts from the source list.
    """
    doc = Document(filepath)
    idx = start_idx
    filled_rows = 0
    detected_rows = 0
    month_cells_filled = 0
    final_rows_filled = 0
    gender_replacements = 0
    removed_holiday_rows = 0
    removed_after_discharge_rows = 0

    data_entries: list[tuple[object, int, int | None, int | None]] = []
    for table in doc.tables:
        diary_col = find_diary_column(table)
        day_col = find_day_column(table)
        month_year_col = find_month_year_column(table)

        if diary_col is None:
            continue

        for row in table.rows:
            if not is_data_row(row, day_col):
                continue
            if len(row.cells) <= diary_col:
                continue
            data_entries.append((row, diary_col, day_col, month_year_col))

    # Calculate the date represented by every template row before any rows are removed.
    dated_entries: list[dict[str, object]] = []
    current_month = start_month
    current_year = start_year
    previous_day: int | None = None

    for entry_index, (row, _diary_col, day_col, _month_year_col) in enumerate(data_entries):
        day_value: int | None = None
        if day_col is not None and len(row.cells) > day_col:
            day_value = cell_int(row.cells[day_col].text)

        if day_value is not None and previous_day is not None and day_value < previous_day:
            current_month, current_year = add_month(current_month, current_year, 1)
        if day_value is not None:
            previous_day = day_value

        row_date = safe_row_date(current_year, current_month, day_value)
        dated_entries.append(
            {
                "month": current_month,
                "year": current_year,
                "day": day_value,
                "date": row_date,
                "after_discharge": False,
                "skip_holiday": False,
                "skip_after_discharge": False,
            }
        )

    # Pick the final data row. With a discharge date this is the last row whose
    # calculated date is not later than discharge. If the template has no exact
    # DD.MM.YYYY row, this row is overwritten with the discharge date.
    final_entry_index: int | None = None
    if discharge_date is not None:
        for entry_index in range(len(data_entries) - 1, -1, -1):
            row_date = dated_entries[entry_index]["date"]
            if isinstance(row_date, date) and row_date <= discharge_date:
                final_entry_index = entry_index
                break
        if final_entry_index is None and data_entries:
            raise ValueError(
                "В выбранной таблице не найдено ни одной строки до даты выписки. "
                "Проверьте месяц/год поступления и дату выписки."
            )
    else:
        for entry_index in range(len(data_entries) - 1, -1, -1):
            day_value = dated_entries[entry_index]["day"]
            row_month = int(dated_entries[entry_index]["month"])
            if not is_holiday_skip_date(day_value if isinstance(day_value, int) else None, row_month):
                final_entry_index = entry_index
                break

    # Now mark rows to remove. With a discharge date, the chosen discharge
    # row must become the last data row in the whole table. Therefore we cut by
    # document order after final_entry_index, not only by calculated row dates.
    # This is important for templates that contain impossible dates such as
    # 29.02 in a non-leap year: such rows do not have a valid date object, but
    # they still must disappear when they stand after the discharge row.
    for entry_index, entry in enumerate(dated_entries):
        day_value = entry["day"]
        row_month = int(entry["month"])
        is_final_row = final_entry_index is not None and entry_index == final_entry_index
        after_final_discharge_row = (
            discharge_date is not None
            and final_entry_index is not None
            and entry_index > final_entry_index
        )
        entry["after_discharge"] = after_final_discharge_row
        entry["skip_after_discharge"] = after_final_discharge_row
        entry["skip_holiday"] = (
            is_holiday_skip_date(day_value if isinstance(day_value, int) else None, row_month)
            and not is_final_row
            and not after_final_discharge_row
        )

    for entry_index, ((row, diary_col, day_col, month_year_col), entry) in enumerate(zip(data_entries, dated_entries)):
        detected_rows += 1
        skip_after_discharge = bool(entry["skip_after_discharge"])
        skip_holiday = bool(entry["skip_holiday"])

        if skip_after_discharge or skip_holiday:
            remove_table_row(row)
            if skip_after_discharge:
                removed_after_discharge_rows += 1
            else:
                removed_holiday_rows += 1
            continue

        row_month = int(entry["month"])
        row_year = int(entry["year"])
        is_final_diary_row = force_final_diary and final_entry_index is not None and entry_index == final_entry_index

        # The discharge date must be visibly written into the final data row.
        # This is intentionally done even if the template did not contain that exact date.
        if is_final_diary_row and discharge_date is not None:
            if day_col is not None and len(row.cells) > day_col:
                fill_text_cell(
                    row.cells[day_col],
                    f"{discharge_date.day:02d}",
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                )
            row_month = discharge_date.month
            row_year = discharge_date.year

        if fill_months and month_year_col is not None and len(row.cells) > month_year_col:
            fill_text_cell(
                row.cells[month_year_col],
                format_month_year(row_month, row_year),
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )
            month_cells_filled += 1

        if is_final_diary_row:
            adapted_final_text, changed = adapt_text_to_patient_gender(final_diary_text, patient_gender)
            adapted_final_text = remove_examinee_words(adapted_final_text)
            gender_replacements += changed
            fill_diary_text_cell(row.cells[diary_col], adapted_final_text, keep_signature=keep_signature)
            filled_rows += 1
            final_rows_filled += 1
            continue

        if not statuses:
            continue
        if idx >= len(statuses):
            if repeat_statuses:
                idx = 0
            else:
                continue

        adapted_status, changed = adapt_text_to_patient_gender(statuses[idx], patient_gender)
        adapted_status = clean_status_text(adapted_status)
        gender_replacements += changed
        fill_diary_text_cell(row.cells[diary_col], adapted_status, keep_signature=keep_signature)
        filled_rows += 1
        idx += 1

    doc.save(filepath)
    return FillResult(
        filled_rows=filled_rows,
        detected_rows=detected_rows,
        month_cells_filled=month_cells_filled,
        final_rows_filled=final_rows_filled,
        next_status_index=idx,
        gender_replacements=gender_replacements,
        removed_holiday_rows=removed_holiday_rows,
        removed_after_discharge_rows=removed_after_discharge_rows,
    )


def open_folder(path: str) -> None:
    try:
        if sys.platform.startswith("win"):
            os.startfile(path)  # type: ignore[attr-defined]
        elif sys.platform == "darwin":
            subprocess.Popen(["open", path])
        else:
            subprocess.Popen(["xdg-open", path])
    except Exception:
        # Opening the folder is a convenience only; processing is already done.
        pass


class App:
    """Compact two-column Tkinter UI without external UI dependencies."""

    BG = "#07111f"
    PANEL = "#0f1b2d"
    PANEL_2 = "#132238"
    FIELD = "#081625"
    LINE = "#24405f"
    TEXT = "#f3f8ff"
    MUTED = "#9fb2c8"
    ACCENT = "#14b8a6"
    ACCENT_HOVER = "#2dd4bf"
    TABLE = "#16a34a"
    TABLE_HOVER = "#22c55e"
    TEXTS = "#0ea5e9"
    TEXTS_HOVER = "#38bdf8"
    SAVE = "#f59e0b"
    SAVE_HOVER = "#fbbf24"
    DANGER = "#fb7185"
    DARK_TEXT = "#051018"

    def __init__(self, root):
        self.root = root
        self.root.title(APP_TITLE)
        self.root.geometry("760x520")
        self.root.minsize(720, 500)
        self.root.configure(bg=self.BG)

        self.statuses: list[str] = []
        self.status_files: list[str] = []
        self.diary_files: list[str] = []

        # Hidden stable defaults. The UI is intentionally clean, but the old behavior is preserved.
        self.repeat_var = tk.BooleanVar(value=True)
        self.reset_each_file_var = tk.BooleanVar(value=True)
        self.keep_signature_var = tk.BooleanVar(value=True)
        self.force_final_diary_var = tk.BooleanVar(value=True)
        self.fill_months_var = tk.BooleanVar(value=True)
        self.start_month_year_var = tk.StringVar(value=datetime.now().strftime("%m.%Y"))
        self.discharge_date_var = tk.StringVar(value="")
        self.patient_name_var = tk.StringVar(value="")
        self.output_dir_var = tk.StringVar(value="")
        self.patient_name_var.trace_add("write", self._update_gender_hint)

        self._setup_style()
        self._build_ui()

        if Document is None:
            messagebox.showerror(
                "Не установлен python-docx",
                "Установите зависимость командой:\n\npip install python-docx",
            )

    def _setup_style(self) -> None:
        self.style = ttk.Style()
        try:
            self.style.theme_use("clam")
        except tk.TclError:
            pass

        self.style.configure("TFrame", background=self.BG)
        self.style.configure(
            "Modern.Horizontal.TProgressbar",
            troughcolor="#0b1828",
            background=self.ACCENT,
            bordercolor="#0b1828",
            lightcolor=self.ACCENT,
            darkcolor=self.ACCENT,
        )
        self.title_font = ("Segoe UI", 20, "bold")
        self.subtitle_font = ("Segoe UI", 9)
        self.label_font = ("Segoe UI", 10, "bold")
        self.small_font = ("Segoe UI", 8)
        self.entry_font = ("Segoe UI", 11, "bold")
        self.button_font = ("Segoe UI", 9, "bold")
        self.big_button_font = ("Segoe UI", 12, "bold")

    def _build_ui(self) -> None:
        shell = tk.Frame(self.root, bg=self.BG)
        shell.pack(fill="both", expand=True, padx=22, pady=16)
        shell.grid_columnconfigure(0, weight=1)
        shell.grid_rowconfigure(1, weight=1)

        header = tk.Frame(shell, bg=self.BG)
        header.grid(row=0, column=0, sticky="ew", pady=(0, 10))
        header.grid_columnconfigure(0, weight=1)
        tk.Label(
            header,
            text="Заполнитель дневников",
            bg=self.BG,
            fg=self.TEXT,
            font=self.title_font,
        ).grid(row=0, column=0, sticky="w")
        tk.Label(
            header,
            text="ФИО · поступление/выписка · таблица · тексты · папка сохранения",
            bg=self.BG,
            fg=self.MUTED,
            font=self.subtitle_font,
        ).grid(row=1, column=0, sticky="w", pady=(1, 0))

        form = tk.Frame(shell, bg=self.BG)
        form.grid(row=1, column=0, sticky="nsew")
        form.grid_columnconfigure(0, weight=1, uniform="cols")
        form.grid_columnconfigure(1, weight=1, uniform="cols")

        # Row 1: patient + month.
        patient_card = self._card(form, row=0, column=0)
        self._card_label(patient_card, "ФИО пациента").grid(row=0, column=0, sticky="w", pady=(0, 6))
        self.patient_entry = self._entry(patient_card, self.patient_name_var)
        self.patient_entry.grid(row=1, column=0, sticky="ew", ipady=5)
        self.gender_hint_label = self._value_label(
            patient_card,
            "Пол определится по фамилии",
            self.MUTED,
            wraplength=300,
        )
        self.gender_hint_label.grid(row=2, column=0, sticky="ew", pady=(5, 0))

        dates_card = self._card(form, row=0, column=1)
        dates_card.grid_columnconfigure(0, weight=1)
        dates_card.grid_columnconfigure(1, weight=1)

        self._card_label(dates_card, "Месяц/год поступления").grid(row=0, column=0, sticky="w", pady=(0, 6))
        self._card_label(dates_card, "Выписка").grid(row=0, column=1, sticky="w", padx=(10, 0), pady=(0, 6))
        self.month_entry = self._entry(dates_card, self.start_month_year_var, justify="center", width=12)
        self.month_entry.grid(row=1, column=0, sticky="ew", ipady=5, padx=(0, 8))
        self.discharge_entry = self._entry(dates_card, self.discharge_date_var, justify="center", width=12)
        self.discharge_entry.grid(row=1, column=1, sticky="ew", ipady=5, padx=(8, 0))
        self.start_month_hint = tk.Label(
            dates_card,
            text="слева берётся месяц перед годом: 02.2026 / 10.02.2026 · справа: дата выписки",
            bg=self.PANEL,
            fg=self.MUTED,
            font=self.small_font,
            anchor="w",
            justify="left",
            wraplength=330,
        )
        self.start_month_hint.grid(row=2, column=0, columnspan=2, sticky="ew", pady=(5, 0))

        # Row 2: diary tables + source texts.
        tables_card = self._card(form, row=1, column=0)
        self._small_button(
            tables_card,
            text="Выбрать таблицы дневников",
            command=self.load_diaries,
            bg=self.TABLE,
            hover=self.TABLE_HOVER,
            fg="#03120a",
        ).grid(row=0, column=0, sticky="ew")
        self.selected_tables_label = self._value_label(
            tables_card,
            "Таблица не выбрана",
            self.DANGER,
        )
        self.selected_tables_label.grid(row=1, column=0, sticky="ew", pady=(8, 0))
        self.diaries_label = self.selected_tables_label

        texts_card = self._card(form, row=1, column=1)
        self._small_button(
            texts_card,
            text="Выбрать тексты дневников",
            command=self.load_statuses,
            bg=self.TEXTS,
            hover=self.TEXTS_HOVER,
            fg="#03101f",
        ).grid(row=0, column=0, sticky="ew")
        self.status_label = self._value_label(texts_card, "Файлы с текстами не выбраны", self.DANGER)
        self.status_label.grid(row=1, column=0, sticky="ew", pady=(8, 0))
        self.status_count = self._value_label(texts_card, "Найдено текстов: 0", self.MUTED)
        self.status_count.grid(row=2, column=0, sticky="ew", pady=(2, 0))

        # Row 3: save folder, wide and visible.
        save_card = self._card(form, row=2, column=0, columnspan=2, pady=(0, 8))
        save_card.grid_columnconfigure(0, weight=0)
        save_card.grid_columnconfigure(1, weight=1)
        self._small_button(
            save_card,
            text="Выбрать куда сохранить файл",
            command=self.choose_output_dir,
            bg=self.SAVE,
            hover=self.SAVE_HOVER,
            fg="#1a0b02",
            width=30,
        ).grid(row=0, column=0, sticky="w")
        self.output_dir_label = self._value_label(
            save_card,
            "Папка не выбрана: сохраню Word-файл рядом с таблицей",
            self.MUTED,
            wraplength=420,
        )
        self.output_dir_label.grid(row=0, column=1, sticky="ew", padx=(14, 0))

        self._build_action_bar(shell)

    def _build_action_bar(self, parent: tk.Frame) -> None:
        action = tk.Frame(parent, bg=self.BG)
        action.grid(row=2, column=0, sticky="ew", pady=(8, 0))
        action.grid_columnconfigure(0, weight=1)

        self.run_button = self._small_button(
            action,
            text="ЗАПОЛНИТЬ ТАБЛИЦЫ ДНЕВНИКОВ",
            command=self.run,
            bg=self.ACCENT,
            hover=self.ACCENT_HOVER,
            fg="#03120f",
            big=True,
        )
        self.run_button.grid(row=0, column=0, sticky="ew")

        footer = tk.Frame(action, bg=self.BG)
        footer.grid(row=1, column=0, sticky="ew", pady=(8, 0))
        footer.grid_columnconfigure(0, weight=1)
        self.progress = ttk.Progressbar(
            footer,
            mode="determinate",
            maximum=100,
            value=0,
            style="Modern.Horizontal.TProgressbar",
        )
        self.progress.grid(row=0, column=0, sticky="ew", padx=(0, 12))
        self.status = tk.Label(
            footer,
            text="Готов",
            bg=self.BG,
            fg=self.MUTED,
            font=self.small_font,
            anchor="e",
        )
        self.status.grid(row=0, column=1, sticky="e")

    def _card(self, parent: tk.Frame, *, row: int, column: int, columnspan: int = 1, pady=(0, 10)) -> tk.Frame:
        card = tk.Frame(parent, bg=self.PANEL, highlightbackground=self.LINE, highlightthickness=1)
        card.grid(row=row, column=column, columnspan=columnspan, sticky="nsew", padx=6, pady=pady)
        card.grid_columnconfigure(0, weight=1)
        inner = tk.Frame(card, bg=self.PANEL)
        inner.grid(row=0, column=0, sticky="nsew", padx=12, pady=10)
        inner.grid_columnconfigure(0, weight=1)
        return inner

    def _card_label(self, parent: tk.Frame, text: str) -> tk.Label:
        return tk.Label(
            parent,
            text=text,
            bg=self.PANEL,
            fg=self.TEXT,
            font=self.label_font,
            anchor="w",
        )

    def _entry(self, parent: tk.Frame, variable: tk.StringVar, *, justify="left", width=None) -> tk.Entry:
        return tk.Entry(
            parent,
            textvariable=variable,
            justify=justify,
            width=width,
            bg=self.FIELD,
            fg=self.TEXT,
            insertbackground=self.TEXT,
            relief="flat",
            font=self.entry_font,
            highlightbackground=self.LINE,
            highlightcolor=self.ACCENT,
            highlightthickness=1,
        )

    def _value_label(self, parent: tk.Frame, text: str, color: str, *, wraplength: int = 300) -> tk.Label:
        return tk.Label(
            parent,
            text=text,
            bg=self.PANEL,
            fg=color,
            font=self.small_font,
            anchor="w",
            justify="left",
            wraplength=wraplength,
        )

    def _small_button(self, parent, *, text: str, command, bg: str, hover: str, fg: str, big: bool = False, width=None):
        button = tk.Button(
            parent,
            text=text,
            command=command,
            bg=bg,
            fg=fg,
            activebackground=hover,
            activeforeground=fg,
            relief="flat",
            bd=0,
            cursor="hand2",
            font=self.big_button_font if big else self.button_font,
            padx=12,
            pady=8 if big else 7,
            width=width,
        )
        button.bind("<Enter>", lambda _event: button.config(bg=hover))
        button.bind("<Leave>", lambda _event: button.config(bg=bg))
        return button

    def _set_footer(self, text: str, color: str | None = None) -> None:
        self.status.config(text=text, fg=color or self.MUTED)
        self.root.update_idletasks()

    def _update_gender_hint(self, *_args) -> None:
        if not hasattr(self, "gender_hint_label"):
            return
        gender = detect_gender_from_patient_name(self.patient_name_var.get())
        if gender is None:
            self.gender_hint_label.config(text="Пол определится по фамилии", fg=self.MUTED)
        elif gender == "male":
            self.gender_hint_label.config(text="Определено: мужской род", fg=self.ACCENT_HOVER)
        else:
            self.gender_hint_label.config(text="Определено: женский род", fg=self.ACCENT_HOVER)

    def choose_output_dir(self):
        directory = filedialog.askdirectory(title="Выберите папку для сохранения результата")
        if not directory:
            return
        self.output_dir_var.set(directory)
        display = directory
        if len(display) > 78:
            display = "…" + display[-77:]
        self.output_dir_label.config(text=display, fg=self.SAVE_HOVER)
        self._set_footer("Сохранение прямо в выбранную папку", self.SAVE_HOVER)

    def load_statuses(self):
        files = filedialog.askopenfilenames(filetypes=[("Word files", "*.docx")])
        if not files:
            return

        try:
            self.status_files = list(files)
            self.statuses = []
            for file in self.status_files:
                self.statuses.extend(extract_statuses_from_docx(file))

            if not self.statuses:
                messagebox.showwarning(
                    "Тексты не найдены",
                    "В выбранных файлах не найдено подходящих длинных текстов дневника.",
                )

            names = ", ".join(os.path.basename(f) for f in self.status_files[:2])
            if len(self.status_files) > 2:
                names += f" и ещё {len(self.status_files) - 2}"
            self.status_label.config(text=f"✓ {names}", fg=self.TABLE_HOVER)
            self.status_count.config(text=f"Найдено текстов: {len(self.statuses)}", fg=self.TEXTS_HOVER)
            self._set_footer(f"Загружено текстов: {len(self.statuses)}", self.TABLE_HOVER)
        except Exception as e:
            messagebox.showerror("Ошибка чтения текстов", str(e))

    def load_diaries(self):
        files = filedialog.askopenfilenames(filetypes=[("Word files", "*.docx")])
        if not files:
            return

        self.diary_files = sorted(list(files), key=lambda p: os.path.basename(p).lower())
        selected_names = [os.path.basename(file) for file in self.diary_files]
        if len(selected_names) == 1:
            selected_title = selected_names[0]
        elif len(selected_names) <= 2:
            selected_title = "; ".join(selected_names)
        else:
            selected_title = f"{'; '.join(selected_names[:2])} и ещё {len(selected_names) - 2}"
        self.selected_tables_label.config(text=selected_title, fg=self.TABLE_HOVER)

        detected_start = detect_first_month_year_from_docx(self.diary_files[0])
        if detected_start is not None:
            self.start_month_year_var.set(format_month_year(*detected_start))
            self.start_month_hint.config(
                text=f"найдено поступление: {format_month_year(*detected_start)} · выписка: ДД.ММ.ГГГГ",
                fg=self.TABLE_HOVER,
            )
        else:
            self.start_month_hint.config(text="слева берётся месяц перед годом: 02.2026 / 10.02.2026 · справа: дата выписки", fg=self.MUTED)

        self._set_footer(f"Выбрана таблица: {selected_title}", self.TEXTS_HOVER)

    def run(self):
        if Document is None:
            messagebox.showerror("Ошибка", "Не установлен python-docx. Выполните: pip install python-docx")
            return

        if not self.statuses and not self.fill_months_var.get() and not self.force_final_diary_var.get():
            messagebox.showerror(
                "Ошибка",
                "Сначала выберите файл(ы) с текстами дневников/статусами, включите заполнение 3-го столбца месяцем/годом или включите финальную запись выписки.",
            )
            return

        if not self.statuses and (self.fill_months_var.get() or self.force_final_diary_var.get()):
            self._set_footer("Тексты не выбраны: будут заполнены только выбранные служебные поля", self.SAVE_HOVER)

        if not self.diary_files:
            messagebox.showerror("Ошибка", "Сначала выберите файлы-таблицы дневников, которые нужно заполнить.")
            return

        try:
            start_month, start_year = parse_admission_month_year(self.start_month_year_var.get())
        except ValueError as e:
            messagebox.showerror("Ошибка в дате поступления", str(e))
            self.month_entry.focus_set()
            return

        try:
            discharge_date_value = parse_optional_discharge_date(self.discharge_date_var.get())
        except ValueError as e:
            messagebox.showerror("Ошибка в дате выписки", str(e))
            self.discharge_entry.focus_set()
            return

        try:
            patient_filename = safe_filename_part(self.patient_name_var.get())
        except ValueError as e:
            messagebox.showerror("Не указано ФИО пациента", str(e))
            self.patient_entry.focus_set()
            return

        patient_gender = detect_gender_from_patient_name(patient_filename)
        if patient_gender is None:
            messagebox.showerror(
                "Не удалось определить пол",
                "Введите ФИО так, чтобы первым словом была фамилия пациента. Например: Иванов И.И. или Петрова А.А.",
            )
            self.patient_entry.focus_set()
            return

        self.run_button.config(state="disabled", text="ИДЁТ ЗАПОЛНЕНИЕ...")
        self.progress.config(value=0, maximum=max(len(self.diary_files), 1))
        self.root.update_idletasks()

        selected_output_dir = normalize_text(self.output_dir_var.get())
        result_dir = Path(selected_output_dir) if selected_output_dir else Path(self.diary_files[0]).parent
        result_dir.mkdir(parents=True, exist_ok=True)

        total_filled = 0
        total_detected = 0
        total_month_cells = 0
        total_final_rows = 0
        total_gender_replacements = 0
        total_removed_holiday_rows = 0
        total_removed_after_discharge_rows = 0
        ok_files = 0
        errors: list[str] = []
        global_idx = 0

        try:
            for pos, src in enumerate(self.diary_files, start=1):
                output_name = make_diary_output_name(patient_filename, file_index=pos, total_files=len(self.diary_files))
                dst = result_dir / output_name
                try:
                    shutil.copy2(src, dst)
                    start_idx = 0 if self.reset_each_file_var.get() else global_idx
                    result = fill_diary_file(
                        str(dst),
                        self.statuses,
                        start_idx=start_idx,
                        repeat_statuses=self.repeat_var.get(),
                        keep_signature=self.keep_signature_var.get(),
                        fill_months=self.fill_months_var.get(),
                        start_month=start_month,
                        start_year=start_year,
                        force_final_diary=self.force_final_diary_var.get(),
                        patient_gender=patient_gender,
                        discharge_date=discharge_date_value,
                    )
                    if not self.reset_each_file_var.get():
                        global_idx = result.next_status_index

                    total_filled += result.filled_rows
                    total_detected += result.detected_rows
                    total_month_cells += result.month_cells_filled
                    total_final_rows += result.final_rows_filled
                    total_gender_replacements += result.gender_replacements
                    total_removed_holiday_rows += result.removed_holiday_rows
                    total_removed_after_discharge_rows += result.removed_after_discharge_rows
                    ok_files += 1
                    self.progress.config(value=pos)
                    self._set_footer(
                        f"Файлы {pos}/{len(self.diary_files)} · дневники {total_filled} · обрезано {total_removed_after_discharge_rows}",
                        self.TEXTS_HOVER,
                    )
                except Exception as e:
                    errors.append(f"{os.path.basename(src)}: {e}")
                    self.progress.config(value=pos)
                    self.root.update_idletasks()
        finally:
            self.run_button.config(state="normal", text="ЗАПОЛНИТЬ ТАБЛИЦЫ ДНЕВНИКОВ")

        self._set_footer(
            f"Готово: файлов {ok_files}, дневников {total_filled}, обрезано {total_removed_after_discharge_rows}",
            self.TABLE_HOVER,
        )

        message = (
            f"Готово.\n\n"
            f"Файл(ы) сохранены в папку:\n{result_dir}\n\n"
            f"ФИО пациента: {patient_filename}\n"
            f"Определённый род: {gender_label(patient_gender)}\n"
            f"Дата/месяц поступления: {self.start_month_year_var.get()}\n"
            f"Дата выписки: {discharge_date_value.strftime('%d.%m.%Y') if discharge_date_value else 'не указана'}\n"
            f"Исправлений рода в текстах: {total_gender_replacements}\n"
            f"Обрезано строк после даты выписки: {total_removed_after_discharge_rows}\n"
            f"Удалено строк 01.01-09.01 и 01.05-09.05: {total_removed_holiday_rows}\n"
            f"Обработано файлов: {ok_files}\n"
            f"Найдено строк таблиц: {total_detected}\n"
            f"Заполнено дневников: {total_filled}\n"
            f"Заполнено ячеек месяц/год: {total_month_cells}\n"
            f"Финальных записей выписки: {total_final_rows}\n"
            f"Начальный месяц/год таблицы: {format_month_year(start_month, start_year)}"
        )
        if errors:
            message += f"\n\nБыли ошибки: {len(errors)}:\n" + "\n".join(errors[:10])
            if len(errors) > 10:
                message += f"\n…и ещё {len(errors) - 10}"
            messagebox.showwarning("Готово с предупреждениями", message)
        else:
            messagebox.showinfo("Готово", message)

        open_folder(str(result_dir))


if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.mainloop()
