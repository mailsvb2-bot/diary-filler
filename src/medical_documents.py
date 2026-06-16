"""
Автозаполнение медицинских DOCX-документов.

Рабочий сценарий:
1) В UI выбирается первичный документ пациента: направление на госпитализацию или первичный осмотр.
2) При необходимости выбирается ЭПИ-файл.
3) Указывается дата выписки.
4) Программа берёт встроенные шаблоны из папки templates или из embedded_templates.py и создаёт все документы.

Пользователь НЕ выбирает шаблоны в UI. Шаблоны являются частью программы.
Для GitHub/EXE сборки шаблоны также доступны как текстовый base64-модуль embedded_templates.py.
"""

from __future__ import annotations

import copy
import os
import re
import sys
import tempfile
import traceback
from dataclasses import asdict, dataclass, field
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

try:
    from diary_filler import adapt_text_to_patient_gender, detect_gender_from_patient_name
except Exception:  # pragma: no cover - защитный fallback для автономного использования модуля
    def detect_gender_from_patient_name(_patient_name: str) -> str | None:
        return None

    def adapt_text_to_patient_gender(text: str, _gender: str | None) -> tuple[str, int]:
        return text, 0


APP_TITLE = "Медицинские документы — автозаполнение"
DATE_FMT = "%d.%m.%Y"

TEMPLATE_FILES: Dict[str, str] = {
    "primary": "2 Первичный.docx",
    "discharge": "3 Выписной.docx",
    "commission": "4 Комиссионный.docx",
    "vk_mse": "5 ВК на МСЭ.docx",
    "sick_leave_vk": "6 ВК больничный.docx",
    "rvk": "6 Акт для рвк.docx",
}

OUTPUT_SUFFIXES: Dict[str, str] = {
    "primary": "Первичный_осмотр",
    "discharge": "Выписной_эпикриз",
    "commission": "Совместный_осмотр",
    "vk_mse": "ВК_на_МСЭ",
    "sick_leave_vk": "ВК_больничный",
    "rvk": "Акт_для_РВК",
}

DOCUMENT_LABELS: Dict[str, str] = {
    "primary": "Первичный осмотр",
    "discharge": "Выписной эпикриз",
    "commission": "Совместный осмотр",
    "vk_mse": "ВК на МСЭ",
    "sick_leave_vk": "ВК больничный",
    "rvk": "Акт для РВК",
}

DOCUMENT_ORDER = ("primary", "discharge", "commission", "vk_mse", "sick_leave_vk", "rvk")


# -----------------------------------------------------------------------------
# Пути к встроенным ресурсам
# -----------------------------------------------------------------------------

def app_dir() -> Path:
    """Папка программы с учётом PyInstaller onefile."""
    if getattr(sys, "frozen", False):
        return Path(getattr(sys, "_MEIPASS", Path(sys.executable).parent))
    return Path(__file__).resolve().parent


def template_dir() -> Path:
    return app_dir() / "templates"


def embedded_template_path(filename: str) -> Optional[Path]:
    """Return a temporary DOCX template path from embedded base64 storage.

    GitHub/EXE builds can work without a physical templates/ folder: the binary
    DOCX templates are stored in embedded_templates.py as UTF-8 text and restored
    into the user temp directory on demand. If templates/ exists next to the app,
    those files still have priority so the developer can edit templates normally.
    """
    try:
        from embedded_templates import TEMPLATE_B64  # type: ignore
    except Exception:
        return None
    raw = TEMPLATE_B64.get(filename)
    if not raw:
        return None
    import base64

    cache_dir = Path(tempfile.gettempdir()) / "medical_diary_autofill_templates"
    cache_dir.mkdir(parents=True, exist_ok=True)
    out = cache_dir / filename
    data = base64.b64decode(raw.encode("ascii") if isinstance(raw, str) else raw)
    if not out.exists() or out.stat().st_size != len(data):
        out.write_bytes(data)
    return out


def bundled_template_path(kind: str) -> Path:
    try:
        filename = TEMPLATE_FILES[kind]
    except KeyError as exc:
        raise KeyError(f"Неизвестный тип документа: {kind}") from exc
    physical = template_dir() / filename
    if physical.exists():
        return physical
    embedded = embedded_template_path(filename)
    return embedded or physical


# -----------------------------------------------------------------------------
# Нормализация текста
# -----------------------------------------------------------------------------

DASHES = {
    "\u2010": "-",
    "\u2011": "-",
    "\u2012": "-",
    "\u2013": "-",
    "\u2014": "-",
    "\u2212": "-",
}


def normalize_text(text: str) -> str:
    if not text:
        return ""
    for src, dst in DASHES.items():
        text = text.replace(src, dst)
    text = text.replace("\xa0", " ")
    text = text.replace("\v", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalize_match(text: str) -> str:
    text = normalize_text(text).lower().replace("ё", "е")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def clean_value(text: str) -> str:
    text = normalize_text(text)
    text = text.strip(" \t:-—–")
    text = re.sub(r"^сюда\s+подстав(?:лять|ляется).*$", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^нужно\s*/\s*не\s*нужно$", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^(состоит\s*/\s*не\s*состоит)\s*[:.-]?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^(нужен\s*/\s*не\s*нужен)\s*[:.-]?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^(нужно\s*/\s*не\s*нужно)\s*[:.-]?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^(да\s*/\s*нет)\s*[:.-]?\s*", "", text, flags=re.IGNORECASE)
    return text.strip()


def looks_like_label(value: str) -> bool:
    value = normalize_match(value)
    if not value:
        return False
    known = [
        "год рождения",
        "дата рождения",
        "зарегистрирован",
        "проживает",
        "на учете",
        "работает",
        "место работы",
        "должность",
        "больничный лист",
        "оформление инвалидности",
        "направление от рвк",
        "в 3 отделение",
        "жалобы",
        "анамнез",
        "психический статус",
        "соматический статус",
        "сомато-неврологический статус",
        "план обследования",
        "план лечения",
        "на основании",
        "диагноз",
        "эпидемиологический",
        "врач",
        "зав",
    ]
    return any(value.startswith(k) for k in known)


# -----------------------------------------------------------------------------
# Модель данных
# -----------------------------------------------------------------------------

@dataclass
class PatientData:
    case_number: str = ""
    fio: str = ""
    # Имя пациента для названия создаваемых файлов.
    # ВАЖНО: это отдельное поле; оно не подменяет ФИО внутри документов.
    output_fio: str = ""
    birth: str = ""
    registered: str = ""
    psych_account: str = ""
    work_org: str = ""
    position: str = ""
    sick_leave: str = ""
    disability: str = ""
    rvk_referral: str = ""
    admission: str = ""

    complaints: str = ""
    life_anamnesis: str = ""
    disease_anamnesis: str = ""
    mental_status: str = ""
    somatic_status: str = ""
    examination_plan: str = ""
    treatment_plan: str = ""
    diagnosis: str = ""
    epidemiology: str = ""

    admission_date: str = ""
    discharge_date: str = ""
    epi_text: str = ""
    input_document_kind: str = ""

    # Ручные реквизиты из UI для отдельных документов.
    rvk_act_number: str = ""
    rvk_military_commissariat: str = ""
    rvk_work_position: str = ""
    vk_date: str = ""
    vk_protocol_number: str = ""
    vk_protocol_date: str = ""
    vk_mse_work_org: str = ""
    vk_mse_position: str = ""
    sick_leave_vk_date: str = ""
    sick_leave_vk_protocol_number: str = ""
    sick_leave_vk_protocol_date: str = ""
    sick_leave_vk_commission_date: str = ""
    sick_leave_vk_work_position: str = ""
    commission_date: str = ""
    commission_number: str = ""

    doctor: str = "Балаганин С.В"
    head: str = "Можарова Е.А."

    warnings: List[str] = field(default_factory=list)

    def lab_dates(self) -> Dict[str, str]:
        result = {"day1": "", "day2": "", "flg": ""}
        dt = parse_date(self.admission_date)
        if not dt:
            return result
        result["day1"] = (dt + timedelta(days=1)).strftime(DATE_FMT)
        result["day2"] = (dt + timedelta(days=2)).strftime(DATE_FMT)
        result["flg"] = (dt - timedelta(days=27)).strftime(DATE_FMT)
        return result

    def missing_critical_fields(self) -> List[str]:
        missing = []
        if not self.fio:
            missing.append("Ф.И.О.")
        if not self.birth:
            missing.append("год/дата рождения")
        if not self.admission_date:
            missing.append("дата госпитализации")
        return missing

    def missing_recommended_fields(self) -> List[str]:
        checks = [
            ("жалобы", self.complaints),
            ("анамнез жизни", self.life_anamnesis),
            ("анамнез заболевания", self.disease_anamnesis),
            ("психический статус", self.mental_status),
            ("диагноз", self.diagnosis),
            ("план лечения", self.treatment_plan),
        ]
        return [name for name, value in checks if not value]


# -----------------------------------------------------------------------------
# Чтение DOCX
# -----------------------------------------------------------------------------

def iter_block_items(parent) -> Iterable[Paragraph | Table]:
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        return

    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def extract_docx_text(path: str | Path) -> str:
    doc = Document(str(path))
    lines: List[str] = []

    def walk(parent):
        for block in iter_block_items(parent):
            if isinstance(block, Paragraph):
                lines.append(block.text)
            elif isinstance(block, Table):
                for row in block.rows:
                    for cell in row.cells:
                        walk(cell)

    walk(doc)
    return normalize_text("\n".join(lines))


# -----------------------------------------------------------------------------
# Парсер первичного документа: направление на госпитализацию или первичный осмотр
# -----------------------------------------------------------------------------

class MedicalTextParser:
    FIELD_ALIASES: Dict[str, Sequence[str]] = {
        "case_number": ("История болезни №", "История болезни N", "ИБ №"),
        "fio": ("Ф.И.О.", "ФИО", "Фамилия Имя Отчество"),
        "birth": ("Год рождения", "Дата рождения", "г.р."),
        "registered": ("Зарегистрирован", "зарегистрирован по адресу", "Проживает", "Адрес регистрации"),
        "psych_account": ("На учёте у психиатров", "На учете у психиатров"),
        "work_org": ("Работает в организации", "Место работы", "Работа"),
        "position": ("Должность",),
        "sick_leave": ("Больничный лист", "ЛН", "Лист нетрудоспособности"),
        "disability": ("Оформление инвалидности", "Инвалидность"),
        "rvk_referral": ("Направление от РВК", "РВК"),
        "admission": ("В 3 отделение КДП поступает", "Поступает", "Поступил"),
        "doctor": ("Врач психиатр", "Врач-психиатр"),
        "head": ("Зав. отделением", "Зав. отд.", "Зав отд"),
    }

    BLOCK_ALIASES: Dict[str, Sequence[str]] = {
        "complaints": ("Жалобы на момент осмотра", "Жалобы при поступлении", "Жалобы"),
        "life_anamnesis": ("Анамнез жизни",),
        "disease_anamnesis": ("Анамнез заболевания",),
        "mental_status": ("Психический статус при поступлении", "Психический статус"),
        "somatic_status": ("Сомато-неврологический статус", "Соматический статус"),
        "examination_plan": ("План обследования",),
        "treatment_plan": ("План лечения", "Лечение"),
        "diagnosis": ("Диагноз", "был выставлен диагноз"),
        "epidemiology": ("Эпидемиологический анамнез",),
    }

    SECTION_MARKERS: Sequence[str] = (
        "Дата, время",
        "История болезни №",
        "Ф.И.О.",
        "ФИО",
        "Год рождения",
        "Дата рождения",
        "Зарегистрирован",
        "Проживает",
        "На учёте у психиатров",
        "На учете у психиатров",
        "Работает в организации",
        "Место работы",
        "Должность",
        "Больничный лист",
        "Оформление инвалидности",
        "Направление от РВК",
        "В 3 отделение КДП поступает",
        "Жалобы на момент осмотра",
        "Жалобы при поступлении",
        "Жалобы",
        "Анамнез жизни",
        "Анамнез заболевания",
        "Психический статус при поступлении",
        "Психический статус",
        "Сомато-неврологический статус",
        "Соматический статус",
        "План обследования",
        "План лечения",
        "На основании данных",
        "Диагноз",
        "Эпидемиологический анамнез",
        "Результаты обследований",
        "Результаты исследований",
        "ЭЭГ",
        "ЭПИ",
        "За время лечения",
        "Рекомендовано",
        "Лечение",
        "Экспертный анамнез",
        "Врач психиатр",
        "Врач-психиатр",
        "Зав. отделением",
        "Зав. отд.",
        "Зав отд",
    )

    @staticmethod
    def _detect_document_kind(text: str) -> str:
        """Определить тип входного первичного документа для статуса/диагностики.

        Это не влияет на схему данных: и направление, и первичный осмотр
        разбираются в один PatientData и затем используются всеми выбранными
        документами.
        """
        low = normalize_match(text)
        if "первичный осмотр" in low:
            return "первичный осмотр"
        if "направление на госпитализацию" in low or "госпитализируется по направлению" in low:
            return "направление на госпитализацию"
        if "в 3 отделение кдп поступает" in low and "анамнез жизни" in low and "психический статус" in low:
            return "первичный документ пациента"
        return "первичный документ"

    def parse_docx(self, path: str | Path) -> PatientData:
        return self.parse_text(extract_docx_text(path))

    def parse_text(self, text: str) -> PatientData:
        text = normalize_text(text)
        data = PatientData()
        data.input_document_kind = self._detect_document_kind(text)

        for field_name, aliases in self.FIELD_ALIASES.items():
            value = self._extract_inline(text, aliases)
            if value:
                setattr(data, field_name, value)

        for field_name, aliases in self.BLOCK_ALIASES.items():
            value = self._extract_block(text, aliases)
            if value:
                setattr(data, field_name, value)

        data.admission_date = self._extract_admission_date(text)

        if not data.diagnosis:
            diagnosis = self._extract_after_phrase(text, r"был\s+выставлен\s+диагноз\s*[:.]?")
            if diagnosis:
                data.diagnosis = diagnosis

        for key, value in list(asdict(data).items()):
            if isinstance(value, str) and looks_like_label(value):
                setattr(data, key, "")

        if not data.registered:
            data.registered = "Нижний Новгород"
        if not data.epidemiology:
            data.epidemiology = (
                "со слов пациента, за пределы Нижегородской области в течение трёх предыдущих месяцев "
                "не выезжал, в контакте с инфекционными больными не был. Венерические заболевания, "
                "туберкулёз, вирусные гепатиты отрицает."
            )
        if not data.doctor:
            data.doctor = "Балаганин С.В"
        if not data.head:
            data.head = "Можарова Е.А."

        for field_name in data.missing_critical_fields():
            data.warnings.append(f"Не найдено критическое поле: {field_name}")
        for field_name in data.missing_recommended_fields():
            data.warnings.append(f"Не найдено рекомендуемое поле: {field_name}")

        return data

    def _extract_inline(self, text: str, aliases: Sequence[str]) -> str:
        lines = text.splitlines()
        for alias in aliases:
            alias_re = self._alias_pattern(alias)
            for line in lines:
                line_norm = normalize_text(line)
                m = re.search(alias_re, line_norm, flags=re.IGNORECASE)
                if not m:
                    continue
                value = line_norm[m.end():]
                value = re.sub(r"^\s*[:№N#.-]*\s*", "", value)
                value = clean_value(value)
                if value and not looks_like_label(value):
                    return value
        return ""

    def _extract_block(self, text: str, aliases: Sequence[str]) -> str:
        for alias in aliases:
            start = self._find_alias_span(text, alias)
            if not start:
                continue
            value_start = start[1]
            remainder = text[value_start:]
            remainder = re.sub(r"^\s*[:.-]*\s*", "", remainder)
            consumed = len(text[value_start:]) - len(remainder)
            value_start += consumed
            value_end = self._find_next_marker_pos(text, value_start, aliases)
            raw = clean_value(text[value_start:value_end])
            raw = self._remove_template_noise(raw)
            if raw and not looks_like_label(raw):
                return raw
        return ""

    def _extract_after_phrase(self, text: str, phrase_pattern: str) -> str:
        m = re.search(phrase_pattern, text, flags=re.IGNORECASE)
        if not m:
            return ""
        value_start = m.end()
        value_end = self._find_next_marker_pos(text, value_start, ())
        return clean_value(text[value_start:value_end])

    def _extract_admission_date(self, text: str) -> str:
        lines = text.splitlines()
        for line in lines[:12]:
            if "дата" in normalize_match(line) or re.search(r"\d{2}\.\d{2}\.\d{4}", line):
                m = re.search(r"\b(\d{2}\.\d{2}\.\d{4})\b", line)
                if m:
                    return m.group(1)
        m = re.search(r"\b(\d{2}\.\d{2}\.\d{4})\b", text)
        return m.group(1) if m else ""

    def _find_alias_span(self, text: str, alias: str) -> Optional[Tuple[int, int]]:
        pattern = self._alias_pattern(alias)
        m = re.search(pattern, text, flags=re.IGNORECASE)
        if not m:
            m = re.search(re.escape(alias), text, flags=re.IGNORECASE)
        return (m.start(), m.end()) if m else None

    def _find_next_marker_pos(self, text: str, start_pos: int, current_aliases: Sequence[str]) -> int:
        best = len(text)
        current_norm = {normalize_match(a) for a in current_aliases}
        for marker in self.SECTION_MARKERS:
            if normalize_match(marker) in current_norm:
                continue
            pattern = self._alias_pattern(marker)
            m = re.search(pattern, text[start_pos:], flags=re.IGNORECASE)
            if not m:
                continue
            pos = start_pos + m.start()
            before = text[max(0, pos - 2):pos]
            if pos == 0 or "\n" in before or marker.lower().startswith("на основании"):
                best = min(best, pos)
        return best

    @staticmethod
    def _alias_pattern(alias: str) -> str:
        alias = re.escape(alias)
        alias = alias.replace(r"\ ", r"\s+")
        alias = alias.replace("ё", "[её]").replace("Ё", "[ЕЁ]")
        return alias

    @staticmethod
    def _remove_template_noise(text: str) -> str:
        noisy_patterns = [
            r"сюда\s+подставлять[^\n]*",
            r"сюда\s+подставляется[^\n]*",
            r"выбирается\s+в\s+ui",
        ]
        for pat in noisy_patterns:
            text = re.sub(pat, "", text, flags=re.IGNORECASE)
        return normalize_text(text)


# -----------------------------------------------------------------------------
# Редактирование DOCX-шаблонов
# -----------------------------------------------------------------------------

PRIMARY_MARKERS = [
    "Дата, время", "История болезни №", "Ф.И.О.", "Год рождения", "Зарегистрирован",
    "На учёте у психиатров", "Работает в организации", "Должность", "Больничный лист",
    "Оформление инвалидности", "Направление от РВК", "В 3 отделение КДП поступает",
    "Жалобы на момент осмотра", "Анамнез жизни", "Анамнез заболевания",
    "Психический статус", "Соматический статус", "План обследования", "План лечения",
    "На основании данных", "Эпидемиологический анамнез", "Врач психиатр", "Зав. отделением",
]

DISCHARGE_MARKERS = [
    "Дата, время", "г.р.,", "Находился на лечении", "В 3 отделение КДП поступает",
    "Жалобы при поступлении", "Анамнез жизни", "Анамнез заболевания", "Психический статус при поступлении",
    "Сомато-неврологический статус", "Результаты обследований", "ОАК", "ОАМ", "RW", "HCV",
    "HBsAg", "ВИЧ", "Биохимия крови", "Глюкоза крови", "Кал на яйца глист", "Флюорография",
    "ЭКГ", "ЭЭГ", "ЭПИ", "За время лечения", "Рекомендовано", "Лечение", "Экспертный анамнез",
    "Зав. отд.",
]

COMMISSION_MARKERS = [
    "г.р.,", "В 3 отделение КДП поступает", "Жалобы при поступлении", "Анамнез жизни",
    "Анамнез заболевания", "Психический статус при поступлении", "Соматический статус",
    "Результаты исследований", "ОАК", "ОАМ", "RW", "HCV", "HBsAg", "ВИЧ",
    "Биохимия крови", "Глюкоза крови", "Кал на яйца глист", "Флюорография", "ЭКГ", "ЭЭГ", "ЭПИ",
    "На основании данных", "Лечение", "Экспертный анамнез", "Эпидемиологический анамнез",
    "Зам глав врача", "Зав отд", "Врач-психиатр",
]

VK_MSE_MARKERS = [
    "Ф.И.О", "Год рождения", "Проживает", "Место работы", "Диагноз", "Жалобы",
    "Анамнез жизни", "Анамнез заболевания", "Психический статус при поступлении", "ЭПИ",
    "Сомато-неврологический статус", "Получает лечение", "Прогноз восстановления", "Цель направления",
    "Зав. отделением", "Лечащий врач", "Выписка из ПРОТОКОЛА", "Клинический и трудовой прогноз", "Решение ВК",
]

SICK_LEAVE_VK_MARKERS = [
    "Ф.И.О", "Год рождения", "Проживает", "Место работы", "Находится на лечении", "Диагноз", "Жалобы",
    "Анамнез жизни", "Анамнез заболевания", "Психический статус при поступлении", "ЭПИ",
    "Сомато-неврологический статус", "Получает лечение", "Прогноз восстановления", "Цель направления",
    "Зав. отделением", "Лечащий врач", "Выписка из ПРОТОКОЛА", "Клинический и трудовой прогноз", "Решение ВК",
]

RVK_MARKERS = [
    "О СОСТОЯНИИ", "История болезни №", "Ф.И.О.", "Год рождения", "Проживает", "Место работы",
    "Находился на обследовании", "Госпитализируется", "На учёте", "Жалобы", "Анамнез жизни",
    "Анамнез заболевания", "Психический статус", "Сомато-неврологический статус", "Результаты обследований",
    "ОАК", "ОАМ", "RW", "HCV", "HBsAg", "ВИЧ", "Биохимия крови", "Глюкоза крови",
    "Кал на яйца глист", "Флюорография", "ЭКГ", "ЭЭГ", "ЭПИ", "Исходя из выше изложенного", "Диагноз",
    "Зам. гл. врача", "Зав. отделением", "Врач-психиатр",
]


class DocxBlockEditor:
    def __init__(self, doc: DocxDocument):
        self.doc = doc

    @property
    def paragraphs(self) -> List[Paragraph]:
        return list(self.doc.paragraphs)

    def replace_first_matching_paragraph(self, markers: Sequence[str], text: str) -> bool:
        idx = self.find_paragraph_index(markers)
        if idx is None:
            return False
        set_paragraph_text(self.paragraphs[idx], text)
        return True

    def replace_all_matching_paragraphs(self, markers: Sequence[str], text: str) -> int:
        count = 0
        for paragraph in self.paragraphs:
            text_norm = normalize_match(paragraph.text)
            if any(paragraph_matches_marker(text_norm, marker) for marker in markers):
                set_paragraph_text(paragraph, text)
                count += 1
        return count

    def remove_all_matching_paragraphs(self, markers: Sequence[str]) -> int:
        """Удалить все абзацы, которые начинаются с указанных маркеров.

        Используется для ЭПИ: если файл ЭПИ не выбран, в итоговых документах
        не должно оставаться ни строки «ЭПИ - ...», ни шаблонной подсказки.
        """
        count = 0
        for paragraph in list(self.paragraphs):
            text_norm = normalize_match(paragraph.text)
            if any(paragraph_matches_marker(text_norm, marker) for marker in markers):
                remove_paragraph(paragraph)
                count += 1
        return count

    def replace_first_matching_regex(self, pattern: str, text: str) -> bool:
        rx = re.compile(pattern, flags=re.IGNORECASE)
        for paragraph in self.paragraphs:
            if rx.search(normalize_text(paragraph.text)):
                set_paragraph_text(paragraph, text)
                return True
        return False

    def replace_block(
        self,
        start_markers: Sequence[str],
        label: str,
        value: str,
        all_markers: Sequence[str],
        *,
        allow_empty: bool = False,
        preserve_when_empty: bool = True,
    ) -> bool:
        value = normalize_text(value)
        if not value and preserve_when_empty and not allow_empty:
            return False

        idx = self.find_paragraph_index(start_markers)
        if idx is None:
            return False

        paragraphs = self.paragraphs
        end_idx = self.find_next_marker_index(idx + 1, all_markers, exclude=start_markers)
        if end_idx is None:
            end_idx = idx + 1

        for remove_idx in range(end_idx - 1, idx, -1):
            remove_paragraph(paragraphs[remove_idx])

        lines = [line.strip() for line in value.splitlines() if line.strip()]
        if not lines:
            lines = [""]

        first = f"{label} {lines[0]}".rstrip()
        set_paragraph_text(paragraphs[idx], first)

        anchor = paragraphs[idx]
        for line in lines[1:]:
            anchor = insert_paragraph_after(anchor, line)
        return True

    def find_paragraph_index(self, markers: Sequence[str]) -> Optional[int]:
        for i, paragraph in enumerate(self.paragraphs):
            text = normalize_match(paragraph.text)
            if any(paragraph_matches_marker(text, marker) for marker in markers):
                return i
        return None

    def find_next_marker_index(self, start: int, markers: Sequence[str], *, exclude: Sequence[str] = ()) -> Optional[int]:
        excluded = {normalize_match(m) for m in exclude}
        for i, paragraph in enumerate(self.paragraphs[start:], start=start):
            text = normalize_match(paragraph.text)
            if not text:
                continue
            for marker in markers:
                norm_marker = normalize_match(marker)
                if norm_marker in excluded:
                    continue
                if paragraph_matches_marker(text, marker):
                    return i
        return None


def paragraph_matches_marker(normalized_paragraph_text: str, marker: str) -> bool:
    marker = normalize_match(marker)
    if not marker or not normalized_paragraph_text:
        return False
    text = normalized_paragraph_text.lstrip()
    if text.startswith(marker):
        if len(text) == len(marker):
            return True
        # Важно: короткие маркеры вроде «ЭПИ» не должны срабатывать на «ЭПИКРИЗ».
        # Разрешаем совпадение только если после маркера идёт разделитель.
        next_char = text[len(marker)]
        return not (next_char.isalnum() or next_char == "_")
    if marker == "зарегистрирован по адресу" and " зарегистрирован по адресу" in text:
        return True
    return False


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    try:
        paragraph.clear()
    except AttributeError:
        for run in list(paragraph.runs):
            run._element.getparent().remove(run._element)
    paragraph.add_run(text)


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_exact_paragraphs(doc: DocxDocument, values: Sequence[str]) -> int:
    """Remove standalone service paragraphs such as a trailing 'ЭЭГ'."""
    normalized_values = {normalize_match(value) for value in values}
    count = 0
    for paragraph in list(iter_all_paragraphs(doc)):
        if normalize_match(paragraph.text) in normalized_values:
            remove_paragraph(paragraph)
            count += 1
    return count


EPI_WORD_RE = re.compile(r"(?<![A-Za-zА-ЯЁа-яё])ЭПИ(?![A-Za-zА-ЯЁа-яё])", re.IGNORECASE)


def iter_all_paragraphs(parent) -> Iterable[Paragraph]:
    """Все абзацы документа, включая абзацы внутри таблиц."""
    if isinstance(parent, DocxDocument):
        for paragraph in parent.paragraphs:
            yield paragraph
        for table in parent.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_all_paragraphs(cell)
    elif isinstance(parent, _Cell):
        for paragraph in parent.paragraphs:
            yield paragraph
        for table in parent.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_all_paragraphs(cell)


def remove_epi_from_text(text: str) -> str:
    """Удаляет именно слово/аббревиатуру «ЭПИ», не трогая «ЭПИКРИЗ»."""
    cleaned = EPI_WORD_RE.sub("", text or "")
    cleaned = re.sub(r"\(\s*\)", "", cleaned)
    cleaned = re.sub(r"\s*,\s*", ", ", cleaned)
    cleaned = re.sub(r",\s*,+", ", ", cleaned)
    cleaned = re.sub(r":\s*,\s*", ": ", cleaned)
    cleaned = re.sub(r",\s*([.;:!?])", r"\1", cleaned)
    cleaned = re.sub(r"\s+([,.;:!?])", r"\1", cleaned)
    cleaned = re.sub(r"[-–—:]\s*$", "", cleaned)
    cleaned = re.sub(r"\s{2,}", " ", cleaned)
    return cleaned.strip()


def remove_epi_mentions_from_document(doc: DocxDocument) -> None:
    """Если ЭПИ-файл не выбран, удаляет все видимые упоминания ЭПИ из результата."""
    for paragraph in list(iter_all_paragraphs(doc)):
        text = paragraph.text
        if not EPI_WORD_RE.search(text or ""):
            continue
        cleaned = remove_epi_from_text(text)
        if cleaned.strip(" ,.;:-–—()"):
            set_paragraph_text(paragraph, cleaned)
        else:
            remove_paragraph(paragraph)


# -----------------------------------------------------------------------------
# Род пациента: общий слой для дневников и медицинских документов
# -----------------------------------------------------------------------------

GENDER_ADAPTED_PATIENT_FIELDS = (
    "complaints",
    "life_anamnesis",
    "disease_anamnesis",
    "mental_status",
    "somatic_status",
    "treatment_plan",
    "epidemiology",
    "admission",
    "psych_account",
    "epi_text",
)


def patient_gender(data: PatientData) -> str | None:
    """Определить род пациента по первой части ФИО, как в заполнителе дневников."""
    return detect_gender_from_patient_name(data.fio or data.output_fio or "")


def adapt_patient_data_to_gender(data: PatientData) -> PatientData:
    """Вернуть копию данных, где клинические текстовые блоки согласованы с родом пациента.

    Диагноз, ФИО, адрес, даты, подписи и служебные реквизиты не трогаем: они
    не являются текстом о пациенте и не должны портиться морфологическим проходом.
    """
    gender = patient_gender(data)
    if gender not in {"male", "female"}:
        return data

    adapted = copy.deepcopy(data)
    for field_name in GENDER_ADAPTED_PATIENT_FIELDS:
        value = getattr(adapted, field_name, "")
        if not isinstance(value, str) or not value:
            continue
        new_value, _changed = adapt_text_to_patient_gender(value, gender)
        setattr(adapted, field_name, new_value)
    return adapted


def adapt_document_to_patient_gender(doc: DocxDocument, data: PatientData) -> None:
    """Применить ту же муж/жен коррекцию к итоговому DOCX.

    Это нужно для шаблонных фраз самих документов: например,
    «Находился на лечении...» -> «Находилась на лечении...» для женской фамилии.
    Диагнозные строки оставляем как есть: диагноз — отдельная медицинская сущность,
    а не грамматическое описание пациента.
    """
    gender = patient_gender(data)
    if gender not in {"male", "female"}:
        return

    for paragraph in list(iter_all_paragraphs(doc)):
        original = paragraph.text or ""
        if not original.strip():
            continue
        # Защита от порчи фраз вида «установлен диагноз: ...» и названий МКБ.
        # Клинические описания вокруг этих строк уже адаптированы на уровне данных.
        if "диагноз" in normalize_match(original):
            continue
        updated, changed = adapt_text_to_patient_gender(original, gender)
        if changed and updated != original:
            set_paragraph_text(paragraph, updated)


def finalize_medical_document(doc: DocxDocument, data: PatientData) -> None:
    """Общие финальные правки перед сохранением любого медицинского документа."""
    adapt_document_to_patient_gender(doc, data)
    if not data.epi_text:
        remove_epi_mentions_from_document(doc)


def format_date_with_russian_year_suffix(value: str) -> str:
    """Return UI date text with exactly one trailing "г." for document headers."""
    value = normalize_text(value)
    if not value:
        return ""
    value = re.sub(r"\s*г\.?\s*$", "", value, flags=re.IGNORECASE).strip()
    return f"{value}г."


# -----------------------------------------------------------------------------
# Расчёт сроков лечения
# -----------------------------------------------------------------------------

def calculate_inclusive_treatment_days(admission_date: str, commission_date: str) -> int | None:
    """Количество дней лечения от даты поступления до даты комиссии включительно."""
    start = parse_date(admission_date)
    finish = parse_date(commission_date)
    if not start or not finish:
        return None
    days = (finish.date() - start.date()).days + 1
    return days if days >= 1 else None


def format_military_commissariat_area(value: str) -> str:
    """Normalize user input for phrase: военного комиссариата <...> района."""
    value = normalize_text(value)
    if not value:
        return ""
    low = value.lower().replace("ё", "е")
    if low.endswith(" района") or low.endswith(" району") or low.endswith(" район"):
        return value
    return f"{value} района"


def treatment_period_text(admission_date: str, commission_date: str) -> str:
    days = calculate_inclusive_treatment_days(admission_date, commission_date)
    if admission_date and days:
        return f"Находится на лечении с {admission_date} ({days} дней)"
    if admission_date:
        return f"Находится на лечении с {admission_date} (всего дней)"
    return "Находится на лечении с (всего дней)"


# -----------------------------------------------------------------------------
# Генерация документов
# -----------------------------------------------------------------------------

class MedicalDocumentRenderer:
    def render(self, kind: str, template_path: str | Path, output_path: str | Path, data: PatientData) -> None:
        methods = {
            "primary": self.render_primary,
            "discharge": self.render_discharge,
            "commission": self.render_commission,
            "vk_mse": self.render_vk_mse,
            "sick_leave_vk": self.render_sick_leave_vk,
            "rvk": self.render_rvk,
        }
        gender_adapted_data = adapt_patient_data_to_gender(data)
        methods[kind](template_path, output_path, gender_adapted_data)

    def render_primary(self, template_path: str | Path, output_path: str | Path, data: PatientData) -> None:
        doc = Document(str(template_path))
        editor = DocxBlockEditor(doc)

        header_date = data.admission_date or "Дата"
        header = f"{header_date} 10:00      Первичный осмотр с зав. отд. Можаровой Е.А."
        editor.replace_first_matching_paragraph(["Дата, время"], header)
        editor.replace_block(["История болезни №"], "История болезни №", data.case_number, PRIMARY_MARKERS, preserve_when_empty=False, allow_empty=True)
        editor.replace_block(["Ф.И.О.", "ФИО"], "Ф.И.О.:", data.fio, PRIMARY_MARKERS, allow_empty=True)
        editor.replace_block(["Год рождения", "Дата рождения"], "Год рождения:", data.birth, PRIMARY_MARKERS, allow_empty=True)
        editor.replace_block(["Зарегистрирован"], "Зарегистрирован:", data.registered, PRIMARY_MARKERS, allow_empty=True)
        editor.replace_block(["На учёте у психиатров", "На учете у психиатров"], "На учёте у психиатров состоит/не состоит:", data.psych_account, PRIMARY_MARKERS)
        editor.replace_block(["Работает в организации"], "Работает в организации:", data.work_org, PRIMARY_MARKERS)
        editor.replace_block(["Должность"], "Должность:", data.position, PRIMARY_MARKERS)
        editor.replace_block(["Больничный лист"], "Больничный лист нужен/ не нужен:", data.sick_leave, PRIMARY_MARKERS)
        editor.replace_block(["Оформление инвалидности"], "Оформление инвалидности нужно/ не нужно:", data.disability, PRIMARY_MARKERS)
        editor.replace_block(["Направление от РВК"], "Направление от РВК да/нет:", data.rvk_referral, PRIMARY_MARKERS)
        editor.replace_block(["В 3 отделение КДП поступает"], "В 3 отделение КДП поступает:", data.admission, PRIMARY_MARKERS)
        editor.replace_block(["Жалобы на момент осмотра", "Жалобы"], "Жалобы на момент осмотра:", data.complaints, PRIMARY_MARKERS)
        editor.replace_block(["Анамнез жизни"], "Анамнез жизни:", data.life_anamnesis, PRIMARY_MARKERS)
        editor.replace_block(["Анамнез заболевания"], "Анамнез заболевания:", data.disease_anamnesis, PRIMARY_MARKERS)
        editor.replace_block(["Психический статус"], "Психический статус:", data.mental_status, PRIMARY_MARKERS)
        editor.replace_block(["Соматический статус"], "Соматический статус:", data.somatic_status, PRIMARY_MARKERS)
        editor.replace_block(["План обследования"], "План обследования:", data.examination_plan, PRIMARY_MARKERS)
        editor.replace_block(["План лечения"], "План лечения:", data.treatment_plan, PRIMARY_MARKERS)

        diagnosis_sentence = ""
        if data.diagnosis:
            diagnosis_sentence = (
                "На основании данных анамнеза жизни и заболевания, психического статуса, "
                f"данных клинических исследований установлен диагноз: {data.diagnosis}"
            )
        editor.replace_block(["На основании данных", "Диагноз"], "", diagnosis_sentence, PRIMARY_MARKERS)
        editor.replace_block(["Эпидемиологический анамнез"], "Эпидемиологический анамнез:", data.epidemiology, PRIMARY_MARKERS)
        editor.replace_block(["Врач психиатр", "Врач-психиатр"], "Врач психиатр", data.doctor, PRIMARY_MARKERS, allow_empty=True)
        editor.replace_block(["Зав. отделением", "Зав. отд."], "Зав. отделением", data.head, PRIMARY_MARKERS, allow_empty=True)
        finalize_medical_document(doc, data)
        doc.save(str(output_path))

    def render_discharge(self, template_path: str | Path, output_path: str | Path, data: PatientData) -> None:
        doc = Document(str(template_path))
        editor = DocxBlockEditor(doc)
        dates = data.lab_dates()

        header_date = data.discharge_date or data.admission_date or "Дата, время"
        header = f"{header_date}      Выписной эпикриз № {data.case_number}".rstrip()
        editor.replace_first_matching_paragraph(["Дата, время"], header)
        person_line = f"{data.fio}, {data.birth} г.р., зарегистрирован по адресу: {data.registered or 'Н. Новгород'}".strip(" ,")
        editor.replace_first_matching_paragraph(["г.р.,", "зарегистрирован по адресу"], person_line)
        period = f"Находился на лечении в ГБУЗ НО «НКЦПЗ» диспансер №2 с {data.admission_date} по {data.discharge_date}".strip()
        editor.replace_first_matching_paragraph(["Находился на лечении"], period)

        editor.replace_block(["В 3 отделение КДП поступает"], "В 3 отделение КДП поступает", data.admission, DISCHARGE_MARKERS)
        editor.replace_block(["Жалобы при поступлении", "Жалобы"], "Жалобы при поступлении:", data.complaints, DISCHARGE_MARKERS)
        editor.replace_block(["Анамнез жизни"], "Анамнез жизни:", data.life_anamnesis, DISCHARGE_MARKERS)
        editor.replace_block(["Анамнез заболевания"], "Анамнез заболевания:", data.disease_anamnesis, DISCHARGE_MARKERS)
        editor.replace_block(["Психический статус при поступлении", "Психический статус"], "Психический статус при поступлении:", data.mental_status, DISCHARGE_MARKERS)
        editor.replace_block(["Сомато-неврологический статус", "Соматический статус"], "Сомато-неврологический статус:", data.somatic_status, DISCHARGE_MARKERS)
        self._replace_lab_lines(editor, dates)
        if data.epi_text:
            editor.replace_block(["ЭПИ"], "ЭПИ –", data.epi_text, DISCHARGE_MARKERS)
        else:
            editor.remove_all_matching_paragraphs(["ЭПИ"])
        if data.treatment_plan:
            editor.replace_block(["Лечение"], "Лечение:", data.treatment_plan, DISCHARGE_MARKERS)
        signature = f"  Зав. отд. {data.head}                                                                                                 Врач-психиатр\t{data.doctor}"
        editor.replace_first_matching_paragraph(["Зав. отд.", "Врач-психиатр"], signature)
        finalize_medical_document(doc, data)
        doc.save(str(output_path))

    def render_commission(self, template_path: str | Path, output_path: str | Path, data: PatientData) -> None:
        """Комиссионный: шапку/первую строку оставляем, заполняем клиническую часть ниже."""
        doc = Document(str(template_path))
        editor = DocxBlockEditor(doc)
        dates = data.lab_dates()

        if data.commission_date or data.commission_number:
            commission_date = format_date_with_russian_year_suffix(data.commission_date or data.admission_date)
            header = (
                f"{commission_date} 10:00      "
                f"Совместный осмотр с зам глав врача Зуйковой А.А. № {data.commission_number}"
            ).rstrip()
            editor.replace_first_matching_regex(r"Совместный\s+осмотр", header)

        person_line = f"{data.fio}, {data.birth} г.р., зарегистрирован по адресу: {data.registered or 'Н. Новгород'}".strip(" ,")
        editor.replace_first_matching_paragraph(["г.р.,", "зарегистрирован по адресу"], person_line)
        editor.replace_block(["В 3 отделение КДП поступает"], "В 3 отделение КДП поступает", data.admission, COMMISSION_MARKERS)
        editor.replace_block(["Жалобы при поступлении", "Жалобы"], "Жалобы при поступлении:", data.complaints, COMMISSION_MARKERS)
        editor.replace_block(["Анамнез жизни"], "Анамнез жизни:", data.life_anamnesis, COMMISSION_MARKERS)
        editor.replace_block(["Анамнез заболевания"], "Анамнез заболевания:", data.disease_anamnesis, COMMISSION_MARKERS)
        editor.replace_block(["Психический статус при поступлении", "Психический статус"], "Психический статус при поступлении:", data.mental_status, COMMISSION_MARKERS)
        editor.replace_block(["Соматический статус", "Сомато-неврологический статус"], "Соматический статус:", data.somatic_status, COMMISSION_MARKERS)

        lab_lines = [
            (["ОАК"], f"ОАК ({dates['day1']}) - в норме"),
            (["ОАМ"], f"ОАМ ({dates['day1']}) - в норме"),
            (["RW"], f"RW (от {dates['day1']}) - в норме"),
            (["HCV"], f"HCV (от {dates['day1']}) - в норме"),
            (["HBsAg"], f"HBsAg (от {dates['day1']}) - в норме"),
            (["ВИЧ"], f"ВИЧ (от {dates['day2']}) - в норме"),
            (["Биохимия крови"], f"Биохимия крови ({dates['day1']}) - в норме"),
            (["Глюкоза крови"], f"Глюкоза крови ({dates['day1']}) – 3,40 ммоль/л"),
            (["Кал на яйца глист"], f"Кал на яйца глист ({dates['day1']}) - не обнаружены."),
            (["Флюорография"], f"Флюорография ({dates['flg']}) - патологии не выявлено."),
            (["ЭКГ"], f"ЭКГ ({dates['day1']}) – ритм синусовый, ЭОС нормальная."),
        ]
        for markers, text in lab_lines:
            editor.replace_first_matching_paragraph(markers, text)
        if data.epi_text:
            editor.replace_block(["ЭПИ"], f"ЭПИ ({dates['day2']}) -", data.epi_text, COMMISSION_MARKERS)
        else:
            editor.remove_all_matching_paragraphs(["ЭПИ"])

        diagnosis_sentence = ""
        if data.diagnosis:
            diagnosis_sentence = (
                "На основании данных анамнеза жизни и заболевания, психического статуса, "
                f"данных клинических исследований установлен диагноз: {data.diagnosis}"
            )
        editor.replace_block(["На основании данных", "Диагноз"], "", diagnosis_sentence, COMMISSION_MARKERS)
        editor.replace_block(["Лечение"], "Лечение:", data.treatment_plan, COMMISSION_MARKERS)
        editor.replace_block(["Эпидемиологический анамнез"], "Эпидемиологический анамнез:", data.epidemiology, COMMISSION_MARKERS)
        finalize_medical_document(doc, data)
        doc.save(str(output_path))

    def render_vk_mse(self, template_path: str | Path, output_path: str | Path, data: PatientData) -> None:
        """ВК на МСЭ: верхнюю шапку и текст решения не трогаем, заполняем поля пациента."""
        doc = Document(str(template_path))
        editor = DocxBlockEditor(doc)

        if data.vk_date:
            editor.replace_first_matching_regex(r"^\s*\.?\s*\d{4}\s*$", data.vk_date)
        if data.vk_protocol_number:
            editor.replace_first_matching_paragraph(["Выписка из ПРОТОКОЛА"], f"Выписка из ПРОТОКОЛА № {data.vk_protocol_number}")
        if data.vk_protocol_date:
            editor.replace_first_matching_regex(r"^\s*От\s*\.?\s*\d{4}\s*г\.?\s*$", f"От {data.vk_protocol_date} г.")

        editor.replace_all_matching_paragraphs(["Ф.И.О", "Ф.И.О:"], f"Ф.И.О: {data.fio}")
        editor.replace_all_matching_paragraphs(["Год рождения"], f"Год рождения: {data.birth}")
        editor.replace_all_matching_paragraphs(["Проживает"], f"Проживает: {data.registered or 'Н. Новгород'}")
        vk_work_parts = [
            (data.vk_mse_work_org or data.work_org).strip(),
            (data.vk_mse_position or data.position).strip(),
        ]
        vk_work_line = ", ".join(part for part in vk_work_parts if part) or "не работает"
        editor.replace_all_matching_paragraphs(["Место работы"], f"Место работы: {vk_work_line}")
        editor.replace_all_matching_paragraphs(["Диагноз"], f"Диагноз: {data.diagnosis}")

        editor.replace_block(["Жалобы"], "Жалобы:", data.complaints, VK_MSE_MARKERS)
        editor.replace_block(["Анамнез жизни"], "Анамнез жизни:", data.life_anamnesis, VK_MSE_MARKERS)
        editor.replace_block(["Анамнез заболевания"], "Анамнез заболевания:", data.disease_anamnesis, VK_MSE_MARKERS)
        editor.replace_block(["Психический статус при поступлении", "Психический статус"], "Психический статус при поступлении:", data.mental_status, VK_MSE_MARKERS)
        if data.epi_text:
            editor.replace_block(["ЭПИ"], "ЭПИ -", data.epi_text, VK_MSE_MARKERS)
        else:
            editor.remove_all_matching_paragraphs(["ЭПИ"])
        editor.replace_block(["Сомато-неврологический статус", "Соматический статус"], "Сомато-неврологический статус:", data.somatic_status, VK_MSE_MARKERS)
        editor.replace_block(["Получает лечение"], "Получает лечение:", data.treatment_plan, VK_MSE_MARKERS)
        finalize_medical_document(doc, data)
        doc.save(str(output_path))

    def render_sick_leave_vk(self, template_path: str | Path, output_path: str | Path, data: PatientData) -> None:
        """ВК больничный: отдельная форма ВК для продления лечения/больничного."""
        doc = Document(str(template_path))
        editor = DocxBlockEditor(doc)

        # Верхняя дата, номер протокола и дата протокола работают так же, как в ВК на МСЭ.
        if data.sick_leave_vk_date:
            editor.replace_first_matching_regex(r"^\s*\.?\s*\d{4}\s*$", data.sick_leave_vk_date)
        if data.sick_leave_vk_protocol_number:
            editor.replace_first_matching_paragraph(["Выписка из ПРОТОКОЛА"], f"Выписка из ПРОТОКОЛА № {data.sick_leave_vk_protocol_number}")
        if data.sick_leave_vk_protocol_date:
            editor.replace_first_matching_regex(r"^\s*От\s*\.?\s*\d{4}\s*г\.?\s*$", f"От {data.sick_leave_vk_protocol_date} г.")

        work_position = data.sick_leave_vk_work_position or ", ".join(part for part in [data.work_org, data.position] if part).strip(", ") or "не работает"
        treatment_line = treatment_period_text(data.admission_date, data.sick_leave_vk_commission_date or data.sick_leave_vk_date)

        editor.replace_all_matching_paragraphs(["Ф.И.О", "Ф.И.О:"], f"Ф.И.О: {data.fio}")
        editor.replace_all_matching_paragraphs(["Год рождения"], f"Год рождения: {data.birth}")
        editor.replace_all_matching_paragraphs(["Проживает"], f"Проживает: {data.registered or 'Н. Новгород'}")
        editor.replace_all_matching_paragraphs(["Место работы"], f"Место работы, должность: {work_position}")
        editor.replace_all_matching_paragraphs(["Находится на лечении"], treatment_line)
        editor.replace_all_matching_paragraphs(["Диагноз"], f"Диагноз: {data.diagnosis}")

        editor.replace_block(["Жалобы"], "Жалобы:", data.complaints, SICK_LEAVE_VK_MARKERS)
        editor.replace_block(["Анамнез жизни"], "Анамнез жизни:", data.life_anamnesis, SICK_LEAVE_VK_MARKERS)
        editor.replace_block(["Анамнез заболевания"], "Анамнез заболевания:", data.disease_anamnesis, SICK_LEAVE_VK_MARKERS)
        editor.replace_block(["Психический статус при поступлении", "Психический статус"], "Психический статус при поступлении:", data.mental_status, SICK_LEAVE_VK_MARKERS)
        if data.epi_text:
            editor.replace_block(["ЭПИ"], "ЭПИ -", data.epi_text, SICK_LEAVE_VK_MARKERS)
        else:
            editor.remove_all_matching_paragraphs(["ЭПИ"])
        editor.replace_block(["Сомато-неврологический статус", "Соматический статус"], "Сомато-неврологический статус:", data.somatic_status, SICK_LEAVE_VK_MARKERS)
        editor.replace_block(["Получает лечение"], "Получает лечение:", data.treatment_plan, SICK_LEAVE_VK_MARKERS)
        finalize_medical_document(doc, data)
        doc.save(str(output_path))

    def render_rvk(self, template_path: str | Path, output_path: str | Path, data: PatientData) -> None:
        """Акт для РВК: нормативную шапку оставляем, заполняем поля ниже."""
        doc = Document(str(template_path))
        editor = DocxBlockEditor(doc)
        dates = data.lab_dates()

        act_number = data.rvk_act_number or data.case_number
        editor.replace_first_matching_paragraph(["О СОСТОЯНИИ"], f"О СОСТОЯНИИ ЗДОРОВЬЯ ГРАЖДАНИНА № {act_number}".rstrip())
        editor.replace_block(["История болезни №"], "История болезни №", data.case_number, RVK_MARKERS, preserve_when_empty=False, allow_empty=True)
        editor.replace_block(["Ф.И.О.", "ФИО"], "Ф.И.О.:", data.fio, RVK_MARKERS, allow_empty=True)
        editor.replace_block(["Год рождения"], "Год рождения:", data.birth, RVK_MARKERS, allow_empty=True)
        editor.replace_block(["Проживает"], "Проживает:", data.registered, RVK_MARKERS, allow_empty=True)
        rvk_work_position = data.rvk_work_position or ", ".join(part for part in [data.work_org, data.position] if part).strip(", ") or "не работает"
        editor.replace_block(["Место работы"], "Место работы:", rvk_work_position, RVK_MARKERS, allow_empty=True)
        period = f"Находился на обследовании в ГБУЗНО «Психиатрическая больница №2» г. Н. Новгорода с {data.admission_date} по {data.discharge_date}".strip()
        editor.replace_first_matching_paragraph(["Находился на обследовании"], period)
        military_area = format_military_commissariat_area(data.rvk_military_commissariat)
        if military_area:
            editor.replace_first_matching_paragraph(
                ["Госпитализируется по направлению военного комиссариата"],
                f"Госпитализируется по направлению военного комиссариата {military_area}."
            )
        if data.psych_account:
            editor.replace_first_matching_paragraph(["На учёте", "На учете"], f"На учёте у психиатров {data.psych_account}.")
        editor.replace_block(["Жалобы"], "Жалобы:", data.complaints or "не предъявляет", RVK_MARKERS, allow_empty=True)
        editor.replace_block(["Анамнез жизни"], "Анамнез жизни:", data.life_anamnesis, RVK_MARKERS)
        editor.replace_block(["Анамнез заболевания"], "Анамнез заболевания:", data.disease_anamnesis, RVK_MARKERS)
        editor.replace_block(["Психический статус"], "Психический статус:", data.mental_status, RVK_MARKERS)
        editor.replace_block(["Сомато-неврологический статус", "Соматический статус"], "Сомато-неврологический статус:", data.somatic_status, RVK_MARKERS)
        self._replace_lab_lines(editor, dates)
        if data.epi_text:
            editor.replace_block(["ЭПИ"], "ЭПИ -", data.epi_text, RVK_MARKERS)
        else:
            editor.remove_all_matching_paragraphs(["ЭПИ"])
        # В шаблоне Акта РВК после блока ЭПИ/ЭЭГ есть служебная одиночная строка "ЭЭГ".
        # Она не относится к результату исследования и должна исчезать из итогового документа.
        remove_exact_paragraphs(doc, ["ЭЭГ", "ЭПИ"])
        editor.replace_first_matching_paragraph(["Диагноз"], f"Диагноз: {data.diagnosis}")
        finalize_medical_document(doc, data)
        doc.save(str(output_path))

    @staticmethod
    def _replace_lab_lines(editor: DocxBlockEditor, dates: Dict[str, str]) -> None:
        replacements = [
            (["ОАК"], f"ОАК - в норме - {dates['day1']}"),
            (["ОАМ"], f"ОАМ - в норме - {dates['day1']}"),
            (["RW"], f"RW - в норме - {dates['day1']}"),
            (["HCV"], f"HCV - в норме - {dates['day1']}"),
            (["HBsAg"], f"HBsAg - в норме - {dates['day1']}"),
            (["ВИЧ"], f"ВИЧ - в норме - {dates['day2']}"),
            (["Биохимия крови"], f"Биохимия крови - в норме - {dates['day1']}"),
            (["Глюкоза крови"], f"Глюкоза крови - 3,40 ммоль/л - {dates['day1']}"),
            (["Кал на яйца глист"], f"Кал на яйца глист - не обнаружены - {dates['day1']}"),
            (["Флюорография"], f"Флюорография - патологии не выявлено - {dates['flg']}"),
            (["ЭКГ"], f"ЭКГ - ритм синусовый, ЧСС 65 ударов в минуту, рисунок ЭКГ в пределах нормы, ЭОС нормальная - {dates['day1']}"),
        ]
        for markers, text in replacements:
            editor.replace_first_matching_paragraph(markers, text)


# -----------------------------------------------------------------------------
# Сервисный слой
# -----------------------------------------------------------------------------

class MedicalDocumentService:
    def __init__(self):
        self.parser = MedicalTextParser()
        self.renderer = MedicalDocumentRenderer()

    def parse_primary_document(self, path: str | Path) -> PatientData:
        """Прочитать входной первичный документ пациента.

Поддерживаются оба рабочих источника данных:
- направление на госпитализацию;
- уже заполненный первичный осмотр.

Оба документа приводятся к единой PatientData, из которой затем создаются
все отмеченные в UI документы.
        """
        return self.parser.parse_docx(path)

    def parse_navigation(self, path: str | Path) -> PatientData:
        # Совместимость со старыми вызовами: раньше входной документ назывался
        # "направление". Теперь это общий первичный документ.
        return self.parse_primary_document(path)

    def load_epi_text(self, path: str | Path) -> str:
        if not path:
            return ""
        path = Path(path)
        if path.suffix.lower() == ".docx":
            text = extract_docx_text(path)
        else:
            text = path.read_text(encoding="utf-8", errors="ignore").strip()
        return strip_leading_epi_label(text)

    def available_templates(self) -> Dict[str, Path]:
        return {kind: bundled_template_path(kind) for kind in DOCUMENT_ORDER}

    def missing_templates(self) -> List[Path]:
        return [path for path in self.available_templates().values() if not path.exists()]

    def create_documents(
        self,
        *,
        navigation_path: str | Path,
        output_dir: str | Path,
        discharge_date: str = "",
        epi_path: str | Path | None = None,
        selected_docs: Sequence[str] = DOCUMENT_ORDER,
        override_data: Optional[PatientData] = None,
    ) -> Tuple[List[Path], PatientData]:
        missing = [bundled_template_path(kind) for kind in selected_docs if not bundled_template_path(kind).exists()]
        if missing:
            missing_text = "\n".join(str(path) for path in missing)
            raise FileNotFoundError(f"Не найдены встроенные шаблоны:\n{missing_text}")

        data = override_data or self.parse_primary_document(navigation_path)
        data.discharge_date = discharge_date.strip() or data.discharge_date
        if epi_path:
            data.epi_text = self.load_epi_text(epi_path)

        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        stem = safe_filename(data.output_fio or data.fio or Path(navigation_path).stem)

        created: List[Path] = []
        for kind in selected_docs:
            template_path = bundled_template_path(kind)
            suffix = OUTPUT_SUFFIXES[kind]
            output_path = available_path(output_dir / f"{stem}_{suffix}.docx")
            self.renderer.render(kind, template_path, output_path, data)
            created.append(output_path)
        return created, data


# -----------------------------------------------------------------------------
# Вспомогательные функции
# -----------------------------------------------------------------------------

def parse_date(value: str) -> Optional[datetime]:
    value = (value or "").strip()
    if not value:
        return None
    try:
        return datetime.strptime(value, DATE_FMT)
    except ValueError:
        return None


def safe_filename(value: str) -> str:
    value = normalize_text(value) or "Пациент"
    value = re.sub(r"[\\/:*?\"<>|]+", "_", value)
    value = re.sub(r"\s+", "_", value)
    return value.strip("._ ")[:80] or "Пациент"


def available_path(path: Path) -> Path:
    if not path.exists():
        return path
    base = path.with_suffix("")
    ext = path.suffix
    counter = 2
    while True:
        candidate = Path(f"{base}_{counter}{ext}")
        if not candidate.exists():
            return candidate
        counter += 1


def strip_leading_epi_label(text: str) -> str:
    """Убирает дублирующую метку из ЭПИ-файла: «ЭПИ: ...» -> «...»."""
    text = normalize_text(text)
    text = re.sub(r"^эпи\s*[:()№N#.-]*\s*", "", text, flags=re.IGNORECASE)
    return text.strip()


def format_preview(data: PatientData) -> str:
    lines = [
        "📋 Извлечённые данные",
        f"Тип первичного документа: {data.input_document_kind or '—'}",
        f"Ф.И.О.: {data.fio or '—'}",
        f"История болезни №: {data.case_number or '—'}",
        f"Год/дата рождения: {data.birth or '—'}",
        f"Адрес: {data.registered or '—'}",
        f"Дата госпитализации: {data.admission_date or '—'}",
        f"Дата выписки: {data.discharge_date or '—'}",
        f"Поступает: {data.admission or '—'}",
        f"Диагноз: {data.diagnosis or '—'}",
        f"План лечения: {shorten(data.treatment_plan, 140) or '—'}",
        f"ЭПИ: {len(data.epi_text)} символов",
        f"Акт РВК №: {data.rvk_act_number or '—'}",
        f"Военкомат РВК: {data.rvk_military_commissariat or '—'}",
        f"Работа/должность РВК: {data.rvk_work_position or '—'}",
        f"ВК на МСЭ: дата {data.vk_date or '—'}, протокол № {data.vk_protocol_number or '—'}, от {data.vk_protocol_date or '—'}",
        f"ВК больничный: дата {data.sick_leave_vk_date or '—'}, протокол № {data.sick_leave_vk_protocol_number or '—'}, от {data.sick_leave_vk_protocol_date or '—'}, комиссия {data.sick_leave_vk_commission_date or '—'}",
        f"Жалобы: {shorten(data.complaints, 180) or '—'}",
        f"Анамнез жизни: {len(data.life_anamnesis)} символов",
        f"Анамнез заболевания: {len(data.disease_anamnesis)} символов",
        f"Психический статус: {len(data.mental_status)} символов",
    ]
    if data.warnings:
        lines.append("")
        lines.append("⚠️ Предупреждения:")
        lines.extend(f"• {w}" for w in data.warnings)
    return "\n".join(lines)


def shorten(text: str, limit: int) -> str:
    text = normalize_text(text)
    if len(text) <= limit:
        return text
    return text[: limit - 1].rstrip() + "…"


# -----------------------------------------------------------------------------
# UI
# -----------------------------------------------------------------------------

class MedicalApp:
    def __init__(self, root: tk.Tk):
        self.root = root
        self.service = MedicalDocumentService()
        self.data = PatientData()

        self.navigation_path = tk.StringVar()
        self.epi_path = tk.StringVar()
        self.output_dir = tk.StringVar()
        self.discharge_date = tk.StringVar(value=datetime.now().strftime(DATE_FMT))
        self.strict_mode = tk.BooleanVar(value=True)
        self.document_vars: Dict[str, tk.BooleanVar] = {
            kind: tk.BooleanVar(value=True) for kind in DOCUMENT_ORDER
        }

        self.root.title(APP_TITLE)
        self.root.geometry("1060x720")
        self.root.minsize(960, 640)
        self._build_ui()
        self._check_bundled_templates()

    def _build_ui(self) -> None:
        self._setup_style()

        outer = ttk.Frame(self.root, padding=14)
        outer.pack(fill="both", expand=True)

        title = ttk.Label(outer, text="🏥 Автозаполнение медицинских документов", style="Title.TLabel")
        title.pack(anchor="w", pady=(0, 10))

        subtitle = ttk.Label(
            outer,
            text="Выберите первичный документ пациента, при необходимости ЭПИ, укажите дату выписки и создайте комплект документов. Шаблоны встроены в программу.",
            style="Muted.TLabel",
        )
        subtitle.pack(anchor="w", pady=(0, 12))

        file_box = ttk.LabelFrame(outer, text="Данные пациента", padding=12)
        file_box.pack(fill="x", pady=(0, 12))

        self._file_row(file_box, "1. Первичный документ", self.navigation_path, self.choose_navigation, [("Word DOCX", "*.docx")])
        self._file_row(file_box, "2. ЭПИ файл / текстовый DOCX", self.epi_path, self.choose_epi, [("Word DOCX", "*.docx"), ("Text", "*.txt"), ("All files", "*.*")], optional=True)

        date_box = ttk.Frame(file_box)
        date_box.pack(fill="x", pady=(8, 0))
        ttk.Label(date_box, text="Дата выписки:", width=34).pack(side="left")
        ttk.Entry(date_box, textvariable=self.discharge_date, width=16).pack(side="left")
        ttk.Label(date_box, text="формат: дд.мм.гггг", style="Muted.TLabel").pack(side="left", padx=10)
        ttk.Checkbutton(date_box, text="строгая проверка ФИО/даты/рождения", variable=self.strict_mode).pack(side="right")

        documents_box = ttk.LabelFrame(outer, text="Какие документы создать", padding=12)
        documents_box.pack(fill="x", pady=(0, 12))
        ttk.Label(
            documents_box,
            text="Отметьте галочками только те документы, которые нужны сейчас. Шаблоны встроены в программу, верхние шапки сохраняются.",
            style="Muted.TLabel",
        ).pack(anchor="w", pady=(0, 8))

        checklist = ttk.Frame(documents_box)
        checklist.pack(fill="x")
        for col, kind in enumerate(DOCUMENT_ORDER):
            cb = ttk.Checkbutton(checklist, text=DOCUMENT_LABELS[kind], variable=self.document_vars[kind])
            cb.grid(row=0, column=col, sticky="w", padx=(0, 18), pady=2)

        actions = ttk.Frame(outer)
        actions.pack(fill="x", pady=(0, 12))

        self.main_button = ttk.Button(actions, text="🚀 Создать выбранные документы", command=self.process_files, style="Accent.TButton")
        self.main_button.pack(side="left", ipady=8, ipadx=14)

        self.progress = ttk.Progressbar(outer, mode="indeterminate")
        self.progress.pack(fill="x", pady=(0, 12))

        log_box = ttk.LabelFrame(outer, text="Проверка и журнал", padding=10)
        log_box.pack(fill="both", expand=True)
        self.log = scrolledtext.ScrolledText(log_box, wrap=tk.WORD, font=("Consolas", 10))
        self.log.pack(fill="both", expand=True)

        self._write_log("Готово. Выберите первичный документ пациента, отметьте нужные документы галочками и нажмите создание. Шаблоны выбирать не нужно — они встроены в программу.\n")

    def _setup_style(self) -> None:
        style = ttk.Style()
        try:
            style.theme_use("clam")
        except tk.TclError:
            pass
        style.configure("Title.TLabel", font=("Segoe UI", 16, "bold"))
        style.configure("Muted.TLabel", foreground="#64748b")
        style.configure("Accent.TButton", font=("Segoe UI", 11, "bold"))

    def _file_row(self, parent: ttk.Frame, label: str, variable: tk.StringVar, command, filetypes, *, optional: bool = False) -> None:
        row = ttk.Frame(parent)
        row.pack(fill="x", pady=4)
        suffix = " (необязательно)" if optional else ""
        ttk.Label(row, text=label + suffix, width=34).pack(side="left")
        ttk.Entry(row, textvariable=variable).pack(side="left", fill="x", expand=True, padx=(0, 8))
        ttk.Button(row, text="Выбрать", command=command).pack(side="right")

    def _check_bundled_templates(self) -> None:
        missing = self.service.missing_templates()
        if missing:
            self._write_log("\n❌ Не найдены встроенные шаблоны:\n")
            for path in missing:
                self._write_log(f"- {path}\n")
            self._write_log("\nПроверьте папку templates рядом с main.py или сборку EXE с --add-data.\n")
        else:
            self._write_log("\n✅ Встроенные шаблоны найдены:\n")
            for kind, path in self.service.available_templates().items():
                self._write_log(f"- {OUTPUT_SUFFIXES[kind]}: {path.name}\n")

    def choose_navigation(self) -> None:
        path = filedialog.askopenfilename(title="Выберите первичный документ", filetypes=[("Word DOCX", "*.docx")])
        if not path:
            return
        self.navigation_path.set(path)
        self.output_dir.set(str(Path(path).parent))
        self.reparse_navigation()

    def choose_epi(self) -> None:
        path = filedialog.askopenfilename(title="Выберите файл ЭПИ", filetypes=[("Word DOCX", "*.docx"), ("Text", "*.txt"), ("All files", "*.*")])
        if path:
            self.epi_path.set(path)
            self.reparse_navigation(silent=True)

    def reparse_navigation(self, *, silent: bool = False) -> None:
        path = self.navigation_path.get().strip()
        if not path:
            if not silent:
                messagebox.showwarning("Нет файла", "Сначала выберите первичный документ.")
            return
        try:
            self.data = self.service.parse_navigation(path)
            self.data.discharge_date = self.discharge_date.get().strip()
            if self.epi_path.get().strip() and Path(self.epi_path.get().strip()).exists():
                self.data.epi_text = self.service.load_epi_text(self.epi_path.get().strip())
            self._write_log("\n" + "=" * 72 + "\n")
            self._write_log(format_preview(self.data) + "\n")
            self._write_log("=" * 72 + "\n")
        except Exception as exc:
            self._show_error("Не удалось прочитать первичный документ", exc)

    def process_files(self) -> None:
        navigation = self.navigation_path.get().strip()
        out_dir = self.output_dir.get().strip() or (str(Path(navigation).parent) if navigation else "")
        discharge_date = self.discharge_date.get().strip()
        selected_docs = self.selected_document_kinds()

        if not navigation or not Path(navigation).exists():
            messagebox.showwarning("Нет первичного документа", "Выберите первичный документ: направление на госпитализацию или первичный осмотр.")
            return
        if not selected_docs:
            messagebox.showwarning("Нет документов", "Отметьте галочками хотя бы один документ для создания.")
            return
        if discharge_date and not parse_date(discharge_date):
            messagebox.showwarning("Неверная дата", "Дата выписки должна быть в формате дд.мм.гггг.")
            return
        missing_templates = [bundled_template_path(kind) for kind in selected_docs if not bundled_template_path(kind).exists()]
        if missing_templates:
            msg = "Не найдены встроенные шаблоны выбранных документов:\n\n" + "\n".join(str(path) for path in missing_templates)
            messagebox.showerror("Нет шаблонов", msg)
            self._write_log("\n❌ " + msg + "\n")
            return

        try:
            self.progress.start()
            self.root.update_idletasks()
            data = self.service.parse_navigation(navigation)
            data.discharge_date = discharge_date
            if self.epi_path.get().strip():
                data.epi_text = self.service.load_epi_text(self.epi_path.get().strip())

            missing = data.missing_critical_fields()
            if missing:
                msg = "Не найдены критические поля:\n\n" + "\n".join(f"• {m}" for m in missing)
                if self.strict_mode.get():
                    messagebox.showerror("Создание остановлено", msg + "\n\nПроверьте, что выбран заполненный первичный документ пациента, а не пустой шаблон.")
                    self._write_log("\n❌ Создание остановлено: " + ", ".join(missing) + "\n")
                    return
                proceed = messagebox.askyesno("Есть пропуски", msg + "\n\nПродолжить всё равно?")
                if not proceed:
                    return

            created, used_data = self.service.create_documents(
                navigation_path=navigation,
                output_dir=out_dir,
                discharge_date=discharge_date,
                epi_path=self.epi_path.get().strip() or None,
                selected_docs=selected_docs,
                override_data=data,
            )

            selected_names = ", ".join(DOCUMENT_LABELS[kind] for kind in selected_docs)
            self._write_log(f"\n✅ Созданы выбранные документы: {selected_names}\n")
            for idx, path in enumerate(created, start=1):
                self._write_log(f"{idx}. {path}\n")
            if used_data.missing_recommended_fields():
                self._write_log("\n⚠️ Есть незаполненные рекомендуемые поля: " + ", ".join(used_data.missing_recommended_fields()) + "\n")
            messagebox.showinfo("Готово", "Документы созданы:\n\n" + "\n".join(path.name for path in created))
        except Exception as exc:
            self._show_error("Не удалось создать документы", exc)
        finally:
            self.progress.stop()

    def selected_document_kinds(self) -> List[str]:
        return [kind for kind in DOCUMENT_ORDER if self.document_vars[kind].get()]

    def _write_log(self, text: str) -> None:
        self.log.insert(tk.END, text)
        self.log.see(tk.END)

    def _show_error(self, title: str, exc: Exception) -> None:
        details = "".join(traceback.format_exception(type(exc), exc, exc.__traceback__))
        self._write_log(f"\n❌ {title}: {exc}\n{details}\n")
        messagebox.showerror(title, str(exc))


def main() -> None:
    root = tk.Tk()
    MedicalApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
